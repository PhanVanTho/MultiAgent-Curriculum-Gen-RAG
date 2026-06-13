# -*- coding: utf-8 -*-
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- IMPORT V21.5 ENTERPRISE CLEANER ---
from dich_vu.xuat_tai_lieu.bo_loc_html import clean_for_docx as _clean_text_docx

def _set_font(run, size=13, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Buộc font EastAsia cũng là Times New Roman (quan trọng cho tiếng Việt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def _sanitize_xml_text(text):
    """Loại bỏ NULL bytes và control characters không hợp lệ trong XML."""
    if not text:
        return text
    # Loại bỏ NULL bytes
    text = text.replace('\x00', '')
    # Loại bỏ control characters (0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F) nhưng giữ \t \n \r
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text

def _add_paragraph(doc, text, style=None, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    if not text:
        return None
    
    # Làm sạch markdown kỹ hơn (dùng bộ lọc V21.5)
    text = _clean_text_docx(text)
    
    # Loại bỏ NULL bytes và control characters không hợp lệ XML
    text = _sanitize_xml_text(text)
    
    text = text.strip()
    
    if not text:
        return None

    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    
    # Thụt đầu dòng 1.27cm (chuẩn giáo trình) NẾU không phải là heading hay list
    if style is None or style == "Normal":
        p.paragraph_format.first_line_indent = Inches(0.5)

    run = p.add_run(text)
    _set_font(run, size=13, bold=bold, italic=italic)
    return p

def xuat_docx(ket_qua: dict, duong_dan_docx: str):
    """
    Xuất nội dung sách ra file DOCX chuẩn định dạng giáo trình.
    """
    book = ket_qua.get("book_vi", {})
    doc = Document()

    # --- THIẾT LẬP LỀ CHUẨN GIÁO TRÌNH (BGDĐT) ---
    # Giấy A4: 21cm x 29.7cm. Lề: Trên/Dưới 2cm, Trái 3cm (đóng gáy), Phải 2cm.
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    # --- STYLE SETUP (Optional hack if styles don't exist) ---
    # (Docx mặc định đã có Normal, Heading 1...)

    from docx.oxml import OxmlElement
    
    # --- CÀI ĐẶT HEADER / FOOTER ---
    for section in doc.sections:
        # Footer: Số trang
        footer = section.footer
        p_footer = footer.paragraphs[0]
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_f = p_footer.add_run("Trang ")
        _set_font(run_f, size=11, italic=True)
        
        run_page = p_footer.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)
        run_page._r.append(fldChar3)
        
        # Header: Tên giáo trình
        title = book.get("title", "GIÁO TRÌNH ĐẠI HỌC")
        header = section.header
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_h = p_header.add_run(title.upper())
        _set_font(run_h, size=10, italic=True)
        
        # Đường viền dưới Header
        p_pr = p_header._element.get_or_add_pPr()
        p_pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_pbdr.append(bottom)
        p_pr.append(p_pbdr)

    # 1. TRANG BÌA
    for _ in range(6): doc.add_paragraph()
    
    # Tên Giáo Trình
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("GIÁO TRÌNH\n" + title.upper())
    _set_font(run_title, size=26, bold=True)
    
    for _ in range(8): doc.add_paragraph()
    
    # Tác giả & Năm
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Tác giả: Hệ thống AI Biên soạn tự động\nNăm xuất bản: 2026")
    _set_font(run_sub, size=14)
    
    doc.add_page_break()

    # 1.5 MỤC LỤC TỰ ĐỘNG (Word TOC Field — có số trang, tự cập nhật)
    # Dùng field TOC chuẩn của Word thay vì dựng thủ công
    from docx.oxml import OxmlElement

    # Tiêu đề "MỤC LỤC"
    h_toc = doc.add_heading("MỤC LỤC", level=1)
    h_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h_toc.runs:
        _set_font(run, size=16, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _insert_toc(doc):
        """Chèn trường TOC tự động của Word (Heading 1–3, có số trang, có leader '...')."""
        # Tạo paragraph chứa TOC field
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after  = Pt(4)
        run = paragraph.add_run()

        # Bắt đầu field
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        fld_begin.set(qn('w:dirty'), 'true')   # Báo Word cần cập nhật khi mở

        # Lệnh TOC: quét Heading 1-3, có tab leader chấm, có số trang bên phải
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = r' TOC \o "1-3" \h \z \u '

        # Kết thúc field
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    _insert_toc(doc)

    # Ghi chú nhỏ bên dưới mục lục
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(8)
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_note = p_note.add_run("[ Mở file trong Microsoft Word → nhấn Ctrl+A → F9 để cập nhật số trang ]")
    _set_font(run_note, size=10, italic=True)
    run_note.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


    # 2. Glossary (V33: Bảng thuật ngữ có định nghĩa)
    glossary = ket_qua.get("glossary", [])
    terms = ket_qua.get("terms", [])
    if glossary:
        h1 = doc.add_heading("BẢNG THUẬT NGỮ", level=1)
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in glossary:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run_term = p.add_run(f"{item.get('term', '')}: ")
            _set_font(run_term, size=13, bold=True)
            run_def = p.add_run(item.get('definition', ''))
            _set_font(run_def, size=13)
        doc.add_page_break()
    elif terms:
        # Fallback: danh sách thuật ngữ cũ (không có định nghĩa)
        h1 = doc.add_heading("DANH MỤC THUẬT NGỮ", level=1)
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, t in enumerate(terms[:100], 1):
            term_text = f"{i}. {t.get('term','')} - {t.get('meaning','')}"
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(term_text)
            _set_font(run, size=13)
        doc.add_page_break()

    # 3. Chapters
    chapters = book.get("chapters", [])
    for idx, ch in enumerate(chapters, 1):
        # Chapter Heading
        h1_text = f"CHƯƠNG {idx}: {ch.get('title','').upper()}"
        h1 = doc.add_heading(h1_text, level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
            _set_font(run, size=16, bold=True)
            # Màu đen cho chuyên nghiệp (mặc định xanh)
            run.font.color.rgb = RGBColor(0, 0, 0) 
        
        # V33: Tóm tắt chương (Đã ẩn theo yêu cầu)
        # summary = ch.get("summary", "")
        # if summary:
        #     p_sum = doc.add_paragraph()
        #     p_sum.paragraph_format.left_indent = Inches(0.3)
        #     p_sum.paragraph_format.right_indent = Inches(0.3)
        #     p_sum.paragraph_format.space_after = Pt(12)
        #     run_label = p_sum.add_run("Tóm tắt chương: ")
        #     _set_font(run_label, size=12, bold=True, italic=True)
        #     run_text = p_sum.add_run(summary)
        #     _set_font(run_text, size=12, italic=True)

        # Sections
        for jdx, sec in enumerate(ch.get("sections", []), 1):
            # Section Heading
            h2_text = f"{idx}.{jdx}. {sec.get('title','')}"
            h2 = doc.add_heading(h2_text, level=2)
            for run in h2.runs:
                _set_font(run, size=14, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

            # --- Lấy danh sách refs và map ID ---
            refs_list = ket_qua.get("references", [])
            ref_dict = {str(r.get("id")): r for r in refs_list if isinstance(r, dict)}
            
            # Content
            content = sec.get("content", "")
            if content:
                sub_idx = 1
                # Xử lý từng dòng để tránh lỗi gộp đoạn và ép Justify sai
                lines = content.split('\n')
                buffer_para = []
                
                def _process_docx_text(clean_para, style="Normal"):
                    # V41.3: Lọc BỎ TẤT CẢ thẻ HTML
                    clean_para = re.sub(r'<[^>]+>', '', clean_para)
                    # Loại bỏ markdown thừa
                    clean_para = clean_para.replace('**', '').replace('`', '')
                    # Xử lý Links
                    clean_para = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1 (\2)', clean_para)
                    
                    # Convert [ID] to APA Inline: (Title, Year)
                    def apa_replacer(match):
                        id_str = match.group(1)
                        ref = ref_dict.get(id_str)
                        if ref:
                            title_short = ref.get('title', 'Tài liệu')
                            if len(title_short) > 30: title_short = title_short[:30] + "..."
                            year = ref.get('year', 2026)
                            return f"({title_short}, {year})"
                        return match.group(0)
                    clean_para = re.sub(r'\[(\d+)\]', apa_replacer, clean_para)
                    
                    _add_paragraph(doc, clean_para, style=style)
                
                def flush_buffer():
                    if buffer_para:
                        text = " ".join(buffer_para).strip()
                        if text:
                            _process_docx_text(text, style="Normal")
                        buffer_para.clear()

                for raw_line in lines:
                    line = raw_line.strip()
                    if not line:
                        flush_buffer()
                        continue
                        
                    if line == "---":
                        flush_buffer()
                        p = doc.add_paragraph()
                        p_pr = p._element.get_or_add_pPr()
                        p_pbdr = OxmlElement('w:pBdr')
                        bottom = OxmlElement('w:bottom')
                        bottom.set(qn('w:val'), 'single')
                        bottom.set(qn('w:sz'), '6')
                        bottom.set(qn('w:space'), '1')
                        bottom.set(qn('w:color'), 'auto')
                        p_pbdr.append(bottom)
                        p_pr.append(p_pbdr)
                    elif line.startswith("### "):
                        flush_buffer()
                        text_clean = line[4:].strip()
                        text_clean = f"{idx}.{jdx}.{sub_idx}. {text_clean}"
                        sub_idx += 1
                        h3 = doc.add_heading(text_clean, level=3)
                        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in h3.runs:
                            _set_font(run, size=13, bold=True)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif line.startswith("#### "):
                        flush_buffer()
                        text_clean = line[5:].strip()
                        h4 = doc.add_heading(text_clean, level=4)
                        h4.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in h4.runs:
                            _set_font(run, size=13, bold=True, italic=True)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif line in ["Câu hỏi Ôn tập", "Bài tập & Ôn tập", "Review Questions"] or line.startswith("**Câu hỏi Ôn tập") or line.startswith("**Review Questions"):
                        flush_buffer()
                        clean_title = line.replace('**', '').replace('###', '').strip()
                        h3 = doc.add_heading(clean_title, level=3)
                        h3.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in h3.runs:
                            _set_font(run, size=13, bold=True, italic=True)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                    elif line.startswith("- ") or line.startswith("* "):
                        flush_buffer()
                        _process_docx_text(line[2:].strip(), style="List Bullet")
                    elif re.match(r'^\d+\.\s', line):
                        flush_buffer()
                        _process_docx_text(re.sub(r'^\d+\.\s', '', line).strip(), style="List Number")
                    else:
                        buffer_para.append(line)
                        
                # Flush cuối cùng nếu còn
                flush_buffer()
        
        doc.add_page_break()

    # 4. Global References
    refs = ket_qua.get("references", [])
    if refs:
        h1 = doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h1.runs:
             _set_font(run, size=16, bold=True)
             run.font.color.rgb = RGBColor(0, 0, 0)
             
        # Sort Alphabetically by Title (APA Standard)
        try:
            refs = sorted(refs, key=lambda x: str(x.get("title", "")).lower() if isinstance(x, dict) else str(x).lower())
        except:
            pass

        for u in refs:
            p = doc.add_paragraph(style="Normal")
            # APA hanging indent (thụt lề dòng thứ 2)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            
            if isinstance(u, dict):
                # APA Format: Title. (Year, Date). In Wikipedia, The Free Encyclopedia. URL
                ref_title = u.get('title', 'Nguồn')
                ref_url = u.get('url', '')
                ref_year = u.get('year', 2026)
                ref_date = u.get('access_date', '')
                
                # Title + năm (Không đánh số thứ tự 1., 2.)
                date_str = f"{ref_year}, {ref_date}" if ref_date else str(ref_year)
                run_body = p.add_run(f"{ref_title}. ({date_str}). In ")
                _set_font(run_body, size=12)
                
                # Wikipedia italic
                run_wiki = p.add_run("Wikipedia, The Free Encyclopedia. ")
                _set_font(run_wiki, size=12, italic=True)
                
                # URL
                run_url = p.add_run(ref_url)
                _set_font(run_url, size=12)
            else:
                run = p.add_run(f"• {u}")
                _set_font(run, size=12)

    doc.save(duong_dan_docx)
