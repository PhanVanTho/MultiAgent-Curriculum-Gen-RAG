# -*- coding: utf-8 -*-
import os
import json
import re
import uuid
import traceback
import logging
import time
import random
import zipfile
import io
from functools import wraps
from datetime import datetime, timezone
from concurrent.futures import ThreadPoolExecutor
import threading
import math

# --- ADAPTIVE MULTI-PORT DATABASE ORCHESTRATOR ---
if __name__ == "__main__":
    import sys
    import socket
    import subprocess

    def is_port_open(host, port):
        for h in [host, "127.0.0.1", "::1"]:
            try:
                with socket.create_connection((h, port), timeout=0.5) as conn:
                    return True
            except OSError:
                continue
        return False

    if not os.getenv("GTAI_SUBPROCESS"):
        mongo_active = is_port_open("localhost", 27017)
        mysql_active = is_port_open("localhost", 3306)

        if mongo_active and mysql_active:
            print("\n" + "="*80)
            print("🚀 PHÁT HIỆN CẢ MONGODB VÀ MYSQL ĐANG BẬT!")
            print("Hệ thống sẽ chạy song song 2 luồng độc lập:")
            print(" - Luồng MongoDB: http://localhost:5000")
            print(" - Luồng MySQL:   http://localhost:5001")
            print("="*80 + "\n")
            
            env_mongo = os.environ.copy()
            env_mongo["GTAI_OVERRIDE_DB_TYPE"] = "mongodb"
            env_mongo["GTAI_OVERRIDE_USE_SQLITE"] = "False"
            env_mongo["GTAI_OVERRIDE_PORT"] = "5000"
            env_mongo["GTAI_SUBPROCESS"] = "True"
            
            env_mysql = os.environ.copy()
            env_mysql["GTAI_OVERRIDE_DB_TYPE"] = "mysql"
            env_mysql["GTAI_OVERRIDE_USE_SQLITE"] = "False"
            env_mysql["GTAI_OVERRIDE_PORT"] = "5001"
            env_mysql["GTAI_SUBPROCESS"] = "True"
            
            p_mongo = subprocess.Popen([sys.executable, __file__], env=env_mongo)
            p_mysql = subprocess.Popen([sys.executable, __file__], env=env_mysql)
            
            try:
                while True:
                    time.sleep(1)
                    if p_mongo.poll() is not None or p_mysql.poll() is not None:
                        break
            except KeyboardInterrupt:
                print("\n🛑 Đang dừng các ứng dụng...")
                p_mongo.terminate()
                p_mysql.terminate()
                p_mongo.wait()
                p_mysql.wait()
            sys.exit(0)

        elif mongo_active:
            print("\n" + "="*80)
            print("🚀 PHÁT HIỆN CHỈ MONGODB ĐANG BẬT!")
            print("Hệ thống chạy với cơ sở dữ liệu: MongoDB tại http://localhost:5000")
            print("="*80 + "\n")
            os.environ["GTAI_OVERRIDE_DB_TYPE"] = "mongodb"
            os.environ["GTAI_OVERRIDE_USE_SQLITE"] = "False"
            os.environ["GTAI_OVERRIDE_PORT"] = "5000"
            os.environ["GTAI_SUBPROCESS"] = "True"

        elif mysql_active:
            print("\n" + "="*80)
            print("🚀 PHÁT HIỆN CHỈ MYSQL (XAMPP) ĐANG BẬT!")
            print("Hệ thống chạy với cơ sở dữ liệu: MySQL tại http://localhost:5000")
            print("="*80 + "\n")
            os.environ["GTAI_OVERRIDE_DB_TYPE"] = "mysql"
            os.environ["GTAI_OVERRIDE_USE_SQLITE"] = "False"
            os.environ["GTAI_OVERRIDE_PORT"] = "5000"
            os.environ["GTAI_SUBPROCESS"] = "True"

        else:
            print("\n" + "="*80)
            print("⚠️ KHÔNG PHÁT HIỆN MONGODB HOẶC MYSQL ĐANG BẬT!")
            print("Hệ thống tự động sử dụng SQLite ngoại tuyến tại http://localhost:5000")
            print("="*80 + "\n")
            os.environ["GTAI_OVERRIDE_USE_SQLITE"] = "True"
            os.environ["GTAI_OVERRIDE_PORT"] = "5000"
            os.environ["GTAI_SUBPROCESS"] = "True"


try:
    import numpy as _np
except ImportError:
    _np = None

def _json_safe_default(obj):
    """Convert numpy and datetime types to native Python types for JSON serialization."""
    if isinstance(obj, datetime):
        return obj.isoformat()
    if _np is not None:
        if isinstance(obj, _np.bool_):
            return bool(obj)
        if isinstance(obj, _np.integer):
            return int(obj)
        if isinstance(obj, _np.floating):
            return float(obj)
        if isinstance(obj, _np.ndarray):
            return obj.tolist()
    raise TypeError(f"Object of type {type(obj).__name__} is not JSON serializable")

def is_valid_query(query):
    # Cho phép chữ, số, khoảng trắng, tiếng Việt, dấu trừ (hyphen, en-dash, em-dash), hai chấm, dấu phẩy, dấu chấm, ngoặc đơn, gạch chéo
    pattern = r"^[\w\sÀ-ỹ\-–—:,\.\(\)/]+$"
    return re.match(pattern, query) is not None

def is_gibberish(text):
    text = text.lower().strip()
    words = text.split()
    
    vowels = set("aeiouyàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ")
    consonants = set("bcdfghjklmnpqrstvwxzđ")
    
    for word in words:
        # Check consecutive consonants
        consec_consonants = 0
        max_consec = 0
        for char in word:
            if char in consonants:
                consec_consonants += 1
                if consec_consonants > max_consec:
                    max_consec = consec_consonants
            else:
                consec_consonants = 0
        
        if max_consec >= 6:
            return True
            
        if len(word) > 5:
            word_vowels = sum(1 for char in word if char in vowels)
            if word_vowels == 0:
                return True
            if len(word) > 8 and (word_vowels / len(word)) < 0.15:
                return True
                
    return False

def is_meaningful(query):
    # Cho phép 1 từ nhưng phải cấu thành từ ít nhất 2 ký tự (ví dụ: AI, IT)
    if len(query.strip()) < 2:
        return False
    return not is_gibberish(query)

def kiem_tra_viet_tat_llm(term):
    """
    Gọi GPT-4o-mini để kiểm tra xem một từ viết tắt/thuật ngữ viết tắt
    có nghĩa và hợp lệ trong bất kỳ ngữ cảnh khoa học, công nghệ, y học, kinh tế, giáo dục, nghệ thuật, xã hội... hay không.
    Trả về True nếu từ viết tắt đó HỢP LỆ (có nghĩa rõ ràng và thông dụng).
    Trả về False nếu từ viết tắt đó VÔ NGHĨA hoặc KHÔNG HỢP LỆ (cần yêu cầu người dùng viết rõ).
    """
    from openai import OpenAI
    from cau_hinh import CauHinh
    
    if not CauHinh.OPENAI_API_KEY:
        # Nếu không có API Key, mặc định cho qua để không chặn người dùng
        return True
        
    try:
        client = OpenAI(api_key=CauHinh.OPENAI_API_KEY, max_retries=1)
        prompt = (
            "Bạn là một chuyên gia ngôn ngữ học và biên tập giáo trình học thuật.\n"
            "Hãy kiểm tra xem từ/cụm từ viết tắt sau đây có phải là một từ viết tắt có nghĩa, hợp lệ và thông dụng trong bất kỳ lĩnh vực nào (Khoa học, Công nghệ, Y tế, Kinh tế, Giáo dục, Nghệ thuật, Đời sống xã hội, v.v.) hay không.\n"
            f"Từ viết tắt cần kiểm tra: \"{term}\"\n\n"
            "Chỉ trả về kết quả dưới định dạng JSON với cấu trúc chính xác như sau:\n"
            "{\n"
            "  \"hop_le\": true hoặc false,\n"
            "  \"giai_thich\": \"giải thích ngắn gọn về nghĩa của từ viết tắt nếu hợp lệ\"\n"
            "}\n"
            "Lưu ý: Chỉ trả về chuỗi JSON thô, không kèm ký tự markdown ```json hay giải thích gì thêm ngoài JSON."
        )
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that evaluates abbreviations and returns raw JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0,
            max_tokens=150
        )
        content = response.choices[0].message.content.strip()
        if "```" in content:
            content = content.replace("```json", "").replace("```", "").strip()
            
        import json
        res = json.loads(content)
        return bool(res.get("hop_le", False))
    except Exception as e:
        logger.error(f"Lỗi kiểm tra viết tắt LLM cho '{term}': {e}")
        # Nếu lỗi API, mặc định cho qua (trả về True) để tránh làm nghẽn/xung đột hệ thống
        return True

def is_abbreviation(text):
    trimmed = text.strip()
    if not trimmed:
        return False
        
    clean_lower = re.sub(r'[^a-zđàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ\s]', '', trimmed.lower())
    
    allowed_common_abbrs = {
        "ai", "iot", "it", "cntt", "csdl", "html", "css", "sql", "dna", "rna", "gdp", 
        "vpn", "wifi", "ip", "cpu", "ram", "rom", "io", "api", "oop", "http", "https", 
        "tcp", "udp", "dns", "lan", "wan", "xml", "nosql", "db", "mvc", "crud", "sdk", 
        "ide", "gui", "cli", "gpu", "ssd", "hdd", "url", "uri", "seo", "rag", "ssl", 
        "tls", "dom", "nlp", "llm", "gpts", "gpt", "json", "rest", "graphql", "aws", 
        "cloud", "saas", "paas", "iaas", "ci", "cd", "cicd", "docker", "k8s", "vm", 
        "npm", "pip", "git", "ssh", "ftp", "smtp", "pop3", "imap", "dhcp",
        "man", "pan", "wlan", "vlan", "nat", "firewall", "ids", 
        "ips", "siem", "soc", "noc", "ddos", "dos", "malware", "virus", "worm", "trojan",
        "ransomware", "phishing", "spam", "adware", "spyware", "rootkit", "exploit", 
        "payload", "shellcode", "backdoor", "honeypot", "cryptography", "encryption", 
        "decryption", "hash", "md5", "sha", "aes", "des", "rsa", "ecc", "dh", "pkc", 
        "pki", "ca", "sftp", "ftps", "ipsec", "l2tp", 
        "pptp", "gre", "openvpn", "wireguard", "sdwan", "mpls", "bgp", "ospf", "rip", 
        "eigrp", "isis", "stp", "vtp", "lacp", "pagp", "hsrp", "vrrp", "glbp", "arp", 
        "icmp", "igmp", "snmp", "ntp", "syslog", "radius", "tacacs", "ldap", "ad", 
        "tftp", "scp", "rsync", "telnet", "soap", "rpc", "grpc", "websocket",
        "ajax", "yaml", "ini", "csv", "tsv", "pdf", "docx", "xlsx", 
        "pptx", "txt", "rtf", "odt", "ods", "odp", "xhtml", "js", "ts", 
        "python", "java", "cpp", "csharp", "php", "ruby", "go", "rust", "swift", "kotlin", 
        "scala", "r", "matlab", "mongodb", "mysql", "postgresql", "sqlite",
        "oracle", "redis", "memcached", "cassandra", "hbase", "neo4j", "elasticsearch",
        "solr", "lucence", "hadoop", "spark", "hive", "pig", "impala", "presto", "druid",
        "flink", "storm", "kafka", "rabbitmq", "activemq", "zeromq", "mqtt", "amqp", 
        "coap", "xmpp", "http2", "http3", "quic", "webrtc", "sip", "rtp", "rtsp", 
        "hls", "dash", "rtmp", "ffmpeg", "opencv", "tensorflow", "pytorch", 
        "keras", "scikit-learn", "numpy", "pandas", "scipy", "matplotlib", "seaborn", 
        "plotly", "bokeh", "jupyter", "colab", "kaggle",
        "fdi", "wto", "who", "asean", "nsnd", "opec", "imf", "wb", "sem", "b2b", "b2c",
        "kpi", "okr", "roi", "hr", "pr", "qa", "qc", "saas", "bi", "erp", "crm", "cms",
        "lms", "pos", "rfid", "gps", "gis", "cad", "cam", "cae"
    }
    
    if clean_lower in allowed_common_abbrs:
        return False
        
    if re.match(r'^[A-ZĐÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝĂĨŨƠƯ]{2,6}$', trimmed):
        if trimmed.lower() in allowed_common_abbrs:
            return False
        if kiem_tra_viet_tat_llm(trimmed):
            return False
        return True
        
    words = trimmed.split()
    if 1 < len(words) <= 4:
        all_uppercase_and_short = True
        for w in words:
            clean_word = re.sub(r'[^a-zA-ZĐÀÁÂÃÈÉÊÌÍÒÓÔÕÙÝĂĨŨƠƯ]', '', w)
            if len(clean_word) > 0:
                if not re.match(r'^[A-ZĐÀÁÂÃÈÉÊÌÍÒÓÔÕÙÝĂĨŨƠƯ]+$', clean_word) or len(clean_word) > 5:
                    all_uppercase_and_short = False
                    break
        if all_uppercase_and_short:
            has_unallowed = False
            for w in words:
                clean_w = re.sub(r'[^a-zđàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]', '', w.lower())
                if len(clean_w) >= 2 and clean_w not in allowed_common_abbrs:
                    if not kiem_tra_viet_tat_llm(w):
                        has_unallowed = True
                        break
            if has_unallowed:
                return True
                
    vowels = set("aeiouyàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ")
    for w in words:
        clean_word = re.sub(r'[^a-zđàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹ]', '', w.lower())
        if not clean_word:
            continue
        if clean_word in allowed_common_abbrs:
            continue
            
        has_vowel = any(char in vowels for char in clean_word)
        if not has_vowel and len(clean_word) >= 2:
            if not kiem_tra_viet_tat_llm(w):
                return True
                
    return False
# --- BỘ HÃM XUNG TOÀN CỤC (GLOBAL THROTTLING V23.1) ---
OPENAI_SEMAPHORE = threading.BoundedSemaphore(6)
TERMS_LOCK = threading.Lock()
GEMINI_SEMAPHORE = threading.BoundedSemaphore(1)

# Lock cho dữ liệu Knowledge Base và Terms (V17.2, V39)
PASSAGES_LOCK = threading.RLock()
TERMS_LOCK = threading.RLock()
MAX_TOTAL_PASSAGES = 2000 

# --- BỘ HÃM XUNG GEMINI (V23 TURBO - 15 RPM SAFE) ---
GEMINI_LOCK = threading.Lock()
LAST_GEMINI_CALL = {"time": 0}

def gemini_throttled_call(func, *args, **kwargs):
    """Bộ điều tiết toàn cục V22.1 - 15 RPM với Jitter tránh xung đột pha."""
    with GEMINI_LOCK:
        # 4.0s base = 15 RPM. Thêm jitter để tránh đồng bộ hóa giữa các thread.
        jitter = random.uniform(0, 0.5)
        last_call = LAST_GEMINI_CALL.get("time", 0)
        wait_time = max(0, 4.0 + jitter - (time.time() - last_call))
        if wait_time > 0:
            time.sleep(wait_time)
        try:
            res = func(*args, **kwargs)
            return res
        finally:
            LAST_GEMINI_CALL["time"] = time.time()

from flask import Flask, request, render_template, jsonify, send_file, url_for, flash, redirect, session
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from dotenv import load_dotenv

# --- NẠP CẤU HÌNH & DB ---
load_dotenv(override=True)
if os.getenv("GTAI_SUBPROCESS") == "True":
    if os.getenv("GTAI_OVERRIDE_DB_TYPE"):
        os.environ["DB_TYPE"] = os.getenv("GTAI_OVERRIDE_DB_TYPE")
    if os.getenv("GTAI_OVERRIDE_USE_SQLITE"):
        os.environ["DB_USE_SQLITE"] = os.getenv("GTAI_OVERRIDE_USE_SQLITE")
    if os.getenv("GTAI_OVERRIDE_PORT"):
        os.environ["PORT"] = os.getenv("GTAI_OVERRIDE_PORT")
from cau_hinh import CauHinh
from mo_hinh import db, NguoiDung, LichSuGiaoTrinh, XacThucOTP, GoiCuoc, TrangThongTin


# --- SERVICE IMPORTS (V23.2 GLOBAL STABILIZATION) ---
from dich_vu.vector_search import tim_kiem_vector, tao_vector_db, tim_kiem_vector_with_llm_rerank
from dich_vu.openai_da_buoc import (
    tao_dan_y as openai_tao_dan_y, 
    viet_noi_dung_chuong as openai_writer, 
    viet_noi_dung_muc,
    viet_noi_dung_batch_sections
)
from dich_vu.gemini_da_buoc import (
    gemini_fix_json, 
    viet_noi_dung_muc_gemini,
    tao_dan_y as gemini_tao_dan_y,
    viet_noi_dung_chuong as gemini_writer
)
from dich_vu.gemini_giam_sat import (
    giam_sat_chuong, 
    giam_sat_outline, 
    giam_sat_quy_mo
)
from dich_vu.xuat_tai_lieu.xuat_docx import xuat_docx
from dich_vu.xuat_tai_lieu.xuat_pdf import xuat_pdf
from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering, fallback_raw_facts

# --- LOGGING ---
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(), logging.FileHandler("app.log", encoding="utf-8")]
)
logger = logging.getLogger(__name__)

# --- FLASK SETUP ---
from dich_vu.xuat_tai_lieu.markdown_parser import parse_markdown
app = Flask(__name__, template_folder="templates", static_folder="static")
app.jinja_env.filters['markdown'] = parse_markdown

def gmt7_filter(dt):
    if not dt:
        return None
    from datetime import timedelta
    return dt + timedelta(hours=7)
app.jinja_env.filters['gmt7'] = gmt7_filter

app.config["SECRET_KEY"] = CauHinh.KHOA_BI_MAT
from mo_hinh import db_type
if db_type == "mongodb":
    app.config["SQLALCHEMY_DATABASE_URI"] = "mongodb://localhost:27017"
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
elif db_type == "sqlite" or os.getenv("DB_USE_SQLITE") == "True":
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///giao_trinh_ai.db"
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
else:
    db_uri = f"mysql+pymysql://{os.getenv('DB_USER', 'root')}:{os.getenv('DB_PASS', '')}@{os.getenv('DB_HOST', 'localhost')}:{os.getenv('DB_PORT', '3306')}/{os.getenv('DB_NAME', 'giao_trinh_ai')}"
    app.config["SQLALCHEMY_DATABASE_URI"] = db_uri
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
        "pool_recycle": 280,
        "pool_pre_ping": True,
    }

db.init_app(app)
login_manager = LoginManager()
login_manager.login_view = "login"
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    # --- SQLAlchemy 2.0 Syntax (V24.6 Fix) ---
    user = db.session.get(NguoiDung, int(user_id))
    if user and user.bi_khoa:
        return None
    return user

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.la_admin:
            flash("Bạn không có quyền truy cập trang này.", "danger")
            return redirect(url_for("trang_chu"))
        return f(*args, **kwargs)
    return decorated_function

# Demo store (In-memory)
CONG_VIEC = {}

def seed_system_config():
    try:
        from mo_hinh import CauHinhHeThong
        from cau_hinh import ma_hoa_key
        if CauHinhHeThong.query.count() == 0:
            logger.info("Đồng bộ cấu hình từ .env vào CSDL lần đầu...")
            keys_to_sync = [
                "OPENAI_API_KEY", "OPENAI_MODEL", 
                "GEMINI_API_KEYS", "GEMINI_MODEL", "GEMINI_MODEL_LITE",
                "WRITER_MODEL", "SEARCH_MODEL", "SUPERVISOR_MODEL_LITE", "SUPERVISOR_MODEL_PRO",
                "PAYMENT_VNPAY_ACTIVE", "PAYMENT_SEPAY_ACTIVE",
                "VNPAY_TMN_CODE", "VNPAY_HASH_SECRET", "VNPAY_PAYMENT_URL", "VNPAY_RETURN_URL",
                "SEPAY_API_KEY", "SEPAY_ACCOUNT_NUMBER", "SEPAY_BANK_BRAND", "SEPAY_WEB_NAME", "SEPAY_XOR_KEY",
                "PHI_TOKEN_AUTO", "PHI_TOKEN_EXPERT", "PHI_TOKEN_CREATIVE"
            ]
            api_keys_to_encrypt = ["OPENAI_API_KEY", "GEMINI_API_KEYS", "SEPAY_API_KEY", "VNPAY_HASH_SECRET"]
            
            for key in keys_to_sync:
                val = os.getenv(key)
                if val is not None:
                    db_val = ma_hoa_key(val) if key in api_keys_to_encrypt else str(val)
                    item = CauHinhHeThong(khoa=key, gia_tri=db_val)
                    db.session.add(item)
            db.session.commit()
            logger.info("Đồng bộ cấu hình lần đầu thành công.")
        
        # Đồng bộ riêng các khoá cấu hình thông tin liên hệ mới nếu chưa có trong DB
        contact_keys = {
            "CONTACT_EMAIL": "phanvantho082019@gmail.com",
            "CONTACT_PHONE": "0327152710",
            "CONTACT_ADDRESS_VI": "Cần Thơ, Việt Nam",
            "CONTACT_ADDRESS_EN": "Can Tho, Vietnam",
            "ADMIN_NOTIFICATION_EMAIL": "phanvantho082019@gmail.com",
            "MAC_DINH_SO_CHUONG_MAX": "15",
            "MAC_DINH_SO_TU_MAX": "1000",
            "GOOGLE_CLIENT_ID": ""
        }
        for k, default_val in contact_keys.items():
            if not CauHinhHeThong.query.filter_by(khoa=k).first():
                val = os.getenv(k, default_val)
                item = CauHinhHeThong(khoa=k, gia_tri=str(val))
                db.session.add(item)
        db.session.commit()
    except Exception as e:
        logger.error(f"Lỗi đồng bộ cấu hình từ .env vào CSDL: {e}")


def seed_trang_thong_tin():
    try:
        from mo_hinh import TrangThongTin
        if TrangThongTin.query.count() == 0:
            logger.info("Seeding default pages into TrangThongTin...")
            
            def extract_lang_content(file_path):
                if not os.path.exists(file_path):
                    return "", ""
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                
                vi_start_marker = '<div class="lang-vi">'
                vi_end_marker = '<!-- ENGLISH -->'
                
                en_start_marker = '<div class="lang-en d-none">'
                en_end_marker = '</article>'
                
                vi_html = ""
                en_html = ""
                
                vi_idx_start = content.find(vi_start_marker)
                if vi_idx_start != -1:
                    vi_idx_start += len(vi_start_marker)
                    vi_idx_end = content.find(vi_end_marker, vi_idx_start)
                    if vi_idx_end != -1:
                        vi_raw = content[vi_idx_start:vi_idx_end].strip()
                        if vi_raw.endswith('</div>'):
                            vi_raw = vi_raw[:-6].strip()
                        vi_html = vi_raw

                en_idx_start = content.find(en_start_marker)
                if en_idx_start != -1:
                    en_idx_start += len(en_start_marker)
                    en_idx_end = content.find(en_end_marker, en_idx_start)
                    if en_idx_end != -1:
                        en_raw = content[en_idx_start:en_idx_end].strip()
                        if en_raw.endswith('</div>'):
                            en_raw = en_raw[:-6].strip()
                        en_html = en_raw
                        
                return vi_html, en_html

            ai_vi_default = """<h2>1. Sở hữu trí tuệ đối với Giáo trình</h2>
<p>Hệ thống <strong>Giáo Trình AI</strong> biên soạn nội dung dựa trên các thuật toán học sâu và cơ sở dữ liệu mở. Bản quyền đối với cấu trúc giáo trình và nội dung biên soạn sau khi được xuất bản thuộc về người dùng đã yêu cầu biên soạn. Người dùng có toàn quyền sử dụng, in ấn, giảng dạy và thương mại hóa sản phẩm giáo trình của mình.</p>

<h2>2. Trách nhiệm về mặt chuyên môn</h2>
<p>Nội dung do AI tạo ra chỉ mang tính chất tham khảo học thuật. Hệ thống không chịu trách nhiệm pháp lý đối với bất kỳ sai sót kiến thức, lỗi logic, hoặc hậu quả phát sinh từ việc áp dụng trực tiếp các kiến thức trong giáo trình vào thực tế giảng dạy mà không qua kiểm duyệt chuyên môn bởi giảng viên hoặc chuyên gia có thẩm quyền.</p>

<h2>3. Nghiêm cấm các hành vi lạm dụng</h2>
<p>Người dùng không được phép:</p>
<ul>
    <li>Sử dụng hệ thống để biên soạn các nội dung vi phạm pháp luật, kích động bạo lực, hoặc vi phạm bản quyền sở hữu trí tuệ của bên thứ ba.</li>
    <li>Tấn công hoặc can thiệp vào hạ tầng máy chủ của Giáo Trình AI bằng các công cụ tự động hoặc bot độc hại.</li>
    <li>Lợi dụng các lỗ hổng hệ thống để trục lợi token hoặc can thiệp vào tài khoản của người dùng khác.</li>
</ul>"""

            ai_en_default = """<h2>1. Intellectual Property Rights</h2>
<p>The <strong>AI Curriculum</strong> compiles content based on deep learning models and open databases. The ownership of the compiled curriculum and outline, once generated and downloaded, belongs entirely to the requesting user. Users retain full rights to use, print, teach, and commercialize their generated curriculum materials.</p>

<h2>2. Disclaimer of Professional Liability</h2>
<p>AI-generated content is intended solely for academic reference and drafting assistance. The system and its developers assume no legal responsibility or liability for any errors, logical inconsistencies, or negative consequences arising from using the generated content in actual teaching environments without prior expert human review.</p>

<h2>3. Prohibited Activities & Abuse</h2>
<p>Users are strictly prohibited from:</p>
<ul>
    <li>Using the system to generate unlawful, hateful, violent, or copyright-infringing content.</li>
    <li>Attacking or interfering with the system infrastructure of AI Curriculum via automated tools or scrapers.</li>
    <li>Exploiting system vulnerabilities for token manipulation or unauthorized access to other user accounts.</li>
</ul>"""

            pages_data = [
                {
                    "ma_trang": "privacy-policy",
                    "tieu_de_vi": "Chính Sách Bảo Mật",
                    "tieu_de_en": "Privacy Policy",
                    "mo_ta_vi": "Cách chúng tôi bảo vệ và xử lý thông tin cá nhân của bạn",
                    "mo_ta_en": "How we collect, use, and safeguard your personal information",
                    "file_name": "templates/privacy_policy.html"
                },
                {
                    "ma_trang": "terms-of-service",
                    "tieu_de_vi": "Điều Khoản Dịch Vụ",
                    "tieu_de_en": "Terms of Service",
                    "mo_ta_vi": "Quy định và điều kiện sử dụng hệ thống Giáo Trình AI",
                    "mo_ta_en": "Terms and conditions for utilizing the AI Curriculum system",
                    "file_name": "templates/terms_of_service.html"
                },
                {
                    "ma_trang": "data-deletion",
                    "tieu_de_vi": "Yêu Cầu Xóa Dữ Liệu",
                    "tieu_de_en": "Data Deletion Request",
                    "mo_ta_vi": "Hướng dẫn và chính sách xóa dữ liệu cá nhân khỏi hệ thống",
                    "mo_ta_en": "Instructions and policies for removing your personal data from the system",
                    "file_name": "templates/data_deletion.html"
                },
                {
                    "ma_trang": "ai-terms",
                    "tieu_de_vi": "Điều khoản AI & Sở hữu trí tuệ",
                    "tieu_de_en": "AI Terms & Intellectual Property",
                    "mo_ta_vi": "Quy định về việc sử dụng nội dung do AI tạo ra và trách nhiệm pháp lý",
                    "mo_ta_en": "Regulations on AI-generated content usage and legal liabilities",
                    "file_name": None
                }
            ]
            
            for p in pages_data:
                vi_html, en_html = "", ""
                if p["file_name"]:
                    try:
                        vi_html, en_html = extract_lang_content(p["file_name"])
                    except Exception as ex:
                        logger.error(f"Error reading file {p['file_name']}: {ex}")
                
                if not vi_html:
                    if p["ma_trang"] == "ai-terms":
                        vi_html = ai_vi_default
                    else:
                        vi_html = f"<p>Nội dung đang được cập nhật cho {p['tieu_de_vi']}.</p>"
                if not en_html:
                    if p["ma_trang"] == "ai-terms":
                        en_html = ai_en_default
                    else:
                        en_html = f"<p>Content is being updated for {p['tieu_de_en']}.</p>"
                
                page_obj = TrangThongTin(
                    ma_trang=p["ma_trang"],
                    tieu_de_vi=p["tieu_de_vi"],
                    tieu_de_en=p["tieu_de_en"],
                    mo_ta_vi=p["mo_ta_vi"],
                    mo_ta_en=p["mo_ta_en"],
                    noi_dung_vi=vi_html,
                    noi_dung_en=en_html
                )
                db.session.add(page_obj)
            db.session.commit()
            logger.info("Successfully seeded default pages into TrangThongTin.")
    except Exception as e:
        logger.error(f"Error seeding TrangThongTin: {e}")

def seed_admin():
    from mo_hinh import db_type
    if db_type == "mongodb":
        with app.app_context():
            db.create_all()
            seed_trang_thong_tin()
            admin = NguoiDung.query.filter_by(ten_dang_nhap="admin").first()

            if not admin:
                hashed_pw = generate_password_hash("admin123")
                new_admin = NguoiDung(ten_dang_nhap="admin", mat_khau=hashed_pw, la_admin=True, email="admin@local")
                db.session.add(new_admin)
                db.session.commit()
                logger.info("Default admin seed complete in MongoDB.")
            try:
                if GoiCuoc.query.count() == 0:
                    logger.info("Seeding default GoiCuoc packages in MongoDB...")
                    default_packages = [
                        GoiCuoc(ten_goi="Trải nghiệm", gia_tien=10000, so_token=10, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7"),
                        GoiCuoc(ten_goi="Khởi đầu", gia_tien=20000, so_token=30, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7"),
                        GoiCuoc(ten_goi="Cơ bản", gia_tien=60000, so_token=100, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7")
                    ]
                    for p in default_packages:
                        db.session.add(p)
                    db.session.commit()
                    logger.info("Default GoiCuoc seed complete.")
            except Exception as e:
                logger.error(f"Error seeding default GoiCuoc: {e}")
            seed_system_config()
        return

    from sqlalchemy import inspect
    with app.app_context():
        # Auto-migration for existing tables
        try:
            inspector = inspect(db.engine)
            
            # 1. Migrate 'lich_su_giao_trinh' columns
            if 'lich_su_giao_trinh' in inspector.get_table_names():
                columns_info = inspector.get_columns('lich_su_giao_trinh')
                cols = [c['name'] for c in columns_info]
                
                # Ensure nguoi_dung_id is nullable (allows guest curriculum)
                nguoi_dung_id_col = next((c for c in columns_info if c['name'] == 'nguoi_dung_id'), None)
                if nguoi_dung_id_col and not nguoi_dung_id_col.get('nullable', True):
                    logger.info("Column 'nguoi_dung_id' is NOT NULL in 'lich_su_giao_trinh'. Modifying to allow NULL...")
                    if 'mysql' in str(db.engine.url):
                        db.session.execute(db.text("SET FOREIGN_KEY_CHECKS=0"))
                        db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh MODIFY COLUMN nguoi_dung_id INT(10) UNSIGNED NULL DEFAULT NULL"))
                        db.session.execute(db.text("SET FOREIGN_KEY_CHECKS=1"))
                        db.session.commit()

                if 'noi_bat' not in cols:
                    logger.info("Column 'noi_bat' is missing from 'lich_su_giao_trinh'. Adding...")
                    db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh ADD COLUMN noi_bat TINYINT(1) NOT NULL DEFAULT 0"))
                    db.session.commit()
                if 'do_dai_ky_tu' not in cols:
                    logger.info("Column 'do_dai_ky_tu' is missing from 'lich_su_giao_trinh'. Adding...")
                    db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh ADD COLUMN do_dai_ky_tu INT DEFAULT 0"))
                    db.session.commit()
                if 'da_xuat_file' not in cols:
                    logger.info("Column 'da_xuat_file' is missing from 'lich_su_giao_trinh'. Adding...")
                    db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh ADD COLUMN da_xuat_file TINYINT(1) DEFAULT 0"))
                    db.session.commit()
                if 'duong_dan_file' not in cols:
                    logger.info("Column 'duong_dan_file' is missing from 'lich_su_giao_trinh'. Adding...")
                    db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh ADD COLUMN duong_dan_file VARCHAR(255)"))
                    db.session.commit()
                if 'mysql' in str(db.engine.url):
                    logger.info("Ensuring 'noi_dung_html' column is LONGTEXT...")
                    db.session.execute(db.text("ALTER TABLE lich_su_giao_trinh MODIFY COLUMN noi_dung_html LONGTEXT"))
                    db.session.commit()

            # 2. Migrate 'nguoi_dung' columns
            if 'nguoi_dung' in inspector.get_table_names():
                user_cols = [c['name'] for c in inspector.get_columns('nguoi_dung')]
                if 'bi_khoa' not in user_cols:
                    logger.info("Column 'bi_khoa' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN bi_khoa TINYINT(1) NOT NULL DEFAULT 0"))
                    db.session.commit()
                if 'email' not in user_cols:
                    logger.info("Column 'email' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN email VARCHAR(150) NULL"))
                    db.session.commit()
                if 'google_id' not in user_cols:
                    logger.info("Column 'google_id' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN google_id VARCHAR(255) NULL"))
                    db.session.commit()
                if 'ho_ten' not in user_cols:
                    logger.info("Column 'ho_ten' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN ho_ten VARCHAR(255) NULL"))
                    db.session.commit()
                if 'anh_dai_dien' not in user_cols:
                    logger.info("Column 'anh_dai_dien' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN anh_dai_dien VARCHAR(500) NULL"))
                    db.session.commit()
                if 'token' not in user_cols:
                    logger.info("Column 'token' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN token INT DEFAULT 10"))
                    db.session.commit()
                if 'ngay_tao' not in user_cols:
                    logger.info("Column 'ngay_tao' is missing from 'nguoi_dung'. Adding...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung ADD COLUMN ngay_tao DATETIME DEFAULT '2026-06-15 00:00:00'"))
                    db.session.commit()
                if 'mysql' in str(db.engine.url):
                    logger.info("Ensuring mat_khau_hash is nullable in MySQL...")
                    db.session.execute(db.text("ALTER TABLE nguoi_dung MODIFY COLUMN mat_khau_hash VARCHAR(255) NULL DEFAULT NULL"))
                    db.session.commit()
        except Exception as e:
            logger.error(f"Error during auto-migration: {e}")
            db.session.rollback()

        db.create_all()
        seed_trang_thong_tin()
        admin = NguoiDung.query.filter_by(ten_dang_nhap="admin").first()
        if not admin:
            hashed_pw = generate_password_hash("admin123")
            new_admin = NguoiDung(ten_dang_nhap="admin", mat_khau=hashed_pw, la_admin=True, email="admin@local")
            db.session.add(new_admin)
            db.session.commit()
            logger.info("Default admin seed complete.")
        
        # Seed default GoiCuoc packages if empty
        try:
            if GoiCuoc.query.count() == 0:
                logger.info("Seeding default GoiCuoc packages...")
                default_packages = [
                    GoiCuoc(ten_goi="Trải nghiệm", gia_tien=10000, so_token=10, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7"),
                    GoiCuoc(ten_goi="Khởi đầu", gia_tien=20000, so_token=30, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7"),
                    GoiCuoc(ten_goi="Cơ bản", gia_tien=60000, so_token=100, mo_ta="Tạo giáo trình chuyên sâu, Tốc độ ưu tiên cao, Hỗ trợ 24/7")
                ]
                db.session.bulk_save_objects(default_packages)
                db.session.commit()
                logger.info("Default GoiCuoc seed complete.")
        except Exception as e:
            logger.error(f"Error seeding default GoiCuoc: {e}")
            
        # Seed dummy curriculum history if empty
        try:
            if LichSuGiaoTrinh.query.count() == 0:
                logger.info("Seeding dummy curriculum history...")
                admin = NguoiDung.query.filter_by(ten_dang_nhap="admin").first()
                admin_id = admin.id if admin else None
                dummy_curriculum = LichSuGiaoTrinh(
                    nguoi_dung_id=admin_id,
                    chu_de="Giáo trình Cơ học lượng tử cơ bản",
                    noi_dung_html="<h1>Cơ học lượng tử</h1><p>Đây là nội dung thử nghiệm.</p>",
                    duong_dan_file="du_lieu/pdf/dummy_uuid.pdf",
                    do_dai_ky_tu=100,
                    da_xuat_file=True
                )
                db.session.add(dummy_curriculum)
                db.session.commit()
                logger.info("Dummy curriculum seed complete.")
        except Exception as e:
            logger.error(f"Error seeding dummy curriculum: {e}")
        seed_system_config()

seed_admin()
os.makedirs(CauHinh.THU_MUC_JSON, exist_ok=True)
os.makedirs(CauHinh.THU_MUC_PDF, exist_ok=True)
os.makedirs(CauHinh.THU_MUC_DOCX, exist_ok=True)

def _luu_json(obj: dict, path: str):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)

def tao_html_hoc_thuat(tieu_de, book_export, references):
    """Tạo mã HTML tối giản (Academic style) cho nội dung giáo trình."""
    from datetime import datetime
    from dich_vu.xuat_tai_lieu.markdown_parser import parse_markdown
    
    html_parts = []
    html_parts.append('<div class="giao-trinh-container">')
    html_parts.append(f'<h1 style="font-size: 2.2rem; text-align: center; margin-bottom: 2rem;">GIÁO TRÌNH: {tieu_de}</h1>')
    html_parts.append(f'<div style="font-size: 1.1rem; margin-bottom: 3rem; text-align: center; color: #555;">'
                      f'<span>Biên soạn tự động bởi Hệ thống AI Data Aggregator</span><br>'
                      f'<span>Ngày tạo: {datetime.now().strftime("%d/%m/%Y")}</span>'
                      f'</div>')

    # Mục lục / TOC
    html_parts.append('<div class="muc-luc-khoi" style="margin-bottom: 3rem; border: 1px solid #ccc; padding: 20px; background-color: #fcfcfc;">')
    html_parts.append('<h2 style="font-size: 1.5rem; margin-top: 0; border-bottom: 1px solid #ddd; padding-bottom: 8px;">Mục Lục</h2>')
    html_parts.append('<ul style="list-style-type: none; padding-left: 0;">')
    for c_idx, chap in enumerate(book_export.get("chapters", []), 1):
        html_parts.append(f'<li style="margin-bottom: 10px;">')
        html_parts.append(f'<strong>Chương {c_idx}: {chap.get("title", "")}</strong>')
        html_parts.append(f'<ul style="list-style-type: none; padding-left: 20px;">')
        for s_idx, sec in enumerate(chap.get("sections", []), 1):
            html_parts.append(f'<li>{c_idx}.{s_idx}. {sec.get("title", "")}</li>')
        html_parts.append('</ul>')
        html_parts.append('</li>')
    html_parts.append('</ul>')
    html_parts.append('</div>')

    # Nội dung chính
    html_parts.append('<div class="noi-dung-chinh">')
    for c_idx, chap in enumerate(book_export.get("chapters", []), 1):
        html_parts.append('<div class="chuong-muc" style="margin-bottom: 3rem;">')
        html_parts.append(f'<h2 style="font-size: 1.8rem; border-bottom: 2px solid #333; padding-bottom: 10px; margin-top: 2rem;">Chương {c_idx}. {chap.get("title", "")}</h2>')
        
        if chap.get("summary"):
            html_parts.append(f'<div style="font-style: italic; background: #f8f9fa; border-left: 4px solid #198754; padding: 10px 15px; margin-bottom: 1.5rem; border-radius: 0 4px 4px 0;">'
                              f'<strong>Tóm tắt chương:</strong> {chap.get("summary")}'
                              f'</div>')
                              
        for s_idx, sec in enumerate(chap.get("sections", []), 1):
            html_parts.append('<div class="muc-con" style="margin-bottom: 2rem;">')
            html_parts.append(f'<h3 style="font-size: 1.4rem; color: #1a4a8a; margin-top: 1.5rem; margin-bottom: 10px;">{c_idx}.{s_idx}. {sec.get("title", "")}</h3>')
            
            # Parse section content to HTML
            sec_content_html = parse_markdown(sec.get("content", ""))
            html_parts.append(f'<div class="content" style="text-align: justify; line-height: 1.8;">{sec_content_html}</div>')
            html_parts.append('</div>')
        html_parts.append('</div>')
    html_parts.append('</div>')

    # Tài liệu tham khảo
    if references:
        html_parts.append('<div class="tai-lieu-tham-khao" style="margin-top: 4rem; border-top: 1px solid #ccc; padding-top: 2rem;">')
        html_parts.append('<h2 style="font-size: 1.6rem; color: #333;">Tài liệu tham khảo (APA)</h2>')
        html_parts.append('<div style="padding-left: 0; line-height: 1.6;">')
        for ref in sorted(references, key=lambda x: x.get('id', 0)):
            html_parts.append(f'<p style="text-indent: -2em; padding-left: 2em; margin-bottom: 0.5rem; font-size: 0.95rem; color: #555;">'
                              f'<span style="font-weight: bold; color: #1a4a8a;">[{ref.get("id")}]</span> '
                              f'{ref.get("title")}. ({ref.get("year", "n.d.")}{", " + ref.get("access_date") if ref.get("access_date") else ""}). '
                              f'In <em>Wikipedia, The Free Encyclopedia</em>. '
                              f'<a href="{ref.get("url")}" target="_blank" style="text-decoration: none; border-bottom: 1px solid #1a4a8a; color: #1a4a8a;">{ref.get("url")}</a>'
                              f'</p>')
        html_parts.append('</div>')
        html_parts.append('</div>')

    html_parts.append('</div>')
    return "\n".join(html_parts)

def _link_guest_curriculum(user):
    guest_ma_cv = session.pop("guest_ma_cv", None)
    if guest_ma_cv and guest_ma_cv in CONG_VIEC:
        thong_tin = CONG_VIEC[guest_ma_cv]
        if thong_tin.get("trang_thai") == "hoan_thanh":
            try:
                path_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{guest_ma_cv}.pdf")
                
                # Check if there is an existing history record for this CV
                ls = LichSuGiaoTrinh.query.filter_by(duong_dan_file=path_pdf).first()
                if ls:
                    # Just update the nguoi_dung_id to prevent duplicates
                    ls.nguoi_dung_id = user.id
                    db.session.commit()
                    logger.info(f"Updated existing guest job {guest_ma_cv} to user {user.id}")
                else:
                    # Parse JSON info and create new entry
                    tong_ky_tu = 0
                    book_data = {}
                    refs_data = []
                    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{guest_ma_cv}.json")
                    if os.path.exists(p_json):
                        try:
                            with open(p_json, 'r', encoding='utf-8') as f:
                                data = json.load(f)
                            book_data = data.get('book_vi', {})
                            refs_data = data.get('references', [])
                            for chap in book_data.get('chapters', []):
                                for sec in chap.get('sections', []):
                                    tong_ky_tu += len(sec.get('content', ''))
                        except Exception:
                            pass
                    
                    noi_dung_html = tao_html_hoc_thuat(thong_tin["tieu_de"], book_data, refs_data)
                    
                    ls = LichSuGiaoTrinh(
                        nguoi_dung_id=user.id,
                        chu_de=thong_tin["tieu_de"],
                        noi_dung_html=noi_dung_html,
                        duong_dan_file=path_pdf,
                        da_xuat_file=True,
                        do_dai_ky_tu=tong_ky_tu
                    )
                    db.session.add(ls)
                    db.session.commit()
                    logger.info(f"Created and linked guest job {guest_ma_cv} to user {user.id}")
            except Exception as e:
                logger.error(f"Link guest job error: {e}")
    return guest_ma_cv

# --- AUTH ROUTES ---
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        u = request.form.get("ten_dang_nhap"); p = request.form.get("mat_khau")
        user = NguoiDung.query.filter_by(ten_dang_nhap=u).first()
        if user and user.mat_khau and check_password_hash(user.mat_khau, p):
            if user.bi_khoa:
                flash("Tài khoản của bạn đã bị khóa bởi quản trị viên.", "danger")
                return redirect(url_for("login"))
            login_user(user); flash("Đăng nhập thành công!", "success")
            
            # Liên kết tài khoản Google chờ xử lý (nếu có)
            google_info = session.get("google_pending")
            if google_info and google_info.get("google_id"):
                user.google_id = google_info["google_id"]
                if not user.ho_ten:
                    user.ho_ten = google_info.get("ho_ten", "")
                if not user.anh_dai_dien:
                    user.anh_dai_dien = google_info.get("anh_dai_dien", "")
                db.session.commit()
                session.pop("google_pending", None)
                flash("Đã liên kết tài khoản Google thành công!", "success")
                logger.info(f"Auto-linked pending Google account to logged-in user: {user.ten_dang_nhap}")

            cv_linked = _link_guest_curriculum(user)
            if cv_linked: return redirect(url_for("ket_qua", ma_cv=cv_linked))
            return redirect(url_for("admin_dashboard") if user.la_admin else url_for("trang_chu"))
        flash("Tên đăng nhập hoặc mật khẩu không đúng.", "danger")
    return render_template("login.html", google_client_id=CauHinh.GOOGLE_CLIENT_ID)

# (Other routes: register, logout, lich-su... omitted for brevity or implementation consistency)
@app.route("/register", methods=["GET", "POST"])
def register():
    google_info = session.get("google_pending", {})
    
    if request.method == "POST":
        u = request.form.get("ten_dang_nhap"); p = request.form.get("mat_khau")
        email = (request.form.get("email") or "").strip().lower()
        if not email:
            flash("Vui lòng nhập địa chỉ email.", "danger"); return redirect(url_for("register"))
        if NguoiDung.query.filter_by(ten_dang_nhap=u).first():
            flash("Tên đăng nhập đã tồn tại.", "danger"); return redirect(url_for("register"))
        if NguoiDung.query.filter_by(email=email).first():
            flash("Email đã được sử dụng.", "danger"); return redirect(url_for("register"))
        
        # Sinh mã OTP ngẫu nhiên 6 chữ số
        otp_code = "".join([str(random.randint(0, 9)) for _ in range(6)])
        from datetime import timedelta
        het_han = datetime.utcnow() + timedelta(minutes=5)
        
        # Hủy kích hoạt các OTP cũ chưa sử dụng của email này
        XacThucOTP.query.filter_by(email=email, da_dung=False).update({"da_dung": True})
        
        # Lưu thông tin đăng ký tạm thời và mã OTP vào bảng xac_thuc_otp
        otp_record = XacThucOTP(
            email=email,
            ten_dang_nhap=u,
            mat_khau_hash=generate_password_hash(p),
            otp=otp_code,
            het_han=het_han,
            da_dung=False,
            google_id=google_info.get("google_id"),
            ho_ten=google_info.get("ho_ten", ""),
            anh_dai_dien=google_info.get("anh_dai_dien", "")
        )
        db.session.add(otp_record)
        db.session.commit()
        
        # Gửi OTP qua Email
        from dich_vu.email_service import gui_email_otp
        success, mail_msg = gui_email_otp(email, otp_code)
        
        if success:
            flash("Mã xác thực OTP đã được gửi tới email của bạn.", "success")
        else:
            flash(mail_msg, "warning")
            
        session["pending_register_email"] = email
        return redirect(url_for("verify_otp"))
        
    return render_template("register.html", google_client_id=CauHinh.GOOGLE_CLIENT_ID, google_info=google_info)

@app.route("/verify-otp", methods=["GET", "POST"])
def verify_otp():
    email = session.get("pending_register_email")
    if not email:
        flash("Vui lòng thực hiện đăng ký trước.", "danger")
        return redirect(url_for("register"))
        
    if request.method == "POST":
        entered_otp = request.form.get("otp", "").strip()
        if not entered_otp:
            flash("Vui lòng nhập mã OTP.", "danger")
            return render_template("verify_otp.html", email=email)
            
        # Tìm bản ghi OTP mới nhất chưa sử dụng của email này
        otp_record = XacThucOTP.query.filter_by(email=email, da_dung=False).order_by(XacThucOTP.id.desc()).first()
        
        if not otp_record:
            flash("Không tìm thấy thông tin xác thực OTP. Vui lòng đăng ký lại.", "danger")
            return redirect(url_for("register"))
            
        # Kiểm tra thời hạn hết hạn
        if datetime.utcnow() > otp_record.het_han:
            flash("Mã OTP đã hết hạn (hạn dùng 5 phút). Vui lòng đăng ký lại để nhận mã mới.", "danger")
            return redirect(url_for("register"))
            
        # So khớp OTP
        if otp_record.otp != entered_otp:
            flash("Mã OTP không chính xác. Vui lòng thử lại.", "danger")
            return render_template("verify_otp.html", email=email)
            
        # OTP chính xác -> Tạo tài khoản NguoiDung chính thức
        # Kiểm tra lại xem username/email có bị ai đăng ký trước trong lúc chờ không
        if NguoiDung.query.filter_by(ten_dang_nhap=otp_record.ten_dang_nhap).first():
            flash("Tên đăng nhập đã bị sử dụng bởi người khác trong lúc chờ xác thực.", "danger")
            return redirect(url_for("register"))
        if NguoiDung.query.filter_by(email=otp_record.email).first():
            flash("Email đã được đăng ký bởi người khác trong lúc chờ xác thực.", "danger")
            return redirect(url_for("register"))
            
        new_user = NguoiDung(
            ten_dang_nhap=otp_record.ten_dang_nhap,
            mat_khau=otp_record.mat_khau_hash,
            email=otp_record.email,
            google_id=otp_record.google_id,
            ho_ten=otp_record.ho_ten,
            anh_dai_dien=otp_record.anh_dai_dien
        )
        
        # Đánh dấu OTP đã dùng
        otp_record.da_dung = True
        db.session.add(new_user)
        db.session.commit()
        
        # Xóa các thông tin tạm trong session
        session.pop("google_pending", None)
        session.pop("pending_register_email", None)
        
        flash("Đăng ký tài khoản thành công! Vui lòng đăng nhập.", "success")
        return redirect(url_for("login"))
        
    return render_template("verify_otp.html", email=email)

@app.route("/resend-otp")
def resend_otp():
    email = session.get("pending_register_email")
    if not email:
        flash("Vui lòng thực hiện đăng ký trước.", "danger")
        return redirect(url_for("register"))
        
    # Lấy thông tin OTP chưa dùng gần nhất của email này để lấy username/mật khẩu
    otp_record = XacThucOTP.query.filter_by(email=email, da_dung=False).order_by(XacThucOTP.id.desc()).first()
    if not otp_record:
        flash("Không tìm thấy yêu cầu đăng ký trước đó. Vui lòng đăng ký lại.", "danger")
        return redirect(url_for("register"))
        
    # Tạo mã OTP mới
    otp_code = "".join([str(random.randint(0, 9)) for _ in range(6)])
    from datetime import timedelta
    het_han = datetime.utcnow() + timedelta(minutes=5)
    
    # Cập nhật bản ghi cũ hoặc tạo bản ghi mới
    otp_record.da_dung = True  # Vô hiệu hóa mã cũ
    
    new_otp_record = XacThucOTP(
        email=email,
        ten_dang_nhap=otp_record.ten_dang_nhap,
        mat_khau_hash=otp_record.mat_khau_hash,
        otp=otp_code,
        het_han=het_han,
        da_dung=False,
        google_id=otp_record.google_id,
        ho_ten=otp_record.ho_ten,
        anh_dai_dien=otp_record.anh_dai_dien
    )
    db.session.add(new_otp_record)
    db.session.commit()
    
    # Gửi lại email
    from dich_vu.email_service import gui_email_otp
    success, mail_msg = gui_email_otp(email, otp_code)
    if success:
        flash("Mã OTP mới đã được gửi lại tới email của bạn.", "success")
    else:
        flash(mail_msg, "warning")
        
    return redirect(url_for("verify_otp"))

@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if current_user.is_authenticated:
        return redirect(url_for("trang_chu"))
        
    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        if not email:
            flash("Vui lòng nhập địa chỉ email.", "danger")
            return redirect(url_for("forgot_password"))
            
        user = NguoiDung.query.filter_by(email=email).first()
        if not user:
            flash("Email không tồn tại trong hệ thống.", "danger")
            return redirect(url_for("forgot_password"))
            
        # Sinh mã OTP 6 chữ số
        otp_code = "".join([str(random.randint(0, 9)) for _ in range(6)])
        from datetime import timedelta
        het_han = datetime.utcnow() + timedelta(minutes=5)
        
        # Hủy OTP cũ của email này
        XacThucOTP.query.filter_by(email=email, da_dung=False).update({"da_dung": True})
        
        otp_record = XacThucOTP(
            email=email,
            ten_dang_nhap=user.ten_dang_nhap,
            mat_khau_hash="",
            otp=otp_code,
            het_han=het_han,
            da_dung=False,
            nguoi_dung_id=user.id
        )
        db.session.add(otp_record)
        db.session.commit()
        
        from dich_vu.email_service import gui_email_otp
        success, mail_msg = gui_email_otp(email, otp_code)
        
        if success:
            session["pending_reset_email"] = email
            flash("Mã xác thực OTP đã được gửi tới email của bạn. Vui lòng kiểm tra hộp thư.", "success")
            return redirect(url_for("reset_password"))
        else:
            flash(f"Không thể gửi email OTP: {mail_msg}", "warning")
            return redirect(url_for("forgot_password"))
            
    return render_template("forgot_password.html")

@app.route("/reset-password", methods=["GET", "POST"])
def reset_password():
    if current_user.is_authenticated:
        return redirect(url_for("trang_chu"))
        
    email = session.get("pending_reset_email")
    if not email:
        flash("Vui lòng gửi yêu cầu khôi phục mật khẩu trước.", "warning")
        return redirect(url_for("forgot_password"))
        
    if request.method == "POST":
        entered_otp = request.form.get("otp", "").strip()
        mat_khau_moi = request.form.get("mat_khau_moi")
        xac_nhan_mat_khau = request.form.get("xac_nhan_mat_khau_moi")
        
        if not entered_otp or not mat_khau_moi or not xac_nhan_mat_khau:
            flash("Vui lòng điền đầy đủ các thông tin.", "danger")
            return render_template("reset_password.html", email=email)
            
        if mat_khau_moi != xac_nhan_mat_khau:
            flash("Mật khẩu mới không trùng khớp.", "danger")
            return render_template("reset_password.html", email=email)
            
        otp_record = XacThucOTP.query.filter_by(email=email, da_dung=False).order_by(XacThucOTP.id.desc()).first()
        
        if not otp_record:
            flash("Không tìm thấy thông tin xác thực OTP. Vui lòng yêu cầu lại.", "danger")
            return redirect(url_for("forgot_password"))
            
        if datetime.utcnow() > otp_record.het_han:
            flash("Mã OTP đã hết hạn (hạn dùng 5 phút). Vui lòng thử lại.", "danger")
            return redirect(url_for("forgot_password"))
            
        if otp_record.otp != entered_otp:
            flash("Mã OTP không chính xác. Vui lòng thử lại.", "danger")
            return render_template("reset_password.html", email=email)
            
        # Cập nhật mật khẩu người dùng
        user = NguoiDung.query.filter_by(email=email).first()
        if not user:
            flash("Người dùng không còn tồn tại.", "danger")
            return redirect(url_for("forgot_password"))
            
        from werkzeug.security import generate_password_hash
        user.mat_khau = generate_password_hash(mat_khau_moi)
        otp_record.da_dung = True
        db.session.commit()
        
        session.pop("pending_reset_email", None)
        flash("Khôi phục mật khẩu thành công! Hãy đăng nhập bằng mật khẩu mới.", "success")
        return redirect(url_for("login"))
        
    return render_template("reset_password.html", email=email)

@app.route("/logout")
@login_required
def logout():
    logout_user(); flash("Đã đăng xuất.", "info"); return redirect(url_for("trang_chu"))

# --- GOOGLE SIGN-IN (GIS JavaScript Library) ---
@app.route("/auth/google", methods=["POST"])
def auth_google():
    """Xác thực Google Sign-In: nhận credential (ID token) từ GIS, xác minh và đăng nhập/đăng ký."""
    import requests as http_requests
    
    data = request.get_json(silent=True) or {}
    credential = data.get("credential", "")
    
    if not credential:
        return jsonify({"error": "Thiếu credential từ Google."}), 400
    
    # Xác minh ID token với Google
    try:
        verify_url = f"https://oauth2.googleapis.com/tokeninfo?id_token={credential}"
        resp = http_requests.get(verify_url, timeout=10)
        if resp.status_code != 200:
            logger.warning(f"Google token verification failed: {resp.status_code}")
            return jsonify({"error": "Token Google không hợp lệ."}), 401
        
        token_info = resp.json()
        
        # Kiểm tra audience (client_id) khớp
        if token_info.get("aud") != CauHinh.GOOGLE_CLIENT_ID:
            logger.warning(f"Google token audience mismatch: {token_info.get('aud')}")
            return jsonify({"error": "Token không hợp lệ cho ứng dụng này."}), 401
        
        google_id = token_info.get("sub")
        email = token_info.get("email", "")
        ho_ten = token_info.get("name", "")
        anh_dai_dien = token_info.get("picture", "")
        
        if not google_id:
            return jsonify({"error": "Không lấy được thông tin từ Google."}), 400
        
    except Exception as e:
        logger.error(f"Google auth verification error: {e}")
        return jsonify({"error": "Lỗi khi xác thực với Google."}), 500
    
    # Tìm hoặc tạo người dùng
    user = NguoiDung.query.filter_by(google_id=google_id).first()
    
    if not user and email:
        # Kiểm tra email đã tồn tại (người dùng đã đăng ký bằng form trước đó)
        user = NguoiDung.query.filter_by(email=email).first()
        if user:
            # Liên kết tài khoản Google với tài khoản hiện có
            user.google_id = google_id
            if not user.ho_ten:
                user.ho_ten = ho_ten
            if not user.anh_dai_dien:
                user.anh_dai_dien = anh_dai_dien
            db.session.commit()
            logger.info(f"Linked Google account to existing user: {user.ten_dang_nhap}")
    
    if not user:
        # Chưa có tài khoản → lưu thông tin Google vào session, chuyển đến trang đăng ký
        session["google_pending"] = {
            "google_id": google_id,
            "email": email,
            "ho_ten": ho_ten,
            "anh_dai_dien": anh_dai_dien
        }
        logger.info(f"Google user not registered, redirecting to register: {email}")
        return jsonify({
            "success": False,
            "need_register": True,
            "message": "Tài khoản Google chưa được đăng ký. Vui lòng đăng ký để tiếp tục.",
            "redirect": url_for("register")
        }), 200
    else:
        # Chỉ tự động điền thông tin nếu chưa được thiết lập trước đó
        if ho_ten and not user.ho_ten:
            user.ho_ten = ho_ten
        if anh_dai_dien and not user.anh_dai_dien:
            user.anh_dai_dien = anh_dai_dien
        db.session.commit()
    
    # Đăng nhập
    if user.bi_khoa:
        return jsonify({
            "success": False,
            "error": "Tài khoản của bạn đã bị khóa bởi quản trị viên."
        }), 403
        
    login_user(user)
    cv_linked = _link_guest_curriculum(user)
    
    redirect_url = url_for("trang_chu")
    if user.la_admin:
        redirect_url = url_for("admin_dashboard")
    if cv_linked:
        redirect_url = url_for("ket_qua", ma_cv=cv_linked)
    
    return jsonify({
        "success": True,
        "message": f"Xin chào, {ho_ten or user.ten_dang_nhap}!",
        "redirect": redirect_url
    })

@app.route("/lich-su")
@login_required
def lich_su():
    # 1. Lấy lịch sử hoàn thành từ DB
    completed = LichSuGiaoTrinh.query.filter_by(nguoi_dung_id=current_user.id).order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    
    items = []
    for item in completed:
        items.append({
            "loai": "completed",
            "id": item.id,
            "ma_cv": item.ma_cv,
            "chu_de": item.chu_de,
            "ngay_tao": item.ngay_tao or datetime.utcnow(),
            "do_dai_ky_tu": item.do_dai_ky_tu or 0,
            "trang_thai": "hoan_thanh"
        })
        
    # 2. Lấy lịch sử đang chạy / lỗi từ bộ nhớ CONG_VIEC
    for ma_cv, job in CONG_VIEC.items():
        if job.get("user_id") == current_user.id:
            # Bỏ qua nếu đã lưu trong DB
            if any(x["ma_cv"] == ma_cv for x in items):
                continue
                
            ngay_tao = job.get("ngay_tao", datetime.utcnow())
            items.append({
                "loai": "active",
                "id": None,
                "ma_cv": ma_cv,
                "chu_de": job.get("tieu_de", "Không rõ chủ đề"),
                "ngay_tao": ngay_tao,
                "do_dai_ky_tu": 0,
                "trang_thai": job.get("trang_thai", "dang_chay"),
                "tien_do": job.get("tien_do", 0),
                "buoc": job.get("buoc", "Đang xử lý"),
                "loi": job.get("loi", "")
            })
            
    # Sắp xếp theo thời gian mới nhất lên đầu
    items.sort(key=lambda x: x["ngay_tao"], reverse=True)
    return render_template("history.html", history=items)

@app.route("/api/chat", methods=["POST"])
@login_required
def api_chat():
    from flask import request, jsonify
    from flask_login import current_user
    from openai import OpenAI
    from cau_hinh import CauHinh
    from mo_hinh import LichSuChatbot, db

    try:
        data = request.get_json() or {}
        user_message_text = data.get("message")
        if not user_message_text:
            return jsonify({"error": "Message is required"}), 400

        # 1. Save user's message to DB
        user_msg = LichSuChatbot(
            nguoi_dung_id=current_user.id,
            role="user",
            content=user_message_text
        )
        db.session.add(user_msg)
        db.session.commit()

        # 2. Fetch full chat history from database for context
        history_records = LichSuChatbot.query.filter_by(nguoi_dung_id=current_user.id).order_by(LichSuChatbot.ngay_tao.asc()).all()

        # 3. Construct system instruction
        system_instruction = (
            "You are a virtual assistant supporting information for the AI Curriculum Compilation System (Giáo Trình AI / AI Curriculum System).\n"
            "You MUST ONLY answer queries directly related to this system (e.g., how to register/login, purchasing tokens, pricing packages, Auto/Expert/Creative compilation modes, featured product display, privacy policy, terms of service, or technical support regarding this software).\n\n"
            "FOR ANY OTHER QUESTIONS OR TOPICS outside this system (e.g., solving math, writing code, general homework, writing essays, cooking recipes, weather, news, general history/geography/science...), you MUST politely but firmly refuse to answer.\n"
            "Example refusal in Vietnamese: 'Tôi xin lỗi, nhưng tôi chỉ được phép trả lời các câu hỏi liên quan đến hệ thống Biên soạn Giáo trình AI.'\n"
            "Example refusal in English: 'I apologize, but I am only allowed to answer questions related to the AI Curriculum Compilation System.'\n\n"
            "Language policy:\n"
            "- If the user queries in Vietnamese, respond in friendly, polite, and concise Vietnamese.\n"
            "- If the user queries in English, respond in friendly, polite, and concise English."
        )

        openai_messages = [{"role": "system", "content": system_instruction}]
        
        # Append history
        for record in history_records:
            openai_messages.append({"role": record.role, "content": record.content})

        # Initialize OpenAI client
        api_key = CauHinh.OPENAI_API_KEY
        if not api_key:
            err_reply = "Tôi xin lỗi, nhưng hệ thống chưa được cấu hình API Key của OpenAI. Vui lòng liên hệ Admin."
            bot_msg = LichSuChatbot(
                nguoi_dung_id=current_user.id,
                role="assistant",
                content=err_reply
            )
            db.session.add(bot_msg)
            db.session.commit()
            return jsonify({"reply": err_reply})

        client = OpenAI(api_key=api_key, max_retries=1)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=openai_messages,
            temperature=0.5,
            max_tokens=600
        )

        reply = response.choices[0].message.content

        # 4. Save bot's reply to DB
        bot_msg = LichSuChatbot(
            nguoi_dung_id=current_user.id,
            role="assistant",
            content=reply
        )
        db.session.add(bot_msg)
        db.session.commit()

        return jsonify({"reply": reply})

    except Exception as e:
        logging.error(f"Chatbot API error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": "Có lỗi xảy ra trong quá trình xử lý phản hồi."}), 500

@app.route("/api/chat/history", methods=["GET"])
@login_required
def api_chat_history():
    from flask import jsonify
    from flask_login import current_user
    from mo_hinh import LichSuChatbot

    try:
        history_records = LichSuChatbot.query.filter_by(nguoi_dung_id=current_user.id).order_by(LichSuChatbot.ngay_tao.asc()).all()
        history = []
        for record in history_records:
            history.append({
                "role": record.role,
                "content": record.content
            })
        return jsonify({"history": history})
    except Exception as e:
        logging.error(f"Chatbot history API error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": "Không thể tải lịch sử chatbot."}), 500

@app.route("/san-pham")
def san_pham():
    from mo_hinh import LichSuGiaoTrinh, db
    import os, json
    from cau_hinh import CauHinh
    from flask_login import current_user
    
    # Chỉ cho phép Admin xem các giáo trình tùy chỉnh đã chọn trong hệ thống.
    # Tài khoản không phải admin (hoặc khách) chỉ có quyền xem các giáo trình mẫu tĩnh.
    is_admin = current_user.is_authenticated and current_user.la_admin
    
    recent_items = LichSuGiaoTrinh.query.filter_by(noi_bat=True).order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    products = []
    
    for item in recent_items:
        chuong = 0
        trich_dan = 0
        chinh_xac = 0
        ma_cv = item.ma_cv
        if ma_cv:
            p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
            if os.path.exists(p_json):
                with open(p_json, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    book = data.get('book_vi', {})
                    if book and 'chapters' in book:
                        chuong = len(book['chapters'])
                    refs = data.get('references', [])
                    trich_dan = len(refs)
                    if 'grounding' in data and 'q_score' in data['grounding']:
                        chinh_xac = data['grounding']['q_score'] * 100
        
        if chuong == 0: chuong = max(3, int(item.do_dai_ky_tu / 5000))
        if trich_dan == 0: trich_dan = max(10, int(item.do_dai_ky_tu / 2000))
        if chinh_xac == 0: chinh_xac = 95.0 + (item.do_dai_ky_tu % 50) / 10.0
        
        products.append({
            "id": item.id,
            "chu_de": item.chu_de,
            "chuong": chuong,
            "chinh_xac": round(chinh_xac, 1),
            "trich_dan": trich_dan,
            "ma_cv": ma_cv,
            "ngay_tao": item.ngay_tao.timestamp() if item.ngay_tao else 0
        })
        
    products.sort(key=lambda x: (x['chinh_xac'], x['ngay_tao']), reverse=True)
    
    all_curriculums = []
    if is_admin:
        all_curriculums = LichSuGiaoTrinh.query.order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
        
    return render_template("showcase.html", products=products, all_curriculums=all_curriculums)

@app.route("/xem-san-pham/<int:id>")
def xem_san_pham(id):
    from mo_hinh import LichSuGiaoTrinh
    import os, json
    from cau_hinh import CauHinh
    from flask_login import current_user
    
    item = db.get_or_404(LichSuGiaoTrinh, id)
    
    is_admin = current_user.is_authenticated and current_user.la_admin
    if not is_admin and not item.noi_bat:
        flash("Bạn không có quyền xem giáo trình này.", "danger")
        return redirect(url_for("san_pham"))
    
    ma_cv = item.ma_cv
    if ma_cv:
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            fake_info = {
                "trang_thai": "hoan_thanh", 
                "tieu_de": item.chu_de,
                "tai_docx": f"/tai/docx/{ma_cv}",
                "tai_pdf": f"/tai/pdf/{ma_cv}",
                "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                "nguon": data.get('references', [])
            }
            return render_template("result.html", ma_cv=ma_cv, thong_tin=fake_info, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}), extracted_terms=data.get('extracted_terms', []), kb_headings=data.get('kb_headings', []))
            
    if item.noi_dung_html:
        if "<html" not in item.noi_dung_html.lower():
            from flask import render_template_string
            return render_template_string("""
            <!doctype html>
            <html lang="vi">
            <head>
                <meta charset="utf-8">
                <title>{{ title }}</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>body { padding: 40px; font-family: 'Times New Roman', serif; max-width: 900px; margin: auto; line-height: 1.6; font-size: 13pt; background: #f8fafc; } .paper { background: white; padding: 50px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border-radius: 8px; }</style>
            </head>
            <body>
                <div class="d-flex justify-content-between align-items-center mb-4">
                    <a href="javascript:history.back()" class="btn btn-outline-secondary">← Quay lại</a>
                    <span class="badge bg-warning text-dark">Chế độ xem tối giản (Bản nháp)</span>
                </div>
                <div class="paper">{{ html_content|safe }}</div>
            </body>
            </html>
            """, title=item.chu_de, html_content=item.noi_dung_html)
        return item.noi_dung_html
        
    return "Nội dung giáo trình không còn khả dụng hoặc đã bị lỗi khi lưu.", 404

@app.get("/")
def trang_chu():
    return render_template("index.html")

@app.get("/privacy-policy")
def privacy_policy():
    page = TrangThongTin.query.filter_by(ma_trang='privacy-policy').first()
    if not page:
        return render_template("privacy_policy.html")
    return render_template("page_dynamic.html", page=page)

@app.get("/terms-of-service")
def terms_of_service():
    page = TrangThongTin.query.filter_by(ma_trang='terms-of-service').first()
    if not page:
        return render_template("terms_of_service.html")
    return render_template("page_dynamic.html", page=page)

@app.get("/data-deletion")
def data_deletion():
    page = TrangThongTin.query.filter_by(ma_trang='data-deletion').first()
    if not page:
        return render_template("data_deletion.html")
    return render_template("page_dynamic.html", page=page)

@app.get("/ai-terms")
def ai_terms():
    page = TrangThongTin.query.filter_by(ma_trang='ai-terms').first()
    if not page:
        return "Trang không tồn tại", 404
    return render_template("page_dynamic.html", page=page)

@app.route("/xoa-tai-khoan", methods=["GET", "POST"])
def xoa_tai_khoan():
    if request.method == "POST":
        username = (request.form.get("username") or "").strip()
        email = (request.form.get("email") or "").strip()
        reason = (request.form.get("reason") or "").strip()
        notes = (request.form.get("notes") or "").strip()
        
        if not username or not email:
            flash("Vui lòng điền đầy đủ tên đăng nhập và email.", "danger")
            return render_template("request_delete_account.html")
            
        account_info = {
            "username": username,
            "email": email,
            "reason": reason,
            "notes": notes
        }
        
        from dich_vu.email_service import gui_email_xoa_tai_khoan
        from cau_hinh import CauHinh
        
        admin_email = CauHinh.ADMIN_NOTIFICATION_EMAIL or "phanvantho082019@gmail.com"
        success, message = gui_email_xoa_tai_khoan(admin_email, account_info)
        
        if success:
            flash(message, "success")
        else:
            flash(message, "warning")
            
        return render_template("request_delete_account.html")
        
    return render_template("request_delete_account.html")

@app.route("/support", methods=["GET", "POST"])
def support():
    if request.method == "POST":
        data = request.get_json() or {}
        name = (data.get("name") or "").strip()
        email = (data.get("email") or "").strip()
        subject = (data.get("subject") or "").strip()
        message = (data.get("message") or "").strip()
        
        if not name or not email or not subject or not message:
            return jsonify({
                "success": False, 
                "message": "Vui lòng điền đầy đủ thông tin yêu cầu hỗ trợ."
            }), 400
            
        ticket_info = {
            "name": name,
            "email": email,
            "subject": subject,
            "message": message
        }
        
        from dich_vu.email_service import gui_email_support_ticket
        from cau_hinh import CauHinh
        
        admin_email = CauHinh.ADMIN_NOTIFICATION_EMAIL or "phanvantho082019@gmail.com"
        success, response_msg = gui_email_support_ticket(admin_email, ticket_info)
        
        return jsonify({
            "success": success,
            "message": response_msg
        })
        
    return render_template("support.html")


@app.get("/tao-giao-trinh")
@login_required
def trang_tao_giao_trinh():
    return render_template(
        "app.html", 
        mac_dinh_top_k=CauHinh.SO_DOAN_THAM_KHAO, 
        mac_dinh_linked=CauHinh.SO_TRANG_LIEN_KET,
        phi_token_auto=getattr(CauHinh, "PHI_TOKEN_AUTO", 1),
        phi_token_expert=getattr(CauHinh, "PHI_TOKEN_EXPERT", 2),
        phi_token_creative=getattr(CauHinh, "PHI_TOKEN_CREATIVE", 3),
        max_so_chuong=getattr(CauHinh, "MAC_DINH_SO_CHUONG_MAX", 15),
        max_so_tu=getattr(CauHinh, "MAC_DINH_SO_TU_MAX", 1000)
    )

@app.route("/admin")
@login_required
@admin_required
def admin_dashboard():
    users = NguoiDung.query.all()
    history = LichSuGiaoTrinh.query.order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    return render_template(
        "admin_overview.html",
        users=users,
        history=history,
        current_section="overview"
    )

@app.route("/admin/users")
@login_required
@admin_required
def admin_users():
    users = NguoiDung.query.all()
    return render_template(
        "admin_users.html",
        users=users,
        current_section="users"
    )

@app.route("/admin/curriculums")
@login_required
@admin_required
def admin_curriculums():
    history = LichSuGiaoTrinh.query.order_by(LichSuGiaoTrinh.ngay_tao.desc()).all()
    return render_template(
        "admin_curriculums.html",
        history=history,
        current_section="curriculums"
    )

@app.route("/admin/curriculums/delete/<int:id>", methods=["POST"])
@login_required
@admin_required
def admin_delete_curriculum(id):
    try:
        from mo_hinh import db, LichSuGiaoTrinh
        import os
        item = db.session.get(LichSuGiaoTrinh, id)
        if item:
            # Delete associated files if they exist
            ma_cv = item.ma_cv
            if ma_cv:
                # File paths to delete
                file_paths = [
                    os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json"),
                    os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}_plain.json"),
                    os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}.pdf"),
                    os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}_plain.pdf"),
                    os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}.docx"),
                    os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}_plain.docx")
                ]
                for p in file_paths:
                    if os.path.exists(p):
                        try:
                            os.remove(p)
                        except Exception as fe:
                            logger.error(f"Error removing file {p}: {fe}")
                            
                # Delete from Azure Blob Storage if configured
                try:
                    from dich_vu.azure_blob import delete_from_blob
                    delete_from_blob(f"json/{ma_cv}.json")
                    delete_from_blob(f"json/{ma_cv}_plain.json")
                    delete_from_blob(f"pdf/{ma_cv}.pdf")
                    delete_from_blob(f"pdf/{ma_cv}_plain.pdf")
                    delete_from_blob(f"docx/{ma_cv}.docx")
                    delete_from_blob(f"docx/{ma_cv}_plain.docx")
                except Exception as blob_err:
                    logger.error(f"Error removing files from Azure Blob Storage: {blob_err}")
            
            db.session.delete(item)
            db.session.commit()
            return jsonify({"success": True, "message": "Xóa giáo trình thành công."})
        return jsonify({"success": False, "error": "Không tìm thấy giáo trình."}), 404
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting curriculum {id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/admin/curriculums/toggle-showcase/<int:id>", methods=["POST"])
@login_required
@admin_required
def admin_toggle_showcase(id):
    try:
        from mo_hinh import db, LichSuGiaoTrinh
        item = db.session.get(LichSuGiaoTrinh, id)
        if not item:
            return jsonify({"success": False, "error": "Không tìm thấy giáo trình."}), 404
            
        data = request.get_json() or {}
        item.noi_bat = bool(data.get("noi_bat", False))
        db.session.commit()
        return jsonify({"success": True, "message": "Đã cập nhật trạng thái trưng bày thành công."})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error toggling showcase for curriculum {id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/admin/settings")
@login_required
@admin_required
def admin_settings():
    masked_openai = CauHinh.OPENAI_API_KEY or ""
    masked_gemini = ", ".join(CauHinh.GEMINI_API_KEYS) if CauHinh.GEMINI_API_KEYS else ""
    masked_vnpay_secret = CauHinh.VNPAY_HASH_SECRET or ""
    masked_sepay_key = CauHinh.SEPAY_API_KEY or ""
    
    return render_template(
        "admin_settings.html",
        cau_hinh=CauHinh,
        masked_openai=masked_openai,
        masked_gemini=masked_gemini,
        masked_vnpay_secret=masked_vnpay_secret,
        masked_sepay_key=masked_sepay_key,
        current_section="settings"
    )

@app.route("/admin/add_user", methods=["POST"])
@login_required
@admin_required
def admin_add_user():
    u = request.form.get("new_username")
    p = request.form.get("new_password")
    is_admin = request.form.get("is_admin") == "on"
    if u and p:
        if NguoiDung.query.filter_by(ten_dang_nhap=u).first():
            flash("Tên đăng nhập đã tồn tại.", "danger")
        else:
            new_user = NguoiDung(ten_dang_nhap=u, mat_khau=generate_password_hash(p), la_admin=is_admin, email=f"{u}@local")
            db.session.add(new_user)
            db.session.commit()
            flash("Thêm người dùng thành công.", "success")
    return redirect(url_for("admin_users"))

@app.route("/admin/update_settings", methods=["POST"])
@login_required
@admin_required
def admin_update_settings():
    from dotenv import set_key
    import os
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
    
    data = request.json
    if not data:
        return jsonify({"success": False, "error": "No data provided"})
        
    keys_to_update = [
        "OPENAI_API_KEY", "OPENAI_MODEL", 
        "GEMINI_API_KEYS", "GEMINI_MODEL", "GEMINI_MODEL_LITE",
        "WRITER_MODEL", "SEARCH_MODEL", "SUPERVISOR_MODEL_LITE", "SUPERVISOR_MODEL_PRO",
        "PAYMENT_VNPAY_ACTIVE", "PAYMENT_SEPAY_ACTIVE",
        "VNPAY_TMN_CODE", "VNPAY_HASH_SECRET", "VNPAY_PAYMENT_URL", "VNPAY_RETURN_URL",
        "SEPAY_API_KEY", "SEPAY_ACCOUNT_NUMBER", "SEPAY_BANK_BRAND", "SEPAY_WEB_NAME", "SEPAY_XOR_KEY",
        "PHI_TOKEN_AUTO", "PHI_TOKEN_EXPERT", "PHI_TOKEN_CREATIVE",
        "CONTACT_EMAIL", "CONTACT_PHONE", "CONTACT_ADDRESS_VI", "CONTACT_ADDRESS_EN",
        "ADMIN_NOTIFICATION_EMAIL", "MAC_DINH_SO_CHUONG_MAX", "MAC_DINH_SO_TU_MAX", "GOOGLE_CLIENT_ID"
    ]


    
    def mask_key(key):
        if not key:
            return ""
        key = key.strip()
        if len(key) <= 10:
            return "********"
        return f"{key[:6]}...{key[-4:]}"
        
    try:
        for key in keys_to_update:
            if key in data:
                new_val = str(data[key])
                
                # Check for masked keys
                if key in ["OPENAI_API_KEY", "VNPAY_HASH_SECRET", "SEPAY_API_KEY"]:
                    if "..." in new_val or "*" in new_val:
                        # Keep original key
                        continue
                elif key == "GEMINI_API_KEYS":
                    tokens = [t.strip() for t in new_val.split(",") if t.strip()]
                    resolved_keys = []
                    for t in tokens:
                        if "..." in t or "*" in t:
                            matched = None
                            for orig in CauHinh.GEMINI_API_KEYS:
                                if mask_key(orig) == t:
                                    matched = orig
                                    break
                            if matched:
                                resolved_keys.append(matched)
                        else:
                            resolved_keys.append(t)
                    new_val = ", ".join(resolved_keys)
                
                # Write to .env
                set_key(env_path, key, new_val)
                
                # Write to DB
                try:
                    from mo_hinh import CauHinhHeThong
                    from cau_hinh import ma_hoa_key
                    api_keys_to_encrypt = ["OPENAI_API_KEY", "GEMINI_API_KEYS", "SEPAY_API_KEY", "VNPAY_HASH_SECRET"]
                    db_val = ma_hoa_key(new_val) if key in api_keys_to_encrypt else new_val
                    
                    config_item = CauHinhHeThong.query.filter_by(khoa=key).first()
                    if not config_item:
                        config_item = CauHinhHeThong(khoa=key, gia_tri=db_val)
                        db.session.add(config_item)
                    else:
                        config_item.gia_tri = db_val
                    db.session.commit()
                except Exception as db_err:
                    db.session.rollback()
                    logger.error(f"Lỗi lưu CSDL cho key {key}: {db_err}")
                
                # Update in memory
                if key == "GEMINI_API_KEYS":
                    CauHinh.GEMINI_API_KEYS = [k.strip() for k in new_val.split(",") if k.strip()]
                elif key in ["PAYMENT_VNPAY_ACTIVE", "PAYMENT_SEPAY_ACTIVE"]:
                    setattr(CauHinh, key, new_val == "True")
                elif key in ["PHI_TOKEN_AUTO", "PHI_TOKEN_EXPERT", "PHI_TOKEN_CREATIVE", "MAC_DINH_SO_CHUONG_MAX", "MAC_DINH_SO_TU_MAX"]:
                    setattr(CauHinh, key, int(new_val))
                elif key == "SEPAY_XOR_KEY":
                    val = int(new_val, 16) if "0x" in new_val.lower() else int(new_val)
                    setattr(CauHinh, key, val)
                else:
                    setattr(CauHinh, key, new_val)
                    
        return jsonify({"success": True})
    except Exception as e:
        logger.error(f"Lỗi cập nhật settings: {e}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/static/i18n.js")
def custom_i18n_js():
    from cau_hinh import CauHinh
    import os
    from flask import Response
    
    js_path = os.path.join(app.static_folder, "i18n.js")
    try:
        with open(js_path, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception as e:
        logger.error(f"Lỗi đọc i18n.js: {e}")
        content = "const TRANSLATIONS = {};"
    
    email = CauHinh.CONTACT_EMAIL or "phanvantho082019@gmail.com"
    phone = CauHinh.CONTACT_PHONE or "0327152710"
    addr_vi = CauHinh.CONTACT_ADDRESS_VI or "Cần Thơ, Việt Nam"
    addr_en = CauHinh.CONTACT_ADDRESS_EN or "Can Tho, Vietnam"
    
    email_js = email.replace("'", "\\'").replace('"', '\\"')
    phone_js = phone.replace("'", "\\'").replace('"', '\\"')
    addr_vi_js = addr_vi.replace("'", "\\'").replace('"', '\\"')
    addr_en_js = addr_en.replace("'", "\\'").replace('"', '\\"')
    
    max_so_chuong = getattr(CauHinh, "MAC_DINH_SO_CHUONG_MAX", 15)
    max_so_tu = getattr(CauHinh, "MAC_DINH_SO_TU_MAX", 1000)

    override_js = f"""
// Dynamic custom configuration override
if (typeof TRANSLATIONS !== 'undefined') {{
  TRANSLATIONS['footer.email_val'] = {{ vi: '{email_js}', en: '{email_js}' }};
  TRANSLATIONS['footer.phone_val'] = {{ vi: '{phone_js}', en: '{phone_js}' }};
  TRANSLATIONS['footer.address_val'] = {{ vi: '{addr_vi_js}', en: '{addr_en_js}' }};
  TRANSLATIONS['app.chapters_count'] = {{ vi: 'SỐ LƯỢNG CHƯƠNG TỰ CHỌN (1-{max_so_chuong})', en: 'NUMBER OF CHAPTERS (1-{max_so_chuong})' }};
  TRANSLATIONS['app.section_words'] = {{ vi: 'ĐỘ DÀI MỖI MỤC CON CẤP 2 (100 - {max_so_tu} từ)', en: 'LEVEL 2 SUBSECTION LENGTH (100 - {max_so_tu} words)' }};
  TRANSLATIONS['app.section_words_help'] = {{ vi: '(Số từ tối thiểu áp dụng cho mỗi mục con cấp 2 của chương, ví dụ: mục 1.1, 1.2)', en: '(Minimum word count applied to each level 2 subsection/topic in a chapter, e.g. section 1.1, 1.2)' }};
  TRANSLATIONS['app.approve_outline'] = {{ vi: 'Tôi muốn duyệt dàn ý chi tiết trước khi sinh nội dung', en: 'Review detailed outline before generating content' }};
  TRANSLATIONS['app.approve_outline_title'] = {{ vi: 'Duyệt dàn ý chi tiết', en: 'Review Detailed Outline' }};
  TRANSLATIONS['app.approve_outline_desc'] = {{ vi: 'Đánh dấu chọn các mục con cấp 2 bạn muốn giữ lại để viết nội dung:', en: 'Select the level 2 sections you want to keep for content generation:' }};
  TRANSLATIONS['app.btn_approve_outline'] = {{ vi: 'Tiến hành viết', en: 'Proceed to Write' }};
  TRANSLATIONS['app.chapters_hint'] = {{ vi: 'Nhập số chương mong muốn (3-{max_so_chuong}). Để trống để tự động phân tích quy mô.', en: 'Enter desired chapters (3-{max_so_chuong}). Leave blank to auto-detect by scale.' }};
}}
"""
    
    full_content = content + "\n" + override_js
    
    response = Response(full_content, mimetype="application/javascript")
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

# --- ADMIN BLOCK USER & DYNAMIC PACKAGES ---
@app.route("/admin/toggle_block_user", methods=["POST"])
@login_required
@admin_required
def admin_toggle_block_user():
    data = request.json or {}
    user_id = data.get("user_id")
    if not user_id:
        return jsonify({"success": False, "error": "Thiếu user_id"}), 400
    
    user = db.session.get(NguoiDung, int(user_id))
    if not user:
        return jsonify({"success": False, "error": "Không tìm thấy người dùng"}), 404
        
    if user.la_admin and user.id == current_user.id:
        return jsonify({"success": False, "error": "Bạn không thể tự khóa tài khoản của chính mình!"}), 400
        
    user.bi_khoa = not user.bi_khoa
    db.session.commit()
    
    action = "khóa" if user.bi_khoa else "mở khóa"
    return jsonify({"success": True, "bi_khoa": user.bi_khoa, "message": f"Đã {action} tài khoản thành viên {user.ten_dang_nhap}!"})

@app.route("/admin/delete_user", methods=["POST"])
@login_required
@admin_required
def admin_delete_user():
    data = request.json or {}
    user_id = data.get("user_id")
    if not user_id:
        return jsonify({"success": False, "error": "Thiếu user_id"}), 400
        
    user = db.session.get(NguoiDung, int(user_id))
    if not user:
        return jsonify({"success": False, "error": "Không tìm thấy người dùng"}), 404
        
    if user.la_admin and user.id == current_user.id:
        return jsonify({"success": False, "error": "Bạn không thể tự xóa tài khoản của chính mình!"}), 400
        
    db.session.delete(user)
    db.session.commit()
    
    return jsonify({"success": True, "message": f"Đã xóa tài khoản thành viên {user.ten_dang_nhap}!"})

@app.route("/admin/pages", methods=["GET"])

@login_required
@admin_required
def admin_pages():
    pages = TrangThongTin.query.all()
    return render_template("admin_pages.html", pages=pages, current_section="pages")

@app.route("/admin/pages/save", methods=["POST"])
@login_required
@admin_required
def admin_save_page():
    try:
        page_id = request.form.get("id")
        tieu_de_vi = request.form.get("tieu_de_vi")
        tieu_de_en = request.form.get("tieu_de_en")
        mo_ta_vi = request.form.get("mo_ta_vi", "")
        mo_ta_en = request.form.get("mo_ta_en", "")
        noi_dung_vi = request.form.get("noi_dung_vi", "")
        noi_dung_en = request.form.get("noi_dung_en", "")
        
        if not page_id:
            flash("Thiếu ID trang.", "danger")
            return redirect(url_for("admin_pages"))
            
        page = db.session.get(TrangThongTin, int(page_id))
        if page:
            page.tieu_de_vi = tieu_de_vi
            page.tieu_de_en = tieu_de_en
            page.mo_ta_vi = mo_ta_vi
            page.mo_ta_en = mo_ta_en
            page.noi_dung_vi = noi_dung_vi
            page.noi_dung_en = noi_dung_en
            db.session.commit()
            flash(f"Cập nhật trang '{tieu_de_vi}' thành công.", "success")
        else:
            flash("Không tìm thấy trang.", "danger")
    except Exception as e:
        db.session.rollback()
        logger.error(f"Lỗi lưu trang: {e}")
        flash(f"Lỗi: {str(e)}", "danger")
        
    return redirect(url_for("admin_pages"))

@app.route("/admin/packages", methods=["GET"])
@login_required
@admin_required
def admin_packages():
    packages = GoiCuoc.query.all()
    return render_template("admin_packages.html", packages=packages, current_section="packages")

@app.route("/admin/packages/save", methods=["POST"])
@login_required
@admin_required
def admin_save_package():
    try:
        pkg_id = request.form.get("id")
        ten_goi = request.form.get("ten_goi")
        gia_tien = int(request.form.get("gia_tien", 0))
        so_token = int(request.form.get("so_token", 0))
        mo_ta = request.form.get("mo_ta", "")
        kich_hoat = request.form.get("kich_hoat") == "on" or request.form.get("kich_hoat") == "true"
        
        if not ten_goi:
            flash("Tên gói không được để trống.", "danger")
            return redirect(url_for("admin_packages"))
            
        if pkg_id:
            # Update existing
            pkg = db.session.get(GoiCuoc, int(pkg_id))
            if pkg:
                pkg.ten_goi = ten_goi
                pkg.gia_tien = gia_tien
                pkg.so_token = so_token
                pkg.mo_ta = mo_ta
                pkg.kich_hoat = kich_hoat
                flash(f"Cập nhật gói cước '{ten_goi}' thành công.", "success")
        else:
            # Create new
            pkg = GoiCuoc(ten_goi=ten_goi, gia_tien=gia_tien, so_token=so_token, mo_ta=mo_ta, kich_hoat=kich_hoat)
            db.session.add(pkg)
            flash(f"Thêm gói cước mới '{ten_goi}' thành công.", "success")
            
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.error(f"Lỗi lưu gói cước: {e}")
        flash(f"Lỗi: {str(e)}", "danger")
        
    return redirect(url_for("admin_packages"))

@app.route("/admin/packages/delete/<int:id>", methods=["POST"])
@login_required
@admin_required
def admin_delete_package(id):
    try:
        pkg = db.session.get(GoiCuoc, id)
        if pkg:
            db.session.delete(pkg)
            db.session.commit()
            return jsonify({"success": True, "message": "Xóa gói cước thành công."})
        return jsonify({"success": False, "error": "Không tìm thấy gói cước."}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/admin/packages/toggle/<int:id>", methods=["POST"])
@login_required
@admin_required
def admin_toggle_package(id):
    try:
        pkg = db.session.get(GoiCuoc, id)
        if pkg:
            pkg.kich_hoat = not pkg.kich_hoat
            db.session.commit()
            action = "kích hoạt" if pkg.kich_hoat else "hủy kích hoạt"
            return jsonify({"success": True, "kich_hoat": pkg.kich_hoat, "message": f"Đã {action} gói cước {pkg.ten_goi}!"})
        return jsonify({"success": False, "error": "Không tìm thấy gói cước."}), 404
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/lich-su/<int:id>")
@login_required
def hien_thi_lich_su(id):
    item = db.get_or_404(LichSuGiaoTrinh, id)
    if not current_user.la_admin and item.nguoi_dung_id != current_user.id:
        flash("Bạn không có quyền xem giáo trình này.", "danger")
        return redirect(url_for("lich_su"))
    
    ma_cv = item.ma_cv
    if ma_cv:
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        # Try downloading json file from Azure Blob Storage if it doesn't exist locally (V34+)
        if not os.path.exists(p_json):
            try:
                from dich_vu.azure_blob import download_from_blob
                download_from_blob(f"json/{ma_cv}.json", p_json)
            except Exception as blob_err:
                logger.error(f"Failed to download JSON file from Azure Blob Storage: {blob_err}")

        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            fake_info = {
                "trang_thai": "hoan_thanh", 
                "tieu_de": item.chu_de,
                "tai_docx": f"/tai/docx/{ma_cv}",
                "tai_pdf": f"/tai/pdf/{ma_cv}",
                "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                "nguon": data.get('references', []),
                "audit_quy_mo": data.get('audit_quy_mo')
            }
            return render_template("result.html", ma_cv=ma_cv, thong_tin=fake_info, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}), extracted_terms=data.get('extracted_terms', []), kb_headings=data.get('kb_headings', []))
            
    if item.noi_dung_html:
        if "<html" not in item.noi_dung_html.lower():
            from flask import render_template_string
            return render_template_string("""
            <!doctype html>
            <html lang="vi">
            <head>
                <meta charset="utf-8">
                <title>{{ title }}</title>
                <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
                <style>body { padding: 40px; font-family: 'Times New Roman', serif; max-width: 900px; margin: auto; line-height: 1.6; font-size: 13pt; background: #f8fafc; } .paper { background: white; padding: 50px; box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1); border-radius: 8px; }</style>
            </head>
            <body>
                <div class="d-flex justify-content-between align-items-center mb-4">
                    <a href="javascript:history.back()" class="btn btn-outline-secondary">← Quay lại</a>
                    <span class="badge bg-warning text-dark">Chế độ xem tối giản (Bản nháp)</span>
                </div>
                <div class="paper">{{ html_content|safe }}</div>
            </body>
            </html>
            """, title=item.chu_de, html_content=item.noi_dung_html)
        return item.noi_dung_html
        
    return "Nội dung giáo trình không còn khả dụng hoặc đã bị lỗi khi lưu.", 404

@app.route("/lich-su/xoa/<int:id>", methods=["POST"])
@login_required
def xoa_lich_su(id):
    try:
        from mo_hinh import db, LichSuGiaoTrinh
        import os
        item = db.session.get(LichSuGiaoTrinh, id)
        if not item:
            return jsonify({"success": False, "error": "Không tìm thấy giáo trình."}), 404
        
        # Verify ownership
        if not current_user.la_admin and item.nguoi_dung_id != current_user.id:
            return jsonify({"success": False, "error": "Bạn không có quyền xóa giáo trình này."}), 403
            
        ma_cv = item.ma_cv
        if ma_cv:
            # File paths to delete
            file_paths = [
                os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json"),
                os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}_plain.json"),
                os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}.pdf"),
                os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}_plain.pdf"),
                os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}.docx"),
                os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}_plain.docx")
            ]
            for p in file_paths:
                if os.path.exists(p):
                    try:
                        os.remove(p)
                    except Exception as fe:
                        logger.error(f"Error removing file {p}: {fe}")
                        
            # Delete from Azure Blob Storage if configured
            try:
                from dich_vu.azure_blob import delete_from_blob
                delete_from_blob(f"json/{ma_cv}.json")
                delete_from_blob(f"json/{ma_cv}_plain.json")
                delete_from_blob(f"pdf/{ma_cv}.pdf")
                delete_from_blob(f"pdf/{ma_cv}_plain.pdf")
                delete_from_blob(f"docx/{ma_cv}.docx")
                delete_from_blob(f"docx/{ma_cv}_plain.docx")
            except Exception as blob_err:
                logger.error(f"Error removing files from Azure Blob Storage: {blob_err}")
                
        db.session.delete(item)
        db.session.commit()
        return jsonify({"success": True, "message": "Xóa lịch sử giáo trình thành công."})
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting curriculum history {id}: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/lich-su/xoa-luu-tam/<string:ma_cv>", methods=["POST"])
@login_required
def xoa_luu_tam(ma_cv):
    if ma_cv in CONG_VIEC:
        job = CONG_VIEC[ma_cv]
        if job.get("user_id") == current_user.id:
            if job.get("trang_thai") in ["that_bai", "huy_bo"]:
                CONG_VIEC.pop(ma_cv)
                return jsonify({"success": True, "message": "Xóa thông tin tiến trình thành công."})
            else:
                return jsonify({"success": False, "error": "Không thể xóa tiến trình đang chạy."}), 400
    return jsonify({"success": False, "error": "Không tìm thấy tiến trình."}), 404


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    if request.method == "POST":
        import os
        import time
        from werkzeug.utils import secure_filename
        from werkzeug.security import generate_password_hash, check_password_hash

        new_ten_dang_nhap = (request.form.get("ten_dang_nhap") or "").strip()
        new_ho_ten = request.form.get("ho_ten")
        new_email = (request.form.get("email") or "").strip().lower()
        mat_khau_cu = request.form.get("mat_khau_cu")
        mat_khau_moi = request.form.get("mat_khau_moi")
        xac_nhan_mat_khau = request.form.get("xac_nhan_mat_khau_moi")
        avatar_file = request.files.get("anh_dai_dien")
        
        profile_updated = False

        # 0. Cập nhật tên đăng nhập
        if new_ten_dang_nhap and new_ten_dang_nhap != current_user.ten_dang_nhap:
            if " " in new_ten_dang_nhap:
                flash("Tên đăng nhập không được chứa khoảng trắng.", "danger")
                return redirect(url_for("profile"))
            if len(new_ten_dang_nhap) < 3:
                flash("Tên đăng nhập phải có ít nhất 3 ký tự.", "danger")
                return redirect(url_for("profile"))
            existing_user = NguoiDung.query.filter_by(ten_dang_nhap=new_ten_dang_nhap).first()
            if existing_user:
                flash("Tên đăng nhập này đã được sử dụng bởi tài khoản khác.", "danger")
                return redirect(url_for("profile"))
            current_user.ten_dang_nhap = new_ten_dang_nhap
            db.session.commit()
            profile_updated = True

        # 1. Cập nhật họ tên
        if new_ho_ten and new_ho_ten != current_user.ho_ten:
            current_user.ho_ten = new_ho_ten
            db.session.commit()
            profile_updated = True
            
        # 2. Cập nhật ảnh đại diện
        if avatar_file and avatar_file.filename != "":
            filename = avatar_file.filename
            ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            if ext not in ["png", "jpg", "jpeg", "gif"]:
                flash("Định dạng ảnh không hợp lệ (chỉ chấp nhận png, jpg, jpeg, gif).", "danger")
                return redirect(url_for("profile"))
                
            new_filename = f"avatar_{current_user.id}_{int(time.time())}.{ext}"
            upload_dir = os.path.join(app.root_path, "static", "uploads", "avatars")
            os.makedirs(upload_dir, exist_ok=True)
            upload_path = os.path.join(upload_dir, new_filename)
            avatar_file.save(upload_path)
            
            # Xóa ảnh cũ
            if current_user.anh_dai_dien and not current_user.anh_dai_dien.startswith("http"):
                old_path = os.path.join(upload_dir, current_user.anh_dai_dien)
                if os.path.exists(old_path):
                    try:
                        os.remove(old_path)
                    except Exception as e:
                        logger.error(f"Error removing old avatar: {e}")
                        
            current_user.anh_dai_dien = new_filename
            db.session.commit()
            profile_updated = True

        # 3. Thay đổi mật khẩu
        if mat_khau_moi:
            if mat_khau_moi != xac_nhan_mat_khau:
                flash("Mật khẩu mới không trùng khớp.", "danger")
                return redirect(url_for("profile"))
                
            if current_user.mat_khau:
                if not mat_khau_cu or not check_password_hash(current_user.mat_khau, mat_khau_cu):
                    flash("Mật khẩu cũ không chính xác.", "danger")
                    return redirect(url_for("profile"))
                    
            current_user.mat_khau = generate_password_hash(mat_khau_moi)
            db.session.commit()
            profile_updated = True

        # 4. Thay đổi Email (luồng OTP)
        if new_email and new_email != (current_user.email or "").strip().lower():
            if NguoiDung.query.filter_by(email=new_email).first():
                flash("Email này đã được sử dụng bởi tài khoản khác.", "danger")
                return redirect(url_for("profile"))
                
            otp_code = "".join([str(random.randint(0, 9)) for _ in range(6)])
            from datetime import timedelta
            het_han = datetime.utcnow() + timedelta(minutes=5)
            
            XacThucOTP.query.filter_by(email=new_email, da_dung=False).update({"da_dung": True})
            otp_record = XacThucOTP(
                email=new_email,
                ten_dang_nhap=current_user.ten_dang_nhap,
                mat_khau_hash="",
                otp=otp_code,
                het_han=het_han,
                da_dung=False,
                nguoi_dung_id=current_user.id
            )
            db.session.add(otp_record)
            db.session.commit()
            
            from dich_vu.email_service import gui_email_otp
            success, mail_msg = gui_email_otp(new_email, otp_code)
            
            if success:
                session["pending_new_email"] = new_email
                flash("Mã xác thực OTP đã được gửi tới email mới của bạn. Vui lòng nhập mã để xác nhận thay đổi.", "success")
                return redirect(url_for("verify_email_change"))
            else:
                flash(f"Không thể gửi email OTP: {mail_msg}", "warning")
                return redirect(url_for("profile"))
                
        if profile_updated:
            flash("Cập nhật thông tin thành công!", "success")
        return redirect(url_for("profile"))
    return render_template("profile.html")

@app.route("/verify-email-change", methods=["GET", "POST"])
@login_required
def verify_email_change():
    new_email = session.get("pending_new_email")
    if not new_email:
        flash("Không có yêu cầu thay đổi email nào đang chờ xử lý.", "warning")
        return redirect(url_for("profile"))
        
    if request.method == "POST":
        entered_otp = request.form.get("otp", "").strip()
        if not entered_otp:
            flash("Vui lòng nhập mã OTP.", "danger")
            return render_template("verify_email_change.html", email=new_email)
            
        # Tìm mã OTP mới nhất của email mới này
        otp_record = XacThucOTP.query.filter_by(email=new_email, da_dung=False).order_by(XacThucOTP.id.desc()).first()
        
        if not otp_record:
            flash("Không tìm thấy thông tin xác thực OTP. Vui lòng thử lại.", "danger")
            return redirect(url_for("profile"))
            
        # Kiểm tra thời hạn hết hạn
        if datetime.utcnow() > otp_record.het_han:
            flash("Mã OTP đã hết hạn (hạn dùng 5 phút). Vui lòng cập nhật lại email để nhận mã mới.", "danger")
            return redirect(url_for("profile"))
            
        # So khớp OTP
        if otp_record.otp != entered_otp:
            flash("Mã OTP không chính xác. Vui lòng thử lại.", "danger")
            return render_template("verify_email_change.html", email=new_email)
            
        # OTP chính xác -> Cập nhật email chính thức cho user
        # Kiểm tra lại xem email có bị ai đăng ký trước trong lúc chờ không
        if NguoiDung.query.filter_by(email=new_email).first():
            flash("Email này đã bị người khác đăng ký trong lúc chờ xác thực.", "danger")
            return redirect(url_for("profile"))
            
        current_user.email = new_email
        otp_record.da_dung = True
        db.session.commit()
        
        session.pop("pending_new_email", None)
        flash("Thay đổi địa chỉ email thành công!", "success")
        return redirect(url_for("profile"))
        
    return render_template("verify_email_change.html", email=new_email)

@app.route("/buy-tokens")
@login_required
def buy_tokens():
    # Hiển thị trang chọn gói cước
    packages = GoiCuoc.query.filter_by(kich_hoat=True).all()
    return render_template("pricing.html", cau_hinh=CauHinh, packages=packages)

@app.post("/payment/create")
@login_required
def create_payment():
    # Lấy dữ liệu từ form
    tokens = request.form.get("tokens", type=int)
    price = request.form.get("price", type=int)
    package_name = request.form.get("package_name", "Gói cước")
    method = request.form.get("method", "VNPAY")
    goi_cuoc_id = request.form.get("goi_cuoc_id", type=int)
    
    if not tokens or not price:
        flash("Thông tin gói cước không hợp lệ.", "danger")
        return redirect(url_for("buy_tokens"))
        
    # Tạo mã đơn hàng duy nhất (chỉ gồm chữ và số để tránh lỗi cổng thanh toán)
    ma_giao_dich = str(uuid.uuid4())[:8] + datetime.now().strftime("%Y%m%d%H%M%S")
    
    # Lưu giao dịch vào DB với trạng thái 'cho_thanh_toan'
    from mo_hinh import GiaoDichNapToken
    giao_dich = GiaoDichNapToken(
        ma_giao_dich=ma_giao_dich,
        nguoi_dung_id=current_user.id,
        so_tien=price,
        so_token=tokens,
        phuong_thuc=method,
        trang_thai="cho_thanh_toan",
        goi_cuoc_id=goi_cuoc_id
    )
    db.session.add(giao_dich)
    db.session.commit()
    
    if method == "SEPAY":
        return redirect(url_for("sepay_checkout", giao_dich_id=giao_dich.id))
        
    # Khởi tạo lớp VNPay
    from dich_vu.vnpay import VNPay
    vnp = VNPay(
        tmn_code=CauHinh.VNPAY_TMN_CODE,
        hash_secret=CauHinh.VNPAY_HASH_SECRET,
        payment_url=CauHinh.VNPAY_PAYMENT_URL
    )
    
    # Lấy IP người dùng (hoặc mặc định 127.0.0.1)
    ip_addr = request.remote_addr or "127.0.0.1"
    
    # Tạo URL thanh toán
    order_info = f"Nap {tokens} token cho tai khoan {current_user.ten_dang_nhap}"
    payment_url = vnp.create_payment_url(
        txn_ref=ma_giao_dich,
        amount=price,
        order_info=order_info,
        return_url=CauHinh.VNPAY_RETURN_URL,
        ip_addr=ip_addr
    )
    
    # Chuyển hướng người dùng sang VNPAY
    return redirect(payment_url)

@app.route("/payment/sepay/<int:giao_dich_id>")
@login_required
def sepay_checkout(giao_dich_id):
    from mo_hinh import GiaoDichNapToken
    from dich_vu.sepay import encode_payment_id
    
    giao_dich = GiaoDichNapToken.query.get_or_404(giao_dich_id)
    
    if giao_dich.nguoi_dung_id != current_user.id or giao_dich.phuong_thuc != "SEPAY":
        abort(403)
        
    if giao_dich.trang_thai != "cho_thanh_toan":
        return redirect(url_for("buy_tokens"))
        
    hex_id = encode_payment_id(giao_dich.id)
    transfer_content = f"{CauHinh.SEPAY_WEB_NAME}NAPTOKEN{hex_id}"
    
    return render_template("sepay_checkout.html", 
                         giao_dich=giao_dich, 
                         transfer_content=transfer_content,
                         bank_brand=CauHinh.SEPAY_BANK_BRAND,
                         account_number=CauHinh.SEPAY_ACCOUNT_NUMBER)

@app.route("/api/payment/sepay/status/<int:giao_dich_id>")
@login_required
def sepay_status(giao_dich_id):
    from mo_hinh import GiaoDichNapToken
    from dich_vu.sepay import check_sepay_transactions
    
    giao_dich = GiaoDichNapToken.query.get_or_404(giao_dich_id)
    
    if giao_dich.nguoi_dung_id != current_user.id:
        return jsonify({"status": "error", "message": "Unauthorized"}), 403
        
    if giao_dich.trang_thai == "thanh_cong":
        return jsonify({"status": "completed"})
        
    if giao_dich.trang_thai == "cho_thanh_toan":
        # Check SePay API
        is_paid = check_sepay_transactions(giao_dich.id, giao_dich.so_tien)
        if is_paid:
            giao_dich.trang_thai = "thanh_cong"
            giao_dich.ngay_hoan_thanh = datetime.utcnow()
            
            # Add tokens to user
            current_user.token += giao_dich.so_token
            db.session.commit()

            # Gửi email thông báo
            try:
                from dich_vu.email_service import gui_email_thanh_toan_thanh_cong
                gui_email_thanh_toan_thanh_cong(
                    to_email=current_user.email,
                    username=current_user.ten_dang_nhap,
                    so_token=giao_dich.so_token,
                    so_tien=giao_dich.so_tien,
                    ma_giao_dich=giao_dich.ma_giao_dich,
                    phuong_thuc=giao_dich.phuong_thuc
                )
            except Exception as e_mail:
                logger.error(f"Error sending payment success email: {e_mail}")

            return jsonify({"status": "completed"})
            
    return jsonify({"status": "pending"})

@app.route("/api/payment/sepay/webhook", methods=["POST"])
def sepay_webhook():
    # Verify auth header if SEPAY_API_KEY is configured
    if CauHinh.SEPAY_API_KEY:
        auth_header = request.headers.get("Authorization")
        expected_token = f"Bearer {CauHinh.SEPAY_API_KEY}"
        if not auth_header or auth_header != expected_token:
            return jsonify({"status": "error", "message": "Unauthorized"}), 401
            
    data = request.get_json(silent=True) or {}
    content = data.get("content", "")
    try:
        amount_in = float(data.get("transferAmount") or data.get("amountIn") or data.get("amount_in") or 0)
    except (ValueError, TypeError):
        amount_in = 0.0
        
    import re
    from dich_vu.sepay import decode_payment_id
    from mo_hinh import GiaoDichNapToken, NguoiDung
    
    prefix = CauHinh.SEPAY_WEB_NAME + "NAPTOKEN"
    pattern = rf"{prefix}([A-Fa-f0-9]+)"
    match = re.search(pattern, content, re.IGNORECASE)
    if not match:
        return jsonify({"status": "ignored", "message": "Content does not match pattern"}), 200
        
    hex_id = match.group(1).upper()
    try:
        giao_dich_id = decode_payment_id(hex_id)
    except Exception:
        return jsonify({"status": "error", "message": "Invalid payment code"}), 200
        
    giao_dich = GiaoDichNapToken.query.get(giao_dich_id)
    if not giao_dich:
        return jsonify({"status": "error", "message": "Transaction not found"}), 200
        
    if giao_dich.trang_thai != "cho_thanh_toan":
        return jsonify({"status": "success", "message": "Already processed"}), 200
        
    if amount_in < giao_dich.so_tien:
        return jsonify({"status": "error", "message": "Insufficient amount"}), 200
        
    # Mark as completed
    giao_dich.trang_thai = "thanh_cong"
    giao_dich.ngay_hoan_thanh = datetime.utcnow()
    
    user = NguoiDung.query.get(giao_dich.nguoi_dung_id)
    if user:
        user.token += giao_dich.so_token
        
    db.session.commit()
    
    # Send email
    try:
        from dich_vu.email_service import gui_email_thanh_toan_thanh_cong
        gui_email_thanh_toan_thanh_cong(
            to_email=user.email if user else None,
            username=user.ten_dang_nhap if user else "N/A",
            so_token=giao_dich.so_token,
            so_tien=giao_dich.so_tien,
            ma_giao_dich=giao_dich.ma_giao_dich,
            phuong_thuc=giao_dich.phuong_thuc
        )
    except Exception as e_mail:
        logger.error(f"Error sending payment success email: {e_mail}")
        
    return jsonify({"status": "success", "message": "Payment completed successfully"}), 200

@app.route("/payment/callback")
@login_required
def payment_callback():
    from dich_vu.vnpay import VNPay
    from mo_hinh import GiaoDichNapToken
    
    # Lấy các tham số VNPAY trả về
    vnp_params = request.args.to_dict()
    
    vnp = VNPay(
        tmn_code=CauHinh.VNPAY_TMN_CODE,
        hash_secret=CauHinh.VNPAY_HASH_SECRET,
        payment_url=CauHinh.VNPAY_PAYMENT_URL
    )
    
    # Xác thực chữ ký số
    is_valid = vnp.verify_payment(vnp_params)
    if not is_valid:
        flash("Chữ ký giao dịch không hợp lệ hoặc dữ liệu bị can thiệp.", "danger")
        return redirect(url_for("payment_failed", message="Chữ ký không hợp lệ"))
        
    # Lấy mã đơn hàng và mã phản hồi từ VNPAY
    ma_giao_dich = vnp_params.get("vnp_TxnRef")
    response_code = vnp_params.get("vnp_ResponseCode")
    
    # Tìm giao dịch trong DB
    giao_dich = GiaoDichNapToken.query.filter_by(ma_giao_dich=ma_giao_dich).first()
    if not giao_dich:
        flash("Không tìm thấy thông tin giao dịch.", "danger")
        return redirect(url_for("payment_failed", message="Không tìm thấy đơn hàng"))
        
    # Tránh xử lý trùng lặp giao dịch (Replay attack / double callback)
    if giao_dich.trang_thai != "cho_thanh_toan":
        if giao_dich.trang_thai == "thanh_cong":
            return redirect(url_for("payment_success", ma_gd=ma_giao_dich))
        else:
            return redirect(url_for("payment_failed", message="Giao dịch đã được xử lý trước đó"))
            
    # Kiểm tra phản hồi thành công từ VNPAY (mã 00)
    if response_code == "00":
        # Cập nhật trạng thái đơn hàng
        giao_dich.trang_thai = "thanh_cong"
        giao_dich.ngay_hoan_thanh = datetime.utcnow()
        
        # Cộng token cho người dùng
        current_user.token += giao_dich.so_token
        db.session.commit()

        # Gửi email thông báo
        try:
            from dich_vu.email_service import gui_email_thanh_toan_thanh_cong
            gui_email_thanh_toan_thanh_cong(
                to_email=current_user.email,
                username=current_user.ten_dang_nhap,
                so_token=giao_dich.so_token,
                so_tien=giao_dich.so_tien,
                ma_giao_dich=giao_dich.ma_giao_dich,
                phuong_thuc=giao_dich.phuong_thuc
            )
        except Exception as e_mail:
            logger.error(f"Error sending payment success email: {e_mail}")

        flash(f"Nạp thành công {giao_dich.so_token} token!", "success")
        return redirect(url_for("payment_success", ma_gd=ma_giao_dich))
    else:
        # Cập nhật trạng thái thất bại / hủy bỏ
        giao_dich.trang_thai = "da_huy" if response_code == "24" else "that_bai"
        giao_dich.ngay_hoan_thanh = datetime.utcnow()
        db.session.commit()
        
        err_msg = "Giao dịch không thành công hoặc đã bị hủy."
        if response_code == "24":
            err_msg = "Khách hàng hủy giao dịch."
        flash(err_msg, "danger")
        return redirect(url_for("payment_failed", message=err_msg))

@app.route("/payment/vnpay_ipn", methods=["GET"])
def vnpay_ipn():
    from dich_vu.vnpay import VNPay
    from mo_hinh import GiaoDichNapToken, NguoiDung
    
    vnp_params = request.args.to_dict()
    
    vnp = VNPay(
        tmn_code=CauHinh.VNPAY_TMN_CODE,
        hash_secret=CauHinh.VNPAY_HASH_SECRET,
        payment_url=CauHinh.VNPAY_PAYMENT_URL
    )
    
    # 1. Xác thực chữ ký số
    is_valid = vnp.verify_payment(vnp_params)
    if not is_valid:
        return jsonify({"RspCode": "97", "Message": "Invalid signature"}), 200
        
    ma_giao_dich = vnp_params.get("vnp_TxnRef")
    response_code = vnp_params.get("vnp_ResponseCode")
    vnp_amount = vnp_params.get("vnp_Amount")
    
    # 2. Tìm giao dịch trong DB
    giao_dich = GiaoDichNapToken.query.filter_by(ma_giao_dich=ma_giao_dich).first()
    if not giao_dich:
        return jsonify({"RspCode": "01", "Message": "Order not found"}), 200
        
    # 3. Kiểm tra số tiền (VNPAY gửi số tiền nhân với 100)
    try:
        expected_amount_cent = int(giao_dich.so_tien) * 100
        if int(vnp_amount) != expected_amount_cent:
            return jsonify({"RspCode": "04", "Message": "Invalid amount"}), 200
    except Exception:
        return jsonify({"RspCode": "04", "Message": "Invalid amount"}), 200
        
    # 4. Kiểm tra trạng thái đơn hàng (Đã xử lý hay chưa)
    if giao_dich.trang_thai != "cho_thanh_toan":
        return jsonify({"RspCode": "02", "Message": "Order already confirmed"}), 200
        
    # 5. Cập nhật kết quả thanh toán
    if response_code == "00":
        giao_dich.trang_thai = "thanh_cong"
        giao_dich.ngay_hoan_thanh = datetime.utcnow()
        
        user = NguoiDung.query.get(giao_dich.nguoi_dung_id)
        if user:
            user.token += giao_dich.so_token
            
        db.session.commit()
        
        # Gửi email thông báo
        try:
            from dich_vu.email_service import gui_email_thanh_toan_thanh_cong
            gui_email_thanh_toan_thanh_cong(
                to_email=user.email if user else None,
                username=user.ten_dang_nhap if user else "N/A",
                so_token=giao_dich.so_token,
                so_tien=giao_dich.so_tien,
                ma_giao_dich=giao_dich.ma_giao_dich,
                phuong_thuc=giao_dich.phuong_thuc
            )
        except Exception as e_mail:
            logger.error(f"Error sending payment success email: {e_mail}")
            
        return jsonify({"RspCode": "00", "Message": "Confirm success"}), 200
    else:
        giao_dich.trang_thai = "da_huy" if response_code == "24" else "that_bai"
        giao_dich.ngay_hoan_thanh = datetime.utcnow()
        db.session.commit()
        return jsonify({"RspCode": "00", "Message": "Confirm success"}), 200

@app.route("/payment/success")
@login_required
def payment_success():
    ma_gd = request.args.get("ma_gd")
    from mo_hinh import GiaoDichNapToken
    
    gd = None
    if ma_gd:
        gd = GiaoDichNapToken.query.filter_by(ma_giao_dich=ma_gd, nguoi_dung_id=current_user.id).first()
        
    if not gd:
        # Phục hồi bằng cách tìm giao dịch thành công mới nhất của người dùng này để tránh lỗi 500
        gd = GiaoDichNapToken.query.filter_by(
            nguoi_dung_id=current_user.id,
            trang_thai="thanh_cong"
        ).order_by(GiaoDichNapToken.id.desc()).first()
        
    if not gd:
        flash("Không tìm thấy thông tin giao dịch.", "warning")
        return redirect(url_for("profile"))
        
    return render_template("payment_success.html", giao_dich=gd)

@app.route("/payment/failed")
@login_required
def payment_failed():
    message = request.args.get("message", "Giao dịch thất bại")
    return render_template("payment_failed.html", message=message)

# -----------------------------------------------------------------------------
from dich_vu.kiem_tra_cau_truc_json import (
    safe_parse_json, 
    safe_section_fix, 
    safe_json_fix, 
    fallback_raw_facts,
    convert_fact_tags_to_html
)

# --- Orchestrator Level Processors ---

def mo_rong_du_lieu_chuong(ma_cv, title, chap_info):
    """
    ACTIVE SEARCH EXPANSION (Placeholder): 
    Dùng để tìm kiếm thêm dữ liệu khi quy mô thực tế không đạt yêu cầu. 
    (Hotfix V5.6: Comment out for safety until full logic is implemented)
    """
    return []

# --- CẤU HÌNH QUY MÔ DỰ ÁN (V17.0+ - Production Grade) ---
CONFIG_QUY_MO = {
    "can_ban": {
        "chapters": (3, 6),       # Căn bản: 3-6 chương
        "parallelism": 5,
        "soft_timeout": 40,
        "hard_timeout": 60,
        "retry": 2
    },
    "tieu_chuan": {
        "chapters": (7, 10),      # Tiêu chuẩn: 7-10 chương
        "parallelism": 6,
        "soft_timeout": 50,
        "hard_timeout": 90,
        "retry": 2
    },
    "chuyen_sau": {
        "chapters": (11, 14),     # Chuyên sâu: 11-14 chương
        "parallelism": 8,
        "soft_timeout": 60,
        "hard_timeout": 120,
        "retry": 3
    }
}

def tinh_so_chuong(quy_mo, documents_count):
    """Tính số chương tối đa có thể theo dữ liệu cào được (Max-First Strategy)."""
    cfg = CONFIG_QUY_MO.get(quy_mo, CONFIG_QUY_MO["tieu_chuan"])
    min_c, max_c = cfg["chapters"]

    # Ít dữ liệu (<6 tài liệu) → chỉ viết min để tránh ảo giác
    if documents_count < 6:
        return min_c
    # Dữ liệu vừa (6-12) → 2/3 khoảng, nghiêng về max
    if documents_count < 12:
        return min_c + (max_c - min_c) * 2 // 3
    # Dữ liệu dồi dào (>=12) → max để tận dụng tối đa
    return max_c


# Quota Guard cho Gemini Free Tier (V21.6: Already initialized at top)

class PipelineContext:
    """
    TRUNG TÂM ĐIỀU PHỐI (CONTEXT OBJECT): 
    Đóng gói toàn bộ metadata của Job để truyền an toàn qua các tầng đa luồng.
    Tránh lỗi NameError và Context Loss vĩnh viễn.
    """
    def __init__(self, ma_cv, tieu_de, quy_mo, api_keys_list, passages_db, global_map, terms, passages, candidates, openai_semaphore, safety_class="SAFE", ngon_ngu="vi", custom_section_words=None, custom_sections_map=None):
        self.ma_cv = ma_cv
        self.tieu_de = tieu_de
        self.quy_mo = quy_mo
        self.api_keys_list = api_keys_list
        self._passages_db = passages_db # Luôn dùng list copy qua property để an toàn (V17.2)
        self.global_map = global_map
        self.terms = terms
        self.passages = passages
        self.candidates = candidates
        self.openai_semaphore = openai_semaphore
        self.safety_class = safety_class
        self.ngon_ngu = ngon_ngu
        self.custom_section_words = custom_section_words
        self.custom_sections_map = custom_sections_map
        self.start_time = time.time()
        self.prefetched_passages = {} # Cache cho Asynchronous Reranking (V35)
    
    @property
    def passages_db(self):
        with PASSAGES_LOCK:
            return list(self._passages_db)
    
    @passages_db.setter
    def passages_db(self, value):
        with PASSAGES_LOCK:
            self._passages_db = value
        
    def get_logger_prefix(self):
        return f"Job {self.ma_cv} | {self.quy_mo.upper()}"

class SectionTaxonomy:
    """
    V8.1: Multi-label Taxonomy Formalization
    """
    FACTUAL_KEYWORDS = ["lịch sử", "diễn biến", "chiến dịch", "kết quả", "nguồn gốc", "nguyên nhân", "cơ sở", "định nghĩa", "tổng quan", "đặc điểm", "thực trạng", "nội dung"]
    ANALYTICAL_KEYWORDS = ["ý nghĩa", "kết luận", "tóm lại", "phân tích", "đánh giá", "bài học", "xu hướng", "nguyên nhân"]
    PROCEDURAL_KEYWORDS = ["cách làm", "phương pháp", "diễn biến", "quy trình", "bước"]

    @classmethod
    def classify(cls, section_title: str) -> list:
        title_lower = section_title.lower()
        labels = []
        if any(k in title_lower for k in cls.FACTUAL_KEYWORDS): labels.append("FACTUAL")
        if any(k in title_lower for k in cls.ANALYTICAL_KEYWORDS): labels.append("ANALYTICAL")
        if any(k in title_lower for k in cls.PROCEDURAL_KEYWORDS): labels.append("PROCEDURAL")
        
        if not labels:
            labels = ["FACTUAL"] # Default an toàn
        return labels

# ─────────────────────────────────────────────────────────────────────────────
# CITATION SAFETY NET (V35.1)
# Quét từng đoạn trong nội dung sau khi AI viết xong.
# Bất kỳ đoạn nào thiếu [ID] sẽ được tự động tra cứu vector và gắn citation phù hợp.
# ─────────────────────────────────────────────────────────────────────────────
def citation_safety_net(section_data: dict, passages: list, api_key: str) -> dict:
    """
    Hậu xử lý (Post-Processing): Quét và tự động gắn citation [ID]
    cho các đoạn văn còn thiếu sau khi AI Writer hoàn thành.
    - Chỉ xử lý đoạn >= 30 ký tự (đủ dài để là đoạn nội dung)
    - Bỏ qua header (###), dòng trống, câu chuyển tiếp ngắn
    - Tìm passage có cosine similarity cao nhất với đoạn văn cần patch
    - Chỉ gắn nếu similarity >= 0.35 (tránh gắn citation sai ngữ nghĩa)
    """
    import re
    import numpy as np
    
    if not passages or not section_data:
        return section_data
    
    content = section_data.get("content", "")
    if not content:
        return section_data
    
    # Chuẩn bị vector pool từ passages đã có sẵn (không gọi API thêm)
    valid_passages = [p for p in passages if "vector" in p and p.get("vector") is not None and p.get("id")]
    if not valid_passages:
        return section_data
    
    passage_vectors = np.array([p["vector"] for p in valid_passages])
    
    def _needs_citation(para: str) -> bool:
        """Kiểm tra đoạn văn có cần gắn thêm citation không."""
        stripped = para.strip()
        if len(stripped) < 30:  # Quá ngắn -> bỏ qua
            return False
        if stripped.startswith("#"):  # Header markdown
            return False
        if re.search(r'\[\d+\]', stripped):  # Đã có citation
            return False
        if re.search(r'<sup class="citation">', stripped):  # APA format
            return False
        if re.search(r'<span class="citation-apa">', stripped):  # APA format v2
            return False
        return True
    
    try:
        from dich_vu.embedding_pool import embedding_pool
        from dich_vu.vector_search import _normalize
        
        lines = content.split("\n")
        patched_lines = []
        patched_count = 0
        
        in_review = False
        for line in lines:
            l_strip = line.strip()
            # Detect review headers
            if re.search(r'(###\s*)?\*{0,2}(Câu hỏi Ôn tập|Bài tập\s*[&＆]\s*Ôn tập|Review Questions)\*{0,2}', l_strip, re.IGNORECASE):
                in_review = True
                
            if in_review:
                # Strip any existing citation tags from the review questions line
                cleaned_line = re.sub(r'\[fact\d+\]|\[\d+\]', '', line)
                patched_lines.append(cleaned_line)
                continue
                
            if not _needs_citation(line):
                patched_lines.append(line)
                continue
            
            # Tạo embedding cho đoạn cần kiểm tra
            try:
                model_name = getattr(CauHinh, "GEMINI_EMBEDDING_MODEL", "gemini-embedding-2")
                embed_res = embedding_pool.embed_content(model_name=model_name, texts=[line[:500]], primary_provider="openai")
                if not embed_res or len(embed_res) == 0:
                    patched_lines.append(line)
                    continue
                
                query_vec = _normalize(np.array(embed_res[0]))
                scores = np.dot(passage_vectors, query_vec)
                best_idx = int(np.argmax(scores))
                best_score = float(scores[best_idx])
                
                # Ngưỡng tin cậy: chỉ gắn nếu đủ độ liên quan
                if best_score >= 0.30:
                    best_id = str(valid_passages[best_idx].get("id", ""))
                    # Gắn citation vào cuối đoạn (trước dấu chấm cuối nếu có)
                    stripped = line.rstrip()
                    if stripped.endswith("."):
                        patched_line = stripped[:-1] + f" [{best_id}]."
                    else:
                        patched_line = stripped + f" [{best_id}]"
                    patched_lines.append(patched_line)
                    patched_count += 1
                else:
                    patched_lines.append(line)
            except Exception:
                patched_lines.append(line)
        
        if patched_count > 0:
            logger.info(f"[CitationNet] Đã tự động gắn {patched_count} citation còn thiếu cho mục '{section_data.get('title', '')}'.")
            section_data["content"] = "\n".join(patched_lines)
    
    except Exception as e:
        logger.warning(f"[CitationNet] Lỗi khi chạy Safety Net: {e}")
    
    return section_data

def process_batch_sections_task(ctx, chap_title, batch_sections_info, mode):
    """Batch-task (V23.2): Biên soạn 3-5 mục cùng lúc để tối ưu API cost và latency."""
    if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
        raise Exception("Tiến trình đã bị người dùng hủy.")
    from dich_vu.kiem_tra_cau_truc_json import safe_section_fix, safe_parse_json
    from dich_vu.openai_da_buoc import viet_noi_dung_batch_sections, viet_rut_gon_rescue
    from dich_vu.audit_service import ScholarlyAuditEngine
    
    # 1. Tìm kiếm facts cho từng section
    relevant_passages_list = []
    dynamic_top_k = {"can_ban": 7, "tieu_chuan": 12, "chuyen_sau": 18}.get(ctx.quy_mo, 12)
    if getattr(ctx, "custom_section_words", 0):
        import math
        dynamic_top_k = max(dynamic_top_k, math.ceil(ctx.custom_section_words / 100))
    
    has_custom = False
    custom_map = getattr(ctx, 'custom_sections_map', None)
    if custom_map:
        has_custom = any(len(secs) > 0 for secs in custom_map.values())
    candidate_count = 50 if has_custom else 30
    
    for s_info in batch_sections_info:
        s_title = s_info.get("title", "Mục mới")
        
        # Thử lấy từ Cache (Asynchronous Pre-fetching V35)
        cache_key = (chap_title, s_title)
        if hasattr(ctx, 'prefetched_passages') and cache_key in ctx.prefetched_passages:
            passages = ctx.prefetched_passages[cache_key]
        else:
            # Fallback nếu RAM rỗng do lỗi
            passages = tim_kiem_vector_with_llm_rerank(
                query=f"{chap_title} {s_title}",
                passages_db=ctx.passages_db,
                api_key=CauHinh.OPENAI_API_KEY,
                top_k=dynamic_top_k,
                candidate_n=min(candidate_count, len(ctx.passages_db)),
                chapter_title=chap_title,
                section_title=s_title
            )
        relevant_passages_list.append(passages)
    
    # 2. Gọi API Batch Writer
    res = {"status": "error"}
    try:
        # 2. Gọi OpenAI Batch API (V23.2 Enhanced)
        logger.info(f"Job {ctx.ma_cv}: BatchWriter active for {len(batch_sections_info)} sections.")
        res = viet_noi_dung_batch_sections(
            chu_de=ctx.tieu_de,
            chapter_title=chap_title,
            sections_info=batch_sections_info,
            relevant_passages_list=relevant_passages_list,
            api_key=CauHinh.OPENAI_API_KEY,
            mode=mode,
            quy_mo=ctx.quy_mo,
            semaphore=ctx.openai_semaphore,
            current_terms=[t.get("term", t) if isinstance(t, dict) else str(t) for t in ctx.terms],
            ngon_ngu=ctx.ngon_ngu,
            custom_section_words=ctx.custom_section_words,
            danh_sach_chuong=getattr(ctx, 'danh_sach_chuong', None),
            custom_sections_map=custom_map
        )
            
        # 3. Fallback logic if batch fails
        if res.get("status") != "success":
            logger.warning(f"Job {ctx.ma_cv}: Batch failed. Falling back to individual.")
            final_batch_data = []
            for i, s_info in enumerate(batch_sections_info):
                sec_data, _ = process_section_task(ctx, chap_title, s_info, "", mode)
                final_batch_data.append(sec_data)
            return final_batch_data, sum(relevant_passages_list, [])
            
    except Exception as e:
        logger.error(f"Critical Batch Error: {e}")

    # 4. Parse & Output
    parsed_batch = []
    if res.get("status") == "success":
        batch_data = safe_parse_json(res["raw_text"])
        if batch_data and "sections" in batch_data:
            llm_sections = batch_data["sections"]
            
            import re
            from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering
            
            def normalize_title_local(t):
                if not t: return ""
                t = clean_title_numbering(t)
                return re.sub(r'[^\w\s]', '', t).lower().strip()

            for i, s_info in enumerate(batch_sections_info):
                s_title = s_info.get("title")
                # 1. Exact Match
                found = next((s for s in llm_sections if s.get("title") == s_title), None)
                # 2. Normalized Title Match
                if not found:
                    norm_target = normalize_title_local(s_title)
                    found = next((s for s in llm_sections if normalize_title_local(s.get("title")) == norm_target), None)
                # 3. Fuzzy/Sub-string Match
                if not found:
                    norm_target = normalize_title_local(s_title)
                    found = next((s for s in llm_sections if norm_target in normalize_title_local(s.get("title")) or normalize_title_local(s.get("title")) in norm_target), None)
                # 4. Index-based Fallback if lengths match
                if not found and len(llm_sections) == len(batch_sections_info):
                    found = llm_sections[i]
                # 5. Index-based Fallback as absolute safety
                if not found and i < len(llm_sections):
                    found = llm_sections[i]

                if found:
                    found["generation_mode"] = mode
                    fixed_sec = safe_section_fix(found, s_title)
                    
                    # Expansion Retry Loop (Giải pháp 1)
                    if getattr(ctx, "custom_section_words", 0) and fixed_sec.get("content"):
                        from dich_vu.gap_filler import dem_so_tu_word
                        from dich_vu.openai_da_buoc import viet_mo_rong_muc
                        
                        target_words = int(ctx.custom_section_words)
                        current_words = dem_so_tu_word(fixed_sec["content"])
                        
                        retry_count = 0
                        while current_words < target_words * 0.75 and retry_count < 3:
                            logger.info(f"Job {ctx.ma_cv}: [{s_title}] Quá ngắn ({current_words}/{target_words}). Kích hoạt Expansion Retry ({retry_count+1}/3).")
                            sec_passages = relevant_passages_list[i] if i < len(relevant_passages_list) else []
                            expanded_text = viet_mo_rong_muc(
                                chu_de=ctx.tieu_de,
                                section_title=s_title,
                                current_content=fixed_sec["content"],
                                relevant_passages=sec_passages,
                                target_words=target_words,
                                api_key=CauHinh.OPENAI_API_KEY,
                                semaphore=ctx.openai_semaphore,
                                ngon_ngu=ctx.ngon_ngu
                            )
                            if expanded_text:
                                fixed_sec["content"] += "\n\n" + expanded_text.strip()
                                current_words = dem_so_tu_word(fixed_sec["content"])
                                logger.info(f"Job {ctx.ma_cv}: [{s_title}] Mở rộng thành công. Số từ mới: {current_words}")
                            else:
                                break
                            retry_count += 1
                                
                    # V35.1: Citation Safety Net - Tự động gắn citation cho đoạn còn thiếu
                    sec_passages = relevant_passages_list[i] if i < len(relevant_passages_list) else []
                    fixed_sec = citation_safety_net(fixed_sec, sec_passages, CauHinh.OPENAI_API_KEY)
                    parsed_batch.append(fixed_sec)
                    
                    # --- V39: DYNAMIC TERM EXTRACTION ---
                    new_terms = found.get("new_terms", [])
                    if new_terms and isinstance(new_terms, list):
                        added_count = 0
                        with TERMS_LOCK:
                            existing_terms_lower = { (t.get("term", t) if isinstance(t, dict) else str(t)).lower() for t in ctx.terms }
                            for nt in new_terms:
                                if isinstance(nt, dict) and "term" in nt and "meaning" in nt:
                                    t_str = str(nt["term"]).strip()
                                    if t_str and t_str.lower() not in existing_terms_lower:
                                        ctx.terms.append({"term": t_str, "meaning": str(nt["meaning"]).strip()})
                                        existing_terms_lower.add(t_str.lower())
                                        added_count += 1
                        if added_count > 0:
                            logger.info(f"Job {ctx.ma_cv}: Added {added_count} new terms during writing of section '{s_title}'.")
                else:
                    logger.warning(f"Section '{s_title}' missing in batch. Rescue triggered.")
                    sec_res = viet_rut_gon_rescue(ctx.tieu_de, s_title, relevant_passages_list[i], CauHinh.OPENAI_API_KEY)
                    rescue_fixed = safe_section_fix(safe_parse_json(sec_res["raw_text"]), s_title)
                    # V35.1: Citation Safety Net cho rescue path
                    sec_passages = relevant_passages_list[i] if i < len(relevant_passages_list) else []
                    rescue_fixed = citation_safety_net(rescue_fixed, sec_passages, CauHinh.OPENAI_API_KEY)
                    parsed_batch.append(rescue_fixed)
                        
    if not parsed_batch:
        for i, s_info in enumerate(batch_sections_info):
            parsed_batch.append({"title": s_info.get("title"), "content": "Lỗi nội dung."})

    # Đảm bảo đúng thứ tự ban đầu và không bao giờ rớt mục nào
    sorted_batch = []
    for i, s_info in enumerate(batch_sections_info):
        found_p = None
        for p in parsed_batch:
            if p.get("title") == s_info.get("title"):
                found_p = p
                break
        if not found_p and i < len(parsed_batch):
            found_p = parsed_batch[i]
            if found_p:
                from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering
                found_p["title"] = clean_title_numbering(s_info.get("title"))
        if found_p:
            sorted_batch.append(found_p)
                
    return sorted_batch, sum(relevant_passages_list, [])

def process_section_task(ctx, chap_title, sec_info, prev_summary, mode):
    """Fallback micro-task: Dùng khi Batch fail hoặc mode đặc biệt."""
    if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
        raise Exception("Tiến trình đã bị người dùng hủy.")
    from dich_vu.kiem_tra_cau_truc_json import safe_section_fix, safe_parse_json
    from dich_vu.openai_da_buoc import viet_noi_dung_muc, viet_rut_gon_rescue
    
    sec_title = sec_info.get("title", "Mục mới")
    
    has_custom = False
    custom_map = getattr(ctx, 'custom_sections_map', None)
    if custom_map:
        has_custom = any(len(secs) > 0 for secs in custom_map.values())
    candidate_count = 50 if has_custom else 30

    # LLM Generative Reranking: Two-Stage Retrieval thay thế Bi-Encoder đơn thuần
    relevant_passages = tim_kiem_vector_with_llm_rerank(
        query=f"{chap_title} {sec_title}",
        passages_db=ctx.passages_db,
        api_key=CauHinh.OPENAI_API_KEY,
        top_k=10,
        candidate_n=min(candidate_count, len(ctx.passages_db)),
        chapter_title=chap_title,
        section_title=sec_title
    )
    
    # 1. Agent 1 (The Writer)
    res = viet_noi_dung_muc(
        ctx.tieu_de, chap_title, sec_title, relevant_passages, CauHinh.OPENAI_API_KEY, 
        mode=mode, quy_mo=ctx.quy_mo, semaphore=ctx.openai_semaphore, ngon_ngu=ctx.ngon_ngu, 
        custom_section_words=ctx.custom_section_words,
        danh_sach_chuong=getattr(ctx, 'danh_sach_chuong', None),
        custom_sections_map=custom_map
    )
    
    if res.get("status") != "success":
        res = viet_rut_gon_rescue(ctx.tieu_de, sec_title, relevant_passages, CauHinh.OPENAI_API_KEY, semaphore=ctx.openai_semaphore)
        
    parsed = safe_parse_json(res["raw_text"])
    found = parsed if isinstance(parsed, dict) else {"title": sec_title, "content": ""}
    
    # Expansion Retry Loop (Giải pháp 1)
    if getattr(ctx, "custom_section_words", 0) and found.get("content"):
        from dich_vu.gap_filler import dem_so_tu_word
        from dich_vu.openai_da_buoc import viet_mo_rong_muc
        
        target_words = int(ctx.custom_section_words)
        current_words = dem_so_tu_word(found["content"])
        
        retry_count = 0
        while current_words < target_words * 0.75 and retry_count < 3:
            logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Quá ngắn ({current_words}/{target_words}). Kích hoạt Expansion Retry ({retry_count+1}/3).")
            expanded_text = viet_mo_rong_muc(
                chu_de=ctx.tieu_de,
                section_title=sec_title,
                current_content=found["content"],
                relevant_passages=relevant_passages,
                target_words=target_words,
                api_key=CauHinh.OPENAI_API_KEY,
                semaphore=ctx.openai_semaphore,
                ngon_ngu=ctx.ngon_ngu
            )
            if expanded_text:
                found["content"] += "\n\n" + expanded_text.strip()
                current_words = dem_so_tu_word(found["content"])
                logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Mở rộng thành công. Số từ mới: {current_words}")
            else:
                break
            retry_count += 1
                
    
    # --- MULTI-AGENT ORCHESTRATION ---
    from dich_vu.gemini_da_buoc import gemini_reviewer_agent
    from dich_vu.openai_da_buoc import openai_editor_agent
    from dich_vu.audit_service import ScholarlyAuditEngine
    
    # 2. Agent 2 (The Reviewer)
    req_cites = [str(p.get("id")) for p in relevant_passages[:10]]
    review_res = gemini_reviewer_agent(
        topic=ctx.tieu_de,
        section_title=sec_title,
        draft_content=json.dumps(found, ensure_ascii=False),
        required_citations=req_cites,
        api_keys=ctx.api_keys_list
    )
    
    if review_res.get("status") == "NEEDS_REVISION":
        logger.warning(f"Job {ctx.ma_cv}: [{sec_title}] Reviewer rejected. Passing to Editor.")
        # 3. Agent 3 (The Editor)
        edit_res = openai_editor_agent(
            chap_title=chap_title,
            section_title=sec_title,
            draft_content=found,
            reviewer_feedback=review_res.get("feedback"),
            passages=relevant_passages,
            api_key=CauHinh.OPENAI_API_KEY,
            semaphore=ctx.openai_semaphore
        )
        if edit_res.get("status") == "success" and "data" in edit_res:
            found = edit_res["data"]
            logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Editor successfully revised the section.")
    
    # Smart Audit (V8.1 Taxonomy)
    labels = SectionTaxonomy.classify(sec_title)
    is_strict_section = "FACTUAL" in labels
    
    if "ANALYTICAL" in labels and not is_strict_section:
        logger.info(f"Job {ctx.ma_cv}: [{sec_title}] Skipping Audit for pure analytical section.")
    else:
        audit_engine = ScholarlyAuditEngine(openai_key=CauHinh.OPENAI_API_KEY, gemini_keys=ctx.api_keys_list)
        audit_engine.run_full_audit(section_data=found, chu_de=ctx.tieu_de, is_strict=is_strict_section)
    
    return safe_section_fix(found, sec_title), relevant_passages

def process_chapter_supervisor(ctx, idx, chap_info, ai_writer_func, giam_sat_func):
    if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
        raise Exception("Tiến trình đã bị người dùng hủy.")
    chap_num = idx + 1; chap_title = chap_info.get("title", f"Chương {chap_num}")
    prefix = ctx.get_logger_prefix()
    logger.info(f"{prefix}: Phase - Chapter {chap_num}")

    # V23.1: Predictive Mode Selection (Tiết kiệm thời gian thử sai)
    kb_density = len(ctx.passages_db) / max(1, CONG_VIEC[ctx.ma_cv].get("tong_chuong", 8))
    # Dùng chuẩn feedback: terms < 6 OR density < 4
    if len(ctx.terms) < 6 or kb_density < 4:
        initial_mode = "SAFE_MINIMAL"
        logger.info(f"{prefix}: Sparse Data (Terms: {len(ctx.terms)}, Density: {kb_density:.1f}). Force SAFE_MINIMAL.")
    else:
        initial_mode = "NORMAL"

    current_mode = initial_mode
    max_attempts = 2
    
    for attempt in range(max_attempts):
        if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
            raise Exception("Tiến trình đã bị người dùng hủy.")
        raw_sections = chap_info.get("sections", [])
        sections = []
        for s in raw_sections:
            if isinstance(s, str):
                sections.append({"title": s, "recommended_pids": []})
            elif isinstance(s, dict):
                sections.append(s)
        
        if getattr(ctx, "custom_section_words", 0) and ctx.custom_section_words >= 600:
            batch_size = 1  # Bắt buộc viết từng mục đơn lẻ để tối ưu độ dài và tránh giới hạn token đầu ra của LLM
        else:
            batch_size = DYNAMIC_BATCH_SIZE.get(ctx.quy_mo, 3)
        section_batches = [sections[i:i + batch_size] for i in range(0, len(sections), batch_size)]
        
        final_chapter_data = []
        all_chapter_passages = []
        
        from concurrent.futures import as_completed, TimeoutError
        
        with ThreadPoolExecutor(max_workers=3) as batch_executor:
            futures = []
            for batch in section_batches:
                f = batch_executor.submit(process_batch_sections_task, ctx, chap_title, batch, current_mode)
                futures.append(f)
            
            try:
                for f in as_completed(futures, timeout=CauHinh.CHAPTER_TIMEOUT):
                    if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
                        raise Exception("Tiến trình đã bị người dùng hủy.")
                    batch_results, passages = f.result()
                    final_chapter_data.extend(batch_results)
                    all_chapter_passages.extend(passages)
            except TimeoutError:
                logger.error(f"{prefix}: Chapter {chap_num} TIMEOUT ({CauHinh.CHAPTER_TIMEOUT}s).")

        # Cleanup & Final Fix
        final_sections = [s for s in final_chapter_data if s]
        fixed = {
            "title": chap_title,
            "sections": final_sections
        }
        
        # Quality Audit Guard (Marginal Pass logic)
        found_ids = set()
        for s in final_sections:
            found_ids.update(re.findall(r'\[(\w+)\]', str(s.get("content", ""))))
        
        # Nếu đạt > 80% coverage hoặc mode là SAFE_MINIMAL -> Accept
        if current_mode == "SAFE_MINIMAL" or len(found_ids) > 5:
            logger.info(f"{prefix}: Chapter {chap_num} SUCCESS (Citations: {len(found_ids)}).")
            CONG_VIEC[ctx.ma_cv]["chuong_hoan_thanh"] = CONG_VIEC[ctx.ma_cv].get("chuong_hoan_thanh", 0) + 1
            return fixed
        else:
            if attempt < max_attempts - 1:
                logger.warning(f"{prefix}: Chapter {chap_num} poor quality (Citations: {len(found_ids)} < 6). Retrying once with SAFE_MINIMAL...")
                current_mode = "SAFE_MINIMAL" # Force for next loop
                continue
            else:
                logger.warning(f"{prefix}: Chapter {chap_num} poor quality after retries.")
                break

    # 3. Kích hoạt Fallback Tối hậu (Nếu Multi-Agent vẫn thất bại)
    logger.warning(f"{prefix}: [FALLBACK] Chapter {chap_num} ('{chap_title}') failed Multi-Agent generation. Falling back to raw facts.")
    try:
        # Nếu có fixed (tức là đã có nội dung nhưng chất lượng kém), ta vẫn trả về fixed
        # Nhưng để an toàn hơn, ta fallback sang raw_facts.
        return fallback_raw_facts(chap_info, all_chapter_passages)
    except Exception as e:
        logger.error(f"{prefix}: [CRITICAL] Fallback for Chapter {chap_num} FAILED: {e}")
        return {"title": chap_title, "sections": []}

def rescue_with_gemini(ctx, chap_info, chap_title, chap_num, prefix, id_to_url=None):
    logger.info(f"{prefix}: OpenAI FAILED. T4 - Gemini Rescue for Chapter {chap_num}")
    CONG_VIEC[ctx.ma_cv]["buoc"] = f"Chương {chap_num}: Cứu hộ định dạng..."
    
    # Circuit Breaker Counter (V17.0+)
    ctx.fail_count = getattr(ctx, 'fail_count', 0) + 1
    total = CONG_VIEC[ctx.ma_cv].get("tong_chuong", 8)
    if ctx.fail_count / total >= 0.5:
        ctx.use_gemini_only = True
        logger.error(f"CIRCUIT BREAKER TRIGGERED: {ctx.fail_count}/{total} fails.")

    rescue_sections = []
    all_passages = [] # Store all found passages for raw fallback (V17.1.7 fix)
    
    has_custom = False
    custom_map = getattr(ctx, 'custom_sections_map', None)
    if custom_map:
        has_custom = any(len(secs) > 0 for secs in custom_map.values())
    candidate_count = 50 if has_custom else 30

    for s_info in chap_info.get("sections", []):
        s_title = s_info.get("title", "Mục mới")
        # LLM Generative Reranking cho Gemini rescue path
        passages = tim_kiem_vector_with_llm_rerank(
            query=f"{chap_title} {s_title}",
            passages_db=ctx.passages_db,
            api_key=CauHinh.OPENAI_API_KEY,
            top_k=10,
            candidate_n=min(candidate_count, len(ctx.passages_db)),
            chapter_title=chap_title,
            section_title=s_title
        )
        all_passages.extend(passages)
        
        # V22 Turbo Throttling
        res = gemini_throttled_call(
            viet_noi_dung_muc_gemini, ctx.tieu_de, chap_title, s_title, passages, 
            api_keys=ctx.api_keys_list,
            danh_sach_chuong=getattr(ctx, 'danh_sach_chuong', None)
        )
        
        parsed = None
        if res["status"] == "success":
            parsed = safe_parse_json(res["raw_text"])
        
        if parsed:
            fixed_sec = safe_section_fix(parsed, s_title)
            if id_to_url:
                sec_cites = set(re.findall(r'\[(\w+)\]', str(fixed_sec.get("content", ""))))
                # id_to_url ở đây là global_map {id: passage_dict}
                fixed_sec["citations"] = [
                    {"id": cid, "url": id_to_url.get(cid, {}).get("url", "")} 
                    for cid in sec_cites if cid in id_to_url
                ]
            rescue_sections.append(fixed_sec)
        else:
            # Local fallback for this specific section
            rescue_sections.append({
                "title": clean_title_numbering(s_title), 
                "content": f"[Nội dung đang được hệ thống xử lý bổ sung từ nguồn cho mục {s_title}...]", 
                "citations": []
            })

    # Return rescued chapter if we have any valid sections, otherwise use raw fallback
    # V29.2: Exclude placeholder text length from total_len calculation to prevent false success.
    total_len = 0
    for s in rescue_sections:
        c = s.get("content", "")
        if "[Nội dung đang được" not in c:
            total_len += len(c)
            
    if rescue_sections and total_len > 50:
        CONG_VIEC[ctx.ma_cv]["chuong_hoan_thanh"] = CONG_VIEC[ctx.ma_cv].get("chuong_hoan_thanh", 0) + 1
        return {
            "title": clean_title_numbering(chap_title),
            "sections": rescue_sections,
            "status": "rescued"
        }
    else:
        logger.warning(f"{prefix}: [RESCUE FAIL] Gemini output too short or empty ({total_len} chars). Falling back to raw facts.")

    # Fallback Logic (Final attempt - Raw Facts)
    return fallback_raw_facts(chap_info, all_passages)

def parallel_generate(ctx, raw_outline, outline_data):
    """Điều phối biên soạn song song các chương với Context Object (ctx)."""
    cfg = CONFIG_QUY_MO.get(ctx.quy_mo, CONFIG_QUY_MO["tieu_chuan"])
    max_workers = cfg["parallelism"]
    
    final_chapters = []
    # Dùng parallelism động theo quy mô (V17.0+)
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = []
        for idx, chap in enumerate(raw_outline):
            time.sleep(random.uniform(0.1, 0.3)) # Giảm jitter để start nhanh hơn
            futures.append(executor.submit(
                process_chapter_supervisor, ctx, idx, chap, openai_writer, giam_sat_chuong
            ))
        final_chapters = [f.result() for f in futures]
    return final_chapters

@app.route('/admin/update_user', methods=['POST'])
@login_required
def admin_update_user():
    if not current_user.la_admin:
        return jsonify({"error": "Unauthorized"}), 403
    
    data = request.get_json()
    u_id = data.get('id')
    user = NguoiDung.query.get(u_id)
    
    if not user:
        return jsonify({"error": "User not found"}), 404
    
    if 'ho_ten' in data: user.ho_ten = data['ho_ten']
    if 'email' in data: user.email = data['email']
    if 'token' in data: user.token = int(data['token'])
    if 'la_admin' in data: user.la_admin = bool(data['la_admin'])
    if 'mat_khau' in data and data['mat_khau'].strip():
        user.mat_khau = generate_password_hash(data['mat_khau'].strip())
    
    db.session.commit()
    return jsonify({"success": True})


@app.post("/huy/<ma_cv>")
def huy_giao_trinh(ma_cv):
    if ma_cv in CONG_VIEC:
        CONG_VIEC[ma_cv]["huy_bo"] = True
        return {"status": "success", "message": "Đã gửi lệnh hủy"}
    return {"status": "error", "message": "Không tìm thấy tiến trình"}, 404

@app.post("/xac_nhan/<ma_cv>/<quyet_dinh>")
def xac_nhan_ha_quy_mo(ma_cv, quyet_dinh):
    if ma_cv not in CONG_VIEC:
        return jsonify({"status": "error", "message": "Không tìm thấy tiến trình"}), 404
    if quyet_dinh == "dong_y":
        data = request.get_json(silent=True) or {}
        if "selected_sections" in data:
            CONG_VIEC[ma_cv]["da_chon_dan_y"] = data.get("selected_sections", [])
        CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = True
        return jsonify({"status": "success", "message": "Đã đồng ý"})
    elif quyet_dinh == "tu_choi":
        CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
        return jsonify({"status": "success", "message": "Đã từ chối"})
    return jsonify({"status": "error", "message": "Quyết định không hợp lệ"}), 400


@app.post("/tao")
def tao_giao_trinh():
    du_lieu = request.get_json(silent=True) or request.form.to_dict()
    tieu_de = (du_lieu.get("tieu_de") or du_lieu.get("title") or "").strip()
    if not tieu_de: return jsonify({"loi": "Thiếu tiêu đề."}), 400
    
    if len(tieu_de) > 100:
        return jsonify({
            "status": "INVALID_INPUT",
            "loi": "Chủ đề quá dài (tối đa 100 ký tự). Vui lòng rút ngắn chủ đề."
        }), 400
    
    if not is_valid_query(tieu_de):
        return jsonify({
            "status": "INVALID_INPUT",
            "loi": "Chủ đề chứa ký tự không hợp lệ. Vui lòng chỉ dùng chữ, số và dấu câu thông thường."
        }), 400
        
    if not is_meaningful(tieu_de):
        if is_gibberish(tieu_de):
            return jsonify({
                "status": "INVALID_INPUT",
                "loi": "Chủ đề dường như được nhập ngẫu nhiên hoặc không có nghĩa. Vui lòng nhập chủ đề rõ ràng."
            }), 400
        return jsonify({
            "status": "INVALID_INPUT",
            "loi": "Chủ đề quá ngắn. Vui lòng nhập từ khóa có ý nghĩa từ 2 ký tự trở lên."
        }), 400

    if is_abbreviation(tieu_de):
        return jsonify({
            "status": "INVALID_INPUT",
            "loi": "Chủ đề chứa từ viết tắt. Hãy nhập đầy đủ nội dung để đạt được kết quả tốt nhất."
        }), 400

    # Lấy các tham số nâng cao (V31+)
    so_chuong_custom = du_lieu.get("so_chuong_custom") or du_lieu.get("custom_so_chuong")
    danh_sach_chuong = du_lieu.get("danh_sach_chuong")
    che_do = du_lieu.get("che_do", "auto")
    approve_outline = bool(du_lieu.get("approve_outline"))

    if danh_sach_chuong:
        if not isinstance(danh_sach_chuong, list):
            return jsonify({
                "status": "INVALID_INPUT",
                "loi": "Danh sách chương phải là một mảng."
            }), 400
        for i, ch in enumerate(danh_sach_chuong, 1):
            if isinstance(ch, dict):
                title = ch.get("title")
                sections = ch.get("sections", [])
                if not isinstance(title, str):
                    return jsonify({
                        "status": "INVALID_INPUT",
                        "loi": f"Tên chương {i} không hợp lệ."
                    }), 400
                ch_strip = title.strip()
                if not isinstance(sections, list):
                    return jsonify({
                        "status": "INVALID_INPUT",
                        "loi": f"Mục cấp 2 của chương {i} phải là một danh sách."
                    }), 400
                for j, s in enumerate(sections, 1):
                    if not isinstance(s, str) or not s.strip():
                        return jsonify({
                            "status": "INVALID_INPUT",
                            "loi": f"Mục cấp 2 thứ {j} của chương {i} không hợp lệ hoặc để trống."
                        }), 400
                    if len(s.strip()) > 150:
                        return jsonify({
                            "status": "INVALID_INPUT",
                            "loi": f"Mục cấp 2 thứ {j} của chương {i} vượt quá giới hạn 150 ký tự."
                        }), 400
                    if is_abbreviation(s.strip()):
                        return jsonify({
                            "status": "INVALID_INPUT",
                            "loi": f"Mục cấp 2 thứ {j} của chương {i} chứa từ viết tắt. Hãy nhập đầy đủ nội dung để đạt được kết quả tốt nhất."
                        }), 400
            elif isinstance(ch, str):
                ch_strip = ch.strip()
            else:
                return jsonify({
                    "status": "INVALID_INPUT",
                    "loi": f"Tên chương {i} không hợp lệ."
                }), 400

            if not ch_strip:
                return jsonify({
                    "status": "INVALID_INPUT",
                    "loi": f"Tên chương {i} không được để trống."
                }), 400
            if len(ch_strip) > 100:
                return jsonify({
                    "status": "INVALID_INPUT",
                    "loi": f"Tên chương {i} vượt quá giới hạn 100 ký tự. Vui lòng rút ngắn lại."
                }), 400
            if is_abbreviation(ch_strip):
                return jsonify({
                    "status": "INVALID_INPUT",
                    "loi": f"Tên chương {i} chứa từ viết tắt. Hãy nhập đầy đủ nội dung để đạt được kết quả tốt nhất."
                }), 400

    from flask_login import current_user
    u_id = current_user.id if current_user.is_authenticated else None

    ma_cv = str(uuid.uuid4())
    CONG_VIEC[ma_cv] = {
        "trang_thai": "dang_chay", 
        "tien_do": 0, 
        "tieu_de": tieu_de, 
        "nhat_ky": [],
        "user_id": u_id,
        "ngay_tao": datetime.utcnow(),
        "cau_hinh_goc": {
            "ngon_ngu": du_lieu.get("ngon_ngu", "vi"),
            "che_do": du_lieu.get("che_do", "auto"),
            "quy_mo": du_lieu.get("quy_mo", "tieu_chuan"),
            "so_chuong_custom": so_chuong_custom,
            "custom_section_words": du_lieu.get("custom_section_words"),
            "manual_titles_toggle": bool(du_lieu.get("danh_sach_chuong")),
            "danh_sach_chuong": du_lieu.get("danh_sach_chuong"),
            "approve_outline": approve_outline
        }
    }

    # Xác định phí token (V32+)
    phi_token = getattr(CauHinh, "PHI_TOKEN_AUTO", 1)
    if che_do == "expert": phi_token = getattr(CauHinh, "PHI_TOKEN_EXPERT", 2)
    elif che_do == "creative": phi_token = getattr(CauHinh, "PHI_TOKEN_CREATIVE", 3)

    # Lấy user_id ngay trong request context trước khi chuyển sang background thread
    from flask_login import current_user
    if current_user.is_authenticated:
        # Đặc quyền Admin: Không giới hạn Token (V33+)
        if not current_user.la_admin:
            if current_user.token < phi_token:
                return jsonify({"loi": f"Bạn cần ít nhất {phi_token} tokens cho chế độ này. Vui lòng mua thêm."}), 403
            current_user.token -= phi_token
            db.session.commit()
    u_id = current_user.id if current_user.is_authenticated else None
    
    # Lấy các tham số nâng cao (V31+)
    so_chuong_custom = du_lieu.get("so_chuong_custom") or du_lieu.get("custom_so_chuong")
    custom_section_words = du_lieu.get("custom_section_words")
    danh_sach_chuong = du_lieu.get("danh_sach_chuong")
    
    # Validate custom_section_words
    if custom_section_words is not None:
        try:
            custom_section_words = int(custom_section_words)
            max_tu = getattr(CauHinh, "MAC_DINH_SO_TU_MAX", 1000)
            if custom_section_words < 100 or custom_section_words > max_tu:
                return jsonify({"loi": f"Độ dài từ mỗi mục phải từ 100 đến {max_tu} từ."}), 400
        except ValueError:
            return jsonify({"loi": "Độ dài từ không hợp lệ."}), 400

    # Validate so_chuong_custom & danh_sach_chuong
    max_so_chuong = getattr(CauHinh, "MAC_DINH_SO_CHUONG_MAX", 15)
    if so_chuong_custom is not None:
        try:
            so_chuong_custom = int(so_chuong_custom)
            if so_chuong_custom < 1 or so_chuong_custom > max_so_chuong:
                return jsonify({"loi": f"Số lượng chương phải từ 1 đến {max_so_chuong} chương."}), 400
        except ValueError:
            return jsonify({"loi": "Số lượng chương không hợp lệ."}), 400

    if danh_sach_chuong:
        if len(danh_sach_chuong) > max_so_chuong:
            return jsonify({"loi": f"Danh sách chương tự chọn vượt quá giới hạn tối đa ({max_so_chuong} chương)."}), 400

    def run_pipeline(user_id, so_chuong_custom=None, danh_sach_chuong=None, custom_section_words=None, approve_outline=False):
        global NEXT_CID
        import time
        start_time = time.time()
        def ghi_nhat_ky(msg):
            from datetime import timedelta
            ts = (datetime.utcnow() + timedelta(hours=7)).strftime("%H:%M:%S")
            elapsed = time.time() - start_time
            log_line = f"[{ts}] {msg} (T+{elapsed:.1f}s)"
            CONG_VIEC[ma_cv]["nhat_ky"].append(log_line)
            logger.info(f"[Job {ma_cv[:8]}] {msg} (T+{elapsed:.1f}s)")
            # Watchdog Check (V22.2)
            if elapsed > 3600: # 60 Minutes Extended Limit for Chuyen Sau + Fallback
                logger.error(f"WATCHDOG TRIGGERED for {ma_cv}: Elapsed {elapsed:.1f}s > 3600s. Emergency termination/fallback.")
                raise TimeoutError("Pipeline Watchdog limit reached.")

        with app.app_context():
            try:
                from dich_vu.meta_controller import meta_controller_instance
                meta_controller_instance.reset_state()
                from dich_vu.embedding_pool import embedding_pool
                embedding_pool.reset_pool_status()
                def check_cancel():
                    if CONG_VIEC.get(ma_cv, {}).get("huy_bo"):
                        raise Exception("Tiến trình đã bị người dùng hủy.")

                
                has_custom_subsections = False
                custom_sections_map = {}
                if danh_sach_chuong:
                    normalized_danh_sach_chuong = []
                    for ch in danh_sach_chuong:
                        if isinstance(ch, dict):
                            ch_title = ch.get("title", "").strip()
                            sections = ch.get("sections", [])
                            normalized_danh_sach_chuong.append(ch_title)
                            if sections:
                                has_custom_subsections = True
                                custom_sections_map[ch_title] = [s.strip() for s in sections if s.strip()]
                        else:
                            normalized_danh_sach_chuong.append(ch.strip())
                    danh_sach_chuong = normalized_danh_sach_chuong

                ghi_nhat_ky(f"Khởi động Pipeline AI cho chủ đề: {tieu_de}")
                from dich_vu.lay_wikipedia import tao_tai_lieu_tu_wikipedia, ekre_discovery_engine
                from dich_vu.lam_sach_van_ban import chia_doan, lam_sach_trang

                quy_mo = du_lieu.get("quy_mo", "tieu_chuan")
                # V39: Map numeric form values to internal string keys
                QUY_MO_MAP = {"1": "can_ban", "2": "tieu_chuan", "3": "chuyen_sau"}
                if quy_mo in QUY_MO_MAP:
                    quy_mo = QUY_MO_MAP[quy_mo]
                ngon_ngu = du_lieu.get("ngon_ngu", "vi")
                CONG_VIEC[ma_cv]["ngon_ngu"] = ngon_ngu
                logger.info(f"Job {ma_cv}: Pipeline started with SCALE={quy_mo}, LANG={ngon_ngu}, CUSTOM_CH={so_chuong_custom}, MANUAL_LIST={'YES' if danh_sach_chuong else 'NO'} for topic '{tieu_de}'")

                # --- Step 0: Safety Classification (V29 — 3-Layer) ---
                from dich_vu.safety_router import classify_topic, reframe_topic, generate_safe_title, get_block_message
                safety_res = classify_topic(tieu_de, CauHinh.OPENAI_API_KEY)
                safety_class = safety_res.get("classification", "SAFE")
                
                if safety_class in ["BLOCK", "BLOCK_LANG"]:
                    logger.error(f"Job {ma_cv}: Topic '{tieu_de}' BLOCKED (Layer: {safety_res.get('layer')}). Reason: {safety_res.get('reason')}")
                    block_msg = get_block_message(safety_res)
                    error_text = block_msg["message"] if block_msg else safety_res.get("reason")
                    if block_msg and block_msg.get("suggestion"):
                        error_text += f"\n\n{block_msg['suggestion']}"
                    CONG_VIEC[ma_cv]["loi"] = error_text
                    CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                    CONG_VIEC[ma_cv]["loai_loi"] = "safety_block"
                    return
                
                ekre_query = tieu_de
                if safety_class == "REFRAME":
                    ekre_query = reframe_topic(tieu_de)
                    CONG_VIEC[ma_cv]["tieu_de"] = generate_safe_title(tieu_de)
                    ghi_nhat_ky(f"Chủ đề nhạy cảm. Đã chuyển sang phân tích học thuật: {ekre_query}")
                    logger.info(f"[SAFETY] Reframed: '{tieu_de}' → '{ekre_query}'")

                # --- V37.2: OUTLINE VALIDATION (Safety & Relevance) ---
                if danh_sach_chuong and isinstance(danh_sach_chuong, list) and len(danh_sach_chuong) > 0:
                    try:
                        _clean_chapters = [ch.strip() for ch in danh_sach_chuong if ch.strip()]
                        
                        # 1. Kiểm tra An toàn (SafetyRouter) chạy song song
                        ghi_nhat_ky("🔍 Đang kiểm tra tính an toàn của danh sách chương tự chọn...")
                        from concurrent.futures import ThreadPoolExecutor
                        
                        unsafe_chapters = []
                        with ThreadPoolExecutor(max_workers=5) as executor:
                            futures = {executor.submit(classify_topic, ch, CauHinh.OPENAI_API_KEY): ch for ch in _clean_chapters}
                            for future in futures:
                                ch_name = futures[future]
                                try:
                                    res = future.result()
                                    if res.get("classification") in ["BLOCK", "BLOCK_LANG"]:
                                        block_msg = get_block_message(res)
                                        reason = res.get("reason", "Nội dung vi phạm.")
                                        if block_msg:
                                            reason = f"{block_msg['title']}: {block_msg['message']}"
                                        unsafe_chapters.append((ch_name, reason))
                                except Exception as e:
                                    logger.warning(f"[ChapterSafety] Error checking '{ch_name}': {e}")
                        
                        if unsafe_chapters:
                            off_names = "\n".join([f"  • '{c[0]}' ({c[1]})" for c in unsafe_chapters])
                            error_msg = f"⛔ Phát hiện tên chương chứa nội dung không an toàn:\n{off_names}\n\nVui lòng chỉnh sửa lại để tiếp tục."
                            ghi_nhat_ky(f"❌ Phát hiện {len(unsafe_chapters)} chương vi phạm chính sách an toàn.")
                            CONG_VIEC[ma_cv]["loi"] = error_msg
                            CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                            CONG_VIEC[ma_cv]["loai_loi"] = "safety_block"
                            return
                        
                        # 2. Kiểm tra Relevance (Topic-Outline Semantic Drift)
                        ghi_nhat_ky("🔍 Đang kiểm tra tên chương có phù hợp với chủ đề...")
                        from openai import OpenAI as _OAI_Validator
                        import json as _json_val
                        
                        _val_client = _OAI_Validator(api_key=CauHinh.OPENAI_API_KEY, max_retries=1)
                        
                        _val_prompt = f"""Bạn là trợ lý kiểm tra nội dung giáo trình.

Chủ đề giáo trình: "{tieu_de}"
Danh sách tên chương do người dùng đặt: {_json_val.dumps(_clean_chapters, ensure_ascii=False)}

Nhiệm vụ: Kiểm tra từng tên chương có liên quan trực tiếp đến chủ đề "{tieu_de}" hay không.
- Liên quan trực tiếp = chương đó là một khía cạnh, thuộc tính, lịch sử, ứng dụng, nền tảng cơ sở (ví dụ: Toán học, Vật lý... nếu áp dụng được), hoặc chủ đề con của "{tieu_de}".
- KHÔNG liên quan = chương đó nói về một lĩnh vực hoàn toàn lạc đề (ví dụ: truyện tranh, giải trí không liên quan), vô nghĩa hoặc không có bất kỳ mối liên hệ học thuật nào với "{tieu_de}".
Lưu ý: Hãy đánh giá thoáng và mở rộng tư duy. Các ngành khoa học nền tảng, công cụ, hoặc khái quát (như Giới thiệu, Toán học, Lịch sử...) luôn được tính là có liên quan.

Trả về ĐÚNG 1 JSON array boolean, mỗi phần tử tương ứng 1 chương. true = liên quan, false = không liên quan.
Ví dụ: [true, false, true]
CHỈ trả về JSON array, KHÔNG giải thích."""

                        _val_resp = _val_client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "user", "content": _val_prompt}],
                            temperature=0,
                            max_tokens=100
                        )
                        import re as _re
                        _val_text = _val_resp.choices[0].message.content.strip()
                        _val_text = _re.sub(r"^```(?:json)?\s*", "", _val_text, flags=_re.IGNORECASE)
                        _val_text = _re.sub(r"\s*```$", "", _val_text).strip()
                        
                        # Parse JSON response
                        _val_results = _json_val.loads(_val_text)
                        
                        if isinstance(_val_results, list) and len(_val_results) == len(_clean_chapters):
                            off_topic_chapters = []
                            for i, (ch_name, is_relevant) in enumerate(zip(_clean_chapters, _val_results)):
                                if is_relevant:
                                    logger.info(f"[SemanticValidator] OK: '{ch_name}' — relevant to '{tieu_de}'")
                                else:
                                    off_topic_chapters.append(ch_name)
                                    logger.warning(f"[SemanticValidator] BLOCKED: '{ch_name}' — NOT relevant to '{tieu_de}'")
                            
                            if off_topic_chapters:
                                off_names = "\n".join([f"  • '{c}'" for c in off_topic_chapters])
                                error_msg = f"⚠️ Một số tên chương không liên quan đến chủ đề '{tieu_de}':\n{off_names}\n\nVui lòng chỉnh sửa tên chương cho phù hợp với chủ đề rồi thử lại."
                                ghi_nhat_ky(f"❌ Tên chương không phù hợp: {', '.join(off_topic_chapters)}")
                                CONG_VIEC[ma_cv]["loi"] = error_msg
                                CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                                CONG_VIEC[ma_cv]["loai_loi"] = "semantic_drift"
                                return
                            else:
                                ghi_nhat_ky("✅ Tất cả tên chương phù hợp với chủ đề.")
                        else:
                            logger.warning(f"[SemanticValidator] Invalid response format: {_val_text}. Skipping.")
                            
                    except Exception as val_err:
                        logger.warning(f"[SemanticValidator] Validation failed safely: {val_err}. Skipping check.")

                # Bước 1: Khám phá tri thức (AKRE Discovery)
                check_cancel()
                CONG_VIEC[ma_cv].update({"tien_do": 10, "buoc": "Đang thực hiện AKRE Discovery..."})
                ghi_nhat_ky("Bắt đầu pha Discovery (AKRE Adaptive Harvesting).")
                # Bắt đầu pha Discovery (EKRE Adaptive Harvesting)
                from dich_vu.lay_wikipedia import ekre_discovery_engine
                
                # Nếu có mục cấp 2 do người dùng tự định nghĩa, dùng cả tên chương và mục nhỏ làm hints tìm kiếm Wikipedia
                if has_custom_subsections:
                    ekre_hints = []
                    for ch_title in danh_sach_chuong:
                        if ch_title not in ekre_hints:
                            ekre_hints.append(ch_title)
                        for sec_title in custom_sections_map.get(ch_title, []):
                            if sec_title not in ekre_hints:
                                ekre_hints.append(sec_title)
                    if not ekre_hints:
                        ekre_hints = danh_sach_chuong
                else:
                    ekre_hints = danh_sach_chuong

                # EKRE trả về passages, candidates, hardened_docs, và xray
                ekre_res = ekre_discovery_engine(
                    ekre_query, 
                    api_keys_list=CauHinh.GEMINI_API_KEYS,
                    quy_mo=quy_mo,
                    api_key_openai=CauHinh.OPENAI_API_KEY,
                    original_topic=tieu_de,
                    chapter_hints=ekre_hints,
                    ngon_ngu=ngon_ngu,  # V44: Language-aware discovery
                    custom_section_words=custom_section_words,
                    check_cancel=check_cancel
                )
                passages = ekre_res.get("passages", [])
                candidates = ekre_res.get("candidates", {})
                hardened_docs = ekre_res.get("hardened_docs", [])
                xray = ekre_res.get("xray", {})
                
                xray["safety_class"] = safety_class
                # V29: Lưu danh sách nguồn với đầy đủ metadata + điểm chất lượng
                source_list = []
                for doc in hardened_docs:
                    source_list.append({
                        "title": doc.get("title", "N/A"),
                        "url": doc.get("url", ""),
                        "lang": doc.get("lang", "vi"),
                        "reason": doc.get("subtopic", ""),
                        "quality_score": round(doc.get("quality_score", 0), 2),
                        "relevance_score": round(doc.get("relevance_score", 0), 3),
                        "text_len": len(doc.get("text", "")),
                        "is_low_priority": doc.get("is_low_priority", False),
                    })
                # Sắp xếp theo quality_score giảm dần
                source_list.sort(key=lambda x: x["quality_score"], reverse=True)
                CONG_VIEC[ma_cv]["top_30_links"] = source_list[:30]
                CONG_VIEC[ma_cv]["discovery_xray"] = xray 
                
                # 💎 STRUCTURE LOGGING (V24.1)
                kb_density = len(passages) / 10 # heuristic density per chapter
                logger.info(f"[STRUCTURE] Scale: {quy_mo} | SearchYield: {xray['stats']['filtered']} | Density: {kb_density:.1f}")
                ghi_nhat_ky(f"Discovery hoàn tất. Tìm thấy {len(candidates)} trang nguồn. (X-Ray Yield: {xray['stats']['filtered']}/{xray['stats']['retrieved']})")
                
                # --- CHUẨN BỊ ĐÁNH GIÁ ĐỦ DỮ LIỆU ---
                from dich_vu.lay_wikipedia import score_knowledge_base
                kb_score = score_knowledge_base(hardened_docs)
                # V35.1: Ngưỡng đồng bộ với công thức mới
                SUFFICIENCY_THRESHOLDS = {
                    "can_ban": 10,
                    "tieu_chuan": 20,
                    "chuyen_sau": 45
                }
                SCALE_DOWNGRADE = {
                    "chuyen_sau": "tieu_chuan",
                    "tieu_chuan": "can_ban"
                }
                SCALE_LABELS = {
                    "can_ban": "Căn bản (4-5 chương)",
                    "tieu_chuan": "Tiêu chuẩn (7-10 chương)",
                    "chuyen_sau": "Chuyên sâu (12-20 chương)"
                }
                if custom_section_words:
                    words_val = int(custom_section_words)
                    if words_val <= 300:
                        min_threshold = 10
                    elif words_val <= 600:
                        min_threshold = 20
                    else:
                        min_threshold = 45
                else:
                    min_threshold = SUFFICIENCY_THRESHOLDS.get(quy_mo, 30)

                # --- PHANH AN TOÀN (RELIABLE SOURCE GATE) ---
                confidence = xray.get("adaptive", {}).get("confidence_score", 0)
                reliable_docs = [
                    d for d in hardened_docs
                    if d.get("quality_score", 0) >= CauHinh.EKRE_MIN_QUALITY_FLOOR
                    and not d.get("is_low_priority", False)
                ]

                reason = None
                if len(reliable_docs) == 0:
                    if xray.get("adaptive", {}).get("stop_reason") == "EMBEDDING_API_FAILED":
                        reason = "API_QUOTA_EXHAUSTED"
                    else:
                        reason = "NO_RELIABLE_DOCS"
                elif confidence < 0.25:
                    reason = "LOW_CONFIDENCE"

                if reason:
                    logger.warning(f"[NO_RELIABLE] Query: {ekre_query} | Confidence: {confidence} | Reason: {reason}")
                    
                    sorted_docs = sorted(hardened_docs, key=lambda x: x.get("quality_score", 0), reverse=True)
                    preview = []
                    for d in sorted_docs[:5]:
                        preview.append({
                            "title": d.get("title", ""),
                            "similarity": round(d.get("relevance_score", 0), 3),
                            "quality": round(d.get("quality_score", 0), 2),
                            "final_score": round(d.get("quality_score", 0), 3)
                        })
                        
                    # Phân tích scale hạ cấp
                    suggested_scale = None
                    if reason not in ("NO_RELIABLE_DOCS", "LOW_CONFIDENCE") and quy_mo in SCALE_DOWNGRADE:
                        lower = SCALE_DOWNGRADE[quy_mo]
                        lower_threshold = SUFFICIENCY_THRESHOLDS[lower]
                        if kb_score >= lower_threshold:
                            suggested_scale = lower
                    
                    if reason == "API_QUOTA_EXHAUSTED":
                        error_msg = "Tất cả API Key Gemini đã cạn kiệt Quota. Không thể chuyển đổi Vector tài liệu. Vui lòng thêm API Key mới hoặc thử lại sau."
                    else:
                        error_msg = "Thông tin tìm kiếm của bạn trên Wikipedia không đủ để viết một bài giáo trình."
                        
                        # --- Gọi OpenAI để gợi ý chủ đề ---
                        try:
                            from openai import OpenAI
                            o_client = OpenAI(api_key=api_key_openai)
                            prompt_suggest = f"Người dùng muốn tạo giáo trình về chủ đề '{tieu_de}', nhưng hệ thống không tìm đủ tài liệu chất lượng trên Wikipedia. Hãy gợi ý 3 chủ đề tương tự, phổ biến hơn, hoặc mang tính học thuật cao hơn để thay thế. Chỉ trả về một mảng JSON các chuỗi: [\"Chủ đề 1\", \"Chủ đề 2\", \"Chủ đề 3\"]. KHÔNG BÌNH LUẬN."
                            resp_suggest = o_client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{"role": "user", "content": prompt_suggest}],
                                temperature=0.7,
                                max_tokens=150
                            )
                            import re as _re
                            text_suggest = resp_suggest.choices[0].message.content
                            text_suggest = _re.sub(r"^```(?:json)?\s*", "", text_suggest.strip(), flags=_re.IGNORECASE)
                            text_suggest = _re.sub(r"\s*```$", "", text_suggest.strip()).strip()
                            suggestions = json.loads(text_suggest)
                            if suggestions and isinstance(suggestions, list):
                                error_msg += f"<br><br>💡 <b>Gợi ý chủ đề tương đương có nhiều dữ liệu hơn:</b><br>"
                                for s in suggestions[:3]:
                                    error_msg += f"- {s}<br>"
                        except Exception as e:
                            logger.error(f"[Suggestion] Error: {e}")
                            pass
                    
                    if suggested_scale:
                        error_msg_html = f"Thông tin tìm kiếm của bạn trên Wikipedia không đủ để viết một bài giáo trình ở mức {SCALE_LABELS.get(quy_mo, quy_mo)}.<br><br>Tuy nhiên, bạn có thể tiếp tục tạo giáo trình ở mức quy mô thấp hơn: <b>{SCALE_LABELS.get(suggested_scale, suggested_scale)}</b>."
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "cho_xac_nhan",
                            "loai_loi": "INSUFFICIENT_DATA",
                            "loi": error_msg_html,
                            "chi_tiet": reason,
                            "suggestions": preview,
                            "suggested_scale": suggested_scale,
                            "suggested_scale_label": SCALE_LABELS.get(suggested_scale, suggested_scale),
                            "kb_score": round(kb_score, 1),
                            "required_score": min_threshold,
                            "current_scale": quy_mo,
                            "current_scale_label": SCALE_LABELS.get(quy_mo, quy_mo),
                            "xac_nhan_cho_phep": None
                        })
                        
                        wait_seconds = 120
                        start_wait = time.time()
                        ghi_nhat_ky(f"Chờ người dùng xác nhận hạ quy mô từ {quy_mo} xuống {suggested_scale}...")
                        while CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is None:
                            time.sleep(1.0)
                            if CONG_VIEC[ma_cv].get("huy_bo"):
                                CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                break
                            if time.time() - start_wait > wait_seconds:
                                CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                break
                                
                        if CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is True:
                            ghi_nhat_ky(f"Người dùng đã đồng ý hạ quy mô xuống {suggested_scale}. Tiếp tục biên soạn.")
                            quy_mo = suggested_scale
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "dang_chay",
                                "quy_mo": quy_mo,
                                "loi": None,
                                "suggested_scale": None,
                                "xac_nhan_cho_phep": None
                            })
                            du_lieu["quy_mo"] = quy_mo
                            min_threshold = SUFFICIENCY_THRESHOLDS.get(quy_mo, 30)
                        else:
                            ghi_nhat_ky("Tiến trình bị dừng do từ chối hạ quy mô hoặc hết thời gian chờ.")
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "that_bai",
                                "loi": "Người dùng đã từ chối hạ quy mô giáo trình. Tiến trình bị hủy."
                            })
                            return
                    else:
                        error_msg += "\n\nHệ thống bắt buộc phải dừng để đảm bảo chất lượng. Vui lòng thử một chủ đề phổ biến hơn."
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "that_bai",
                            "loai_loi": "NO_RELIABLE_SOURCES",
                            "loi": error_msg,
                            "chi_tiet": reason,
                            "suggestions": preview,
                            "suggested_scale": None,
                            "suggested_scale_label": None,
                            "kb_score": round(kb_score, 1),
                            "required_score": min_threshold,
                            "current_scale": quy_mo,
                            "current_scale_label": SCALE_LABELS.get(quy_mo, quy_mo),
                            "query_suggestions": [
                                f"{tieu_de} là gì",
                                f"{tieu_de} cơ bản",
                                f"Tổng quan về {tieu_de}"
                            ]
                        })
                        return
                
                # --- DATA SUFFICIENCY GATE (V31 — Kiểm tra đủ dữ liệu cho quy mô) ---
                
                # --- TOPIC PRESENCE CHECK (V31.8 — Diacritics-aware) ---
                # Loại bỏ từ đệm phổ biến để lấy cụm danh từ cốt lõi
                from dich_vu.lam_sach_van_ban import remove_diacritics
                VIET_STOPWORDS = {
                    "tổng", "quan", "về", "giới", "thiệu", "cơ", "bản", "nâng", "cao",
                    "chuyên", "sâu", "nhập", "môn", "đại", "cương", "khái", "niệm",
                    "của", "và", "các", "trong", "cho", "với", "từ", "đến", "là",
                    "một", "những", "được", "có", "này", "theo", "tại", "trên",
                    # Bản không dấu của stopwords
                    "tong", "ve", "gioi", "co", "ban", "nang",
                    "chuyen", "nhap", "dai", "cuong", "khai",
                    "cua", "va", "cac", "trong", "voi", "tu", "den", "la",
                    "mot", "nhung", "duoc", "nay", "tai"
                }
                topic_lower = tieu_de.lower().strip()
                # Lấy cụm danh từ cốt lõi (bỏ stopwords + từ quá ngắn)
                core_words = [w for w in topic_lower.split() if w not in VIET_STOPWORDS and len(w) > 1]
                core_phrase = " ".join(core_words)
                core_phrase_nodiac = remove_diacritics(core_phrase).lower()
                
                # Sinh tất cả subphrases liên tiếp ≥2 từ (sắp xếp từ dài → ngắn)
                # Ví dụ: "nguyên lí bom nguyên tử" → ["nguyên lí bom nguyên tử", "nguyên lí bom nguyên", "lí bom nguyên tử", "nguyên lí bom", "bom nguyên tử", ...]
                subphrases = []
                if len(core_words) >= 2:
                    for length in range(len(core_words), 1, -1):  # Từ dài đến ngắn
                        for start in range(len(core_words) - length + 1):
                            sp = " ".join(core_words[start:start+length])
                            subphrases.append(sp)
                elif core_phrase:
                    subphrases = [core_phrase]
                
                subphrases_nodiac = [remove_diacritics(sp).lower() for sp in subphrases]
                
                logger.info(f"[DATA-GATE] Core: '{core_phrase}' | Subphrases: {subphrases[:5]}...")
                
                docs_mentioning_topic = 0
                for d in hardened_docs:
                    doc_text = (d.get("text", "") + " " + d.get("title", "")).lower()
                    doc_text_nodiac = remove_diacritics(doc_text)
                    # Match BẤT KỲ subphrase nào (có dấu hoặc không dấu)
                    matched = any(
                        sp in doc_text or sp_nd in doc_text_nodiac
                        for sp, sp_nd in zip(subphrases, subphrases_nodiac)
                    )
                    if matched:
                        docs_mentioning_topic += 1
                
                topic_presence_ratio = docs_mentioning_topic / max(len(hardened_docs), 1)
                avg_sim = xray.get("stats", {}).get("avg_sim", 1.0)
                
                logger.info(
                    f"[DATA-GATE] kb_score={kb_score:.1f} | threshold={min_threshold} | scale={quy_mo} | "
                    f"topic_presence={docs_mentioning_topic}/{len(hardened_docs)} ({topic_presence_ratio:.0%}) | "
                    f"avg_sim={avg_sim:.3f}"
                )
                
                # DỪNG chỉ khi CẢ HAI: (1) phrase không tìm thấy VÀ (2) avg_sim thấp
                # → "phim ma" tìm được "phim kinh dị" (avg_sim cao) → KHÔNG chặn
                # → "phi công" tìm được IT (avg_sim thấp + phrase 0%) → CHẶN
                # V31.8.1: Nới lỏng - Nếu có ít nhất 1 bài khớp topic, chỉ chặn nếu tập tài liệu quá tệ (avg_sim < 0.35)
                is_topic_absent = (docs_mentioning_topic == 0) or (topic_presence_ratio < 0.10 and avg_sim < 0.35)
                if is_topic_absent and avg_sim < 0.45:
                    reason_detail = "TOPIC_NOT_FOUND" if avg_sim < 0.35 else "LOW_RELEVANCE"
                    logger.warning(
                        f"[DATA-GATE] {reason_detail}: '{tieu_de}' (core='{core_phrase}') "
                        f"presence={docs_mentioning_topic}/{len(hardened_docs)}, avg_sim={avg_sim:.3f}"
                    )
                    kb_score = 0  # Ép điểm về 0 để kích hoạt gate bên dưới
                
                if kb_score < min_threshold:
                    # Tùy chọn hạ độ dài mục tự chọn (custom_section_words)
                    suggested_words = None
                    if custom_section_words:
                        words_val = int(custom_section_words)
                        if words_val > 600 and kb_score >= 20:
                            suggested_words = 400
                        elif words_val > 300 and kb_score >= 10:
                            suggested_words = 150
                    
                    if custom_section_words and suggested_words:
                        error_msg_html = (
                            f"Chủ đề <b>'{tieu_de}'</b> khan hiếm dữ liệu trên Wikipedia.<br>"
                            f"Số lượng tài liệu hiện tại không đủ để viết bài chi tiết ở mức <b>{custom_section_words} từ mỗi mục</b> như yêu cầu.<br><br>"
                            f"Bạn có đồng ý hạ độ dài xuống <b>{suggested_words} từ mỗi mục</b> (khoảng {suggested_words * 6} ký tự) không?"
                        )
                        if alt_topics:
                            error_msg_html += f"<br><br>💡 <b>Gợi ý chủ đề tương đương có nhiều dữ liệu hơn:</b><br>"
                            for s in alt_topics:
                                error_msg_html += f"- {s}<br>"
                                
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "cho_xac_nhan",
                            "loai_loi": "CUSTOM_WORDS_DOWNGRADE",
                            "loi": error_msg_html,
                            "suggested_words": suggested_words,
                            "requested_words": custom_section_words,
                            "xac_nhan_cho_phep": None
                        })
                        
                        wait_seconds = 120
                        start_wait = time.time()
                        ghi_nhat_ky(f"Chờ người dùng xác nhận hạ độ dài mục từ {custom_section_words} từ xuống {suggested_words} từ...")
                        while CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is None:
                            time.sleep(1.0)
                            if CONG_VIEC[ma_cv].get("huy_bo"):
                                CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                break
                            if time.time() - start_wait > wait_seconds:
                                CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                break
                                
                        if CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is True:
                            ghi_nhat_ky(f"Người dùng đã đồng ý hạ độ dài mục xuống {suggested_words} từ. Tiếp tục biên soạn.")
                            custom_section_words = suggested_words
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "dang_chay",
                                "loi": None,
                                "suggested_words": None,
                                "xac_nhan_cho_phep": None
                            })
                            ctx.custom_section_words = custom_section_words
                        else:
                            ghi_nhat_ky("Tiến trình bị dừng do từ chối hạ độ dài mục hoặc hết thời gian chờ.")
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "that_bai",
                                "loi": "Người dùng đã từ chối hạ độ dài giáo trình. Tiến trình bị hủy."
                            })
                            return
                    else:
                        suggested_scale = None
                        if not is_topic_absent and topic_presence_ratio >= 0.10:
                            if quy_mo in SCALE_DOWNGRADE:
                                lower = SCALE_DOWNGRADE[quy_mo]
                                lower_threshold = SUFFICIENCY_THRESHOLDS[lower]
                                if kb_score >= lower_threshold:
                                    suggested_scale = lower
                                    
                        if suggested_scale:
                            error_msg_html = f"Thông tin tìm kiếm của bạn trên Wikipedia không đủ để viết một bài giáo trình ở mức {SCALE_LABELS.get(quy_mo, quy_mo)}.<br><br>Tuy nhiên, bạn có thể tiếp tục tạo giáo trình ở mức quy mô thấp hơn: <b>{SCALE_LABELS.get(suggested_scale, suggested_scale)}</b>."
                            if alt_topics:
                                error_msg_html += f"<br><br>💡 <b>Gợi ý chủ đề tương đương có nhiều dữ liệu hơn:</b><br>"
                                for s in alt_topics:
                                    error_msg_html += f"- {s}<br>"
                                    
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "cho_xac_nhan",
                                "loai_loi": "INSUFFICIENT_DATA",
                                "loi": error_msg_html,
                                "suggested_scale": suggested_scale,
                                "suggested_scale_label": SCALE_LABELS.get(suggested_scale, suggested_scale),
                                "kb_score": round(kb_score, 1),
                                "required_score": min_threshold,
                                "current_scale": quy_mo,
                                "current_scale_label": SCALE_LABELS.get(quy_mo, quy_mo),
                                "xac_nhan_cho_phep": None
                            })
                            
                            wait_seconds = 120
                            start_wait = time.time()
                            ghi_nhat_ky(f"Chờ người dùng xác nhận hạ quy mô từ {quy_mo} xuống {suggested_scale}...")
                            while CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is None:
                                time.sleep(1.0)
                                if CONG_VIEC[ma_cv].get("huy_bo"):
                                    CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                    break
                                if time.time() - start_wait > wait_seconds:
                                    CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                                    break
                                    
                            if CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is True:
                                ghi_nhat_ky(f"Người dùng đã đồng ý hạ quy mô xuống {suggested_scale}. Tiếp tục biên soạn.")
                                quy_mo = suggested_scale
                                if custom_section_words:
                                    if quy_mo == "can_ban":
                                        custom_section_words = 150
                                    elif quy_mo == "tieu_chuan":
                                        custom_section_words = 400
                                    ctx.custom_section_words = custom_section_words
                                    
                                CONG_VIEC[ma_cv].update({
                                    "trang_thai": "dang_chay",
                                    "quy_mo": quy_mo,
                                    "loi": None,
                                    "suggested_scale": None,
                                    "xac_nhan_cho_phep": None
                                })
                                du_lieu["quy_mo"] = quy_mo
                                min_threshold = SUFFICIENCY_THRESHOLDS.get(quy_mo, 30)
                            else:
                                ghi_nhat_ky("Tiến trình bị dừng do từ chối hạ quy mô hoặc hết thời gian chờ.")
                                CONG_VIEC[ma_cv].update({
                                    "trang_thai": "that_bai",
                                    "loi": "Người dùng đã từ chối hạ quy mô giáo trình. Tiến trình bị hủy."
                                })
                                return
                        else:
                            CONG_VIEC[ma_cv].update({
                                "trang_thai": "that_bai",
                                "loai_loi": "INSUFFICIENT_DATA",
                                "loi": error_msg,
                                "kb_score": round(kb_score, 1),
                                "required_score": min_threshold,
                                "current_scale": quy_mo,
                                "current_scale_label": SCALE_LABELS.get(quy_mo, quy_mo),
                                "query_suggestions": alt_topics,
                            })
                            return
                
                # Mọi thứ an toàn, tiếp tục tạo Vector DB
                if not passages:
                    logger.warning(f"Job {ma_cv}: EKRE found no documents. Triggering fallback search...")
                
                # Tạo Vector DB từ EKRE Passages
                passages_db = tao_vector_db(passages, api_key=CauHinh.OPENAI_API_KEY, check_cancel=check_cancel)
                global_map = {p['id']: p for p in passages_db}
                
                # 🛠️ KHỞI TẠO PipelineContext: Trung tâm điều phối Metadata
                ctx = PipelineContext(
                    ma_cv=ma_cv,
                    tieu_de=tieu_de,
                    quy_mo=quy_mo,
                    api_keys_list=CauHinh.GEMINI_API_KEYS,
                    passages_db=passages_db,
                    global_map=global_map,
                    terms=du_lieu.get("terms", []),
                    passages=passages,
                    candidates=candidates,
                    openai_semaphore=OPENAI_SEMAPHORE,
                    safety_class=safety_class,
                    ngon_ngu=ngon_ngu,
                    custom_section_words=custom_section_words,
                    custom_sections_map=custom_sections_map
                )
                prefix = ctx.get_logger_prefix()
                
                # Bước 2: Dàn ý (Outline) & Thuật ngữ
                CONG_VIEC[ctx.ma_cv].update({"tien_do": 30, "buoc": "Đang phân tích tri thức & trích xuất thuật ngữ..."})
                
                try:
                    from dich_vu.openai_da_buoc import (
                        xay_dung_metadata_toan_dien, 
                        trich_xuat_thuat_ngu,
                        nhom_thuat_ngu_va_tao_dan_y, classify_domain,
                        tao_dan_y_tu_passages,
                        xac_dinh_ngan_sach_thuat_ngu,
                        get_structure_config,
                        InsufficientDataError
                    )
                    
                    # V23.5.4: Logging and Budgeting fix (Synchronized with Scales)
                    num_articles = len(ctx.candidates) if ctx.candidates else 1
                    so_chuong_yeu_cau = int(so_chuong_custom) if so_chuong_custom else 0
                    
                    # 1. Budgeting (Xác định quy mô và định mức thuật ngữ)
                    budget = xac_dinh_ngan_sach_thuat_ngu(num_articles, so_chuong_yeu_cau, quy_mo=ctx.quy_mo)
                    
                    target_ch_log = so_chuong_yeu_cau if so_chuong_yeu_cau > 0 else f"{get_structure_config(ctx.quy_mo)['ch'][0]}-{get_structure_config(ctx.quy_mo)['ch'][1]}"
                    logger.info(f"[STRUCTURE] Scale: {ctx.quy_mo} | TermBudget: {budget['core_count']} | ChapterGoal: {target_ch_log}")
                    
                    ghi_nhat_ky(f"Ngân sách thuật ngữ: {budget['core_count']} core. Mục tiêu: ~{target_ch_log} chương.")

                    # 2. Metadata Builder
                    metadata_list = xay_dung_metadata_toan_dien(ctx.passages)
                    
                    # 2. Term Extraction (Adaptive V23.3) - 🧵 Buffered by Semaphore
                    ghi_nhat_ky(f"Bắt đầu trích xuất thuật ngữ khoa học (Target: {budget['core_count'] + budget['support_count']}).")
                    step_start = time.time()
                    terms_data = trich_xuat_thuat_ngu(
                        ctx.passages, api_key=CauHinh.OPENAI_API_KEY, 
                        target_core=budget["core_count"], 
                        target_support=budget["support_count"],
                        semaphore=ctx.openai_semaphore,
                        check_cancel=check_cancel
                    )
                    logger.info(f"{prefix}: Term Extraction completed in {time.time()-step_start:.2f}s")
                    ctx.terms = terms_data.get("core_terms", []) + terms_data.get("supporting_terms", [])
                    CONG_VIEC[ma_cv]["terms_detail"] = terms_data
                    ghi_nhat_ky(f"Đã trích xuất {len(ctx.terms)} thuật ngữ khoa học.")
                    CONG_VIEC[ctx.ma_cv]["terms_count"] = len(ctx.terms)
                    
                    # 4. Validation Layer (Term Intensity Check)
                    expected_min = budget["core_count"] * 0.6
                    if len(ctx.terms) < 15:
                        logger.warning(f"{prefix}: Critically low core terms ({len(ctx.terms)} < 15). Triggering fallback.")
                        raise ValueError("Insufficient core terms extracted.")
                    elif len(ctx.terms) < expected_min:
                        logger.warning(f"{prefix}: Low core terms ({len(ctx.terms)} < {expected_min}). Letting Architect attempt outline anyway.")

                    # 4.5. Adaptive Routing (Complexity Scoring V36.0)
                    complexity_score = 0
                    disc_xray = CONG_VIEC[ma_cv].get("discovery_xray", {})
                    if disc_xray.get("complexity") == "high": complexity_score += 4
                    elif disc_xray.get("complexity") == "medium": complexity_score += 2
                    
                    if len(ctx.terms) > 80: complexity_score += 3
                    elif len(ctx.terms) > 50: complexity_score += 1
                    
                    if any(d.get("lang") == "en" for d in CONG_VIEC[ma_cv].get("top_30_links", [])):
                        complexity_score += 3
                        
                    is_hard_case = complexity_score >= 5
                    logger.info(f"[ROUTER] Topic Complexity Score: {complexity_score} | Routing to: {'Gemini 3.1' if is_hard_case else 'GPT-4o-mini'}")

                    # 4.6. V38: Domain Classification (Adaptive Framework Selection)
                    from dich_vu.openai_da_buoc import classify_domain
                    domain_info = classify_domain(tieu_de, CauHinh.OPENAI_API_KEY)
                    ghi_nhat_ky(f"📚 Lĩnh vực: {domain_info.get('label', 'N/A')} (confidence={domain_info.get('confidence', 0):.0%})")
                    CONG_VIEC[ma_cv]["domain_info"] = domain_info
                    CONG_VIEC[ma_cv]["domain_label"] = domain_info.get("label", "")

                    # 5. Advanced Outline Generation (Phases & Uniqueness) - 🧵 Buffered by Semaphore
                    step_start = time.time()
                    outline_data = nhom_thuat_ngu_va_tao_dan_y(
                        terms_data, 
                        api_key=CauHinh.OPENAI_API_KEY, 
                        chu_de=ctx.tieu_de,
                        so_chuong=so_chuong_yeu_cau,
                        quy_mo=ctx.quy_mo,
                        semaphore=ctx.openai_semaphore,
                        ngon_ngu=ngon_ngu,
                        so_chuong_custom=so_chuong_custom,
                        danh_sach_chuong=danh_sach_chuong,
                        passages=ctx.passages,
                        is_hard_case=is_hard_case,
                        domain_info=domain_info,
                        custom_section_words=custom_section_words,
                        custom_sections_map=custom_sections_map
                    )
                    actual_chapters = len(outline_data.get("outline", []))
                    logger.info(f"{prefix}: Advanced Outline completed in {time.time()-step_start:.2f}s. Chapters: {actual_chapters}")
                    
                    # 💎 Structural Audit (V24.6)
                    target_ch_audit = so_chuong_yeu_cau if so_chuong_yeu_cau > 0 else (get_structure_config(ctx.quy_mo).get("ch", (4, 8))[0] + get_structure_config(ctx.quy_mo).get("ch", (4, 8))[1]) // 2
                    if actual_chapters < target_ch_audit:
                        logger.warning(f"[STRUCTURE] Job {ma_cv}: Structural Violation (Got {actual_chapters}, Expected {target_ch_audit}).")
                    
                    if not outline_data or not outline_data.get("outline"):
                        raise ValueError("Advanced Outline failed or empty.")
                    
                    # --- PER-CHAPTER OUTLINE GUARD (V23.5.1 Synchronized) ---
                    # V40: Skip thin chapter guard when user provides custom titles + custom word count
                    is_focused_custom = bool((danh_sach_chuong and len(danh_sach_chuong) > 0 and custom_section_words) or has_custom_subsections)
                    raw_outline = outline_data.get("outline", [])
                    struct_cfg = get_structure_config(ctx.quy_mo)
                    min_required = struct_cfg["sec"][0] # Lấy số mục tối thiểu từ cấu hình chung
                    
                    thin_chapters = [c.get("title") for c in raw_outline if len(c.get("sections", [])) < min_required]
                    if thin_chapters and not is_focused_custom:
                        logger.warning(f"{prefix}: Found {len(thin_chapters)} thin chapters. Retrying with STRICT_FRAGMENTATION...")
                        retry_outline = nhom_thuat_ngu_va_tao_dan_y(
                            terms_data, api_key=CauHinh.OPENAI_API_KEY, 
                            chu_de=f"{ctx.tieu_de} (STRICT_FRAGMENTATION: Each chapter MUST have {min_required} to {struct_cfg['sec'][1]} sections)",
                            so_chuong=so_chuong_yeu_cau,
                            quy_mo=ctx.quy_mo,
                            semaphore=ctx.openai_semaphore,
                            ngon_ngu=ngon_ngu,
                            so_chuong_custom=so_chuong_custom,
                            danh_sach_chuong=danh_sach_chuong,
                            passages=ctx.passages,
                            is_hard_case=is_hard_case,
                            domain_info=domain_info,
                            custom_section_words=custom_section_words
                        )
                        if retry_outline and retry_outline.get("outline"):
                            outline_data = retry_outline
                            actual_ch_retry = len(outline_data.get("outline", []))
                            logger.info(f"{prefix}: Strict Fragmentation Outline completed in {time.time()-step_start:.2f}s. Chapters: {actual_ch_retry}")
                    elif thin_chapters and is_focused_custom:
                        logger.info(f"{prefix}: Skipping thin chapter guard for focused custom mode (custom titles + custom words). {len(thin_chapters)} chapters have fewer sections, which is expected.")

                except Exception as ex_prod:
                    logger.warning(f"{prefix}: Tier 1 Outline failed ({ex_prod}). Attempting Tier 2 (Smart Passages-to-Outline)...")
                    try:
                        # Tier 2: Smart Passages-to-Outline Fallback
                        step_start = time.time()
                        outline_data = tao_dan_y_tu_passages(
                            ctx.tieu_de, ctx.passages, 
                            api_key=CauHinh.OPENAI_API_KEY, 
                            quy_mo=ctx.quy_mo,
                            semaphore=ctx.openai_semaphore
                        )
                        logger.info(f"{prefix}: Tier 2 Smart Outline completed in {time.time()-step_start:.2f}s")
                    except Exception as ex_tier2:
                        logger.warning(f"{prefix}: Tier 2 Outline failed ({ex_tier2}). Falling back to Tier 3 (Legacy Rescue)...")
                        # Tier 3: Legacy Outline Rescue (Oldest Path)
                        top_pass = tim_kiem_vector(ctx.tieu_de, ctx.passages_db, api_key=CauHinh.OPENAI_API_KEY, top_k=60)
                        che_do = du_lieu.get("che_do", "auto")
                        outline_data = openai_tao_dan_y(
                            ctx.tieu_de, top_pass, api_key=CauHinh.OPENAI_API_KEY, 
                            quy_mo=ctx.quy_mo, che_do=che_do, so_chuong_max=8,
                            semaphore=ctx.openai_semaphore
                        )
                
                raw_outline = outline_data.get("outline", [])
                actual_ch = len(raw_outline)
                
                # Xác định số chương tối thiểu mong đợi (nếu custom thì theo custom, nếu auto thì theo quy mô)
                expected_ch_min = so_chuong_yeu_cau
                if expected_ch_min <= 0:
                    struct_cfg = get_structure_config(ctx.quy_mo)
                    expected_ch_min = struct_cfg.get("ch", (4, 8))[0]
                
                if actual_ch < expected_ch_min:
                    if so_chuong_yeu_cau > 0:
                        detail_msg = f"<b>{so_chuong_yeu_cau} chương</b> như yêu cầu."
                    else:
                        scale_name = SCALE_LABELS.get(ctx.quy_mo, ctx.quy_mo)
                        detail_msg = f"mức tối thiểu <b>{expected_ch_min} chương</b> của quy mô '{scale_name}'."
                        
                    error_msg_html = (
                        f"Chủ đề <b>'{ctx.tieu_de}'</b> khan hiếm dữ liệu trên Wikipedia.<br>"
                        f"Số chương tối đa có thể biên soạn dựa trên dữ liệu hiện có là <b>{actual_ch} chương</b> thay vì {detail_msg}<br><br>"
                        f"Bạn có đồng ý hạ số chương xuống <b>{actual_ch} chương</b> không?"
                    )
                    
                    CONG_VIEC[ma_cv].update({
                        "trang_thai": "cho_xac_nhan",
                        "loai_loi": "CUSTOM_CHAPTER_DOWNGRADE",
                        "loi": error_msg_html,
                        "suggested_chapters": actual_ch,
                        "requested_chapters": expected_ch_min,
                        "xac_nhan_cho_phep": None
                    })
                    
                    wait_seconds = 120
                    start_wait = time.time()
                    ghi_nhat_ky(f"Chờ người dùng xác nhận hạ số chương từ {expected_ch_min} xuống {actual_ch}...")
                    while CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is None:
                        time.sleep(1.0)
                        if CONG_VIEC[ma_cv].get("huy_bo"):
                            CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                            break
                        if time.time() - start_wait > wait_seconds:
                            CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                            break
                            
                    if CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is True:
                        ghi_nhat_ky(f"Người dùng đã đồng ý hạ số chương xuống {actual_ch}. Tiếp tục biên soạn.")
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "dang_chay",
                            "loi": None,
                            "xac_nhan_cho_phep": None
                        })
                    else:
                        ghi_nhat_ky("Tiến trình bị dừng do từ chối hạ số chương hoặc hết thời gian chờ.")
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "that_bai",
                            "loi": "Người dùng đã từ chối hạ số chương giáo trình. Tiến trình bị hủy."
                        })
                        return
                
                total_sections = sum([len(c.get("sections", [])) for c in raw_outline])
                CONG_VIEC[ma_cv]["tong_chuong"] = len(raw_outline)
                CONG_VIEC[ma_cv]["tong_muc"] = total_sections
                ghi_nhat_ky(f"Dàn ý hoàn tất: {len(raw_outline)} chương, {total_sections} mục con.")
                logger.info(f"{prefix}: Outline created with {len(raw_outline)} chapters and {total_sections} sections.")

                # --- DUYỆT DÀN Ý TÙY CHỌN (OPT-IN OUTLINE APPROVAL) ---
                if approve_outline:
                    outline_temp = []
                    sec_global_idx = 0
                    for ch_idx, ch in enumerate(raw_outline):
                        ch_data = {
                            "ch_idx": ch_idx,
                            "title": ch.get("title", ""),
                            "sections": []
                        }
                        for sec in ch.get("sections", []):
                            ch_data["sections"].append({
                                "idx": sec_global_idx,
                                "title": sec.get("title", "")
                            })
                            sec_global_idx += 1
                        outline_temp.append(ch_data)
                    
                    CONG_VIEC[ma_cv].update({
                        "trang_thai": "cho_xac_nhan",
                        "loai_loi": "APPROVE_OUTLINE",
                        "loi": "Chờ người dùng duyệt dàn ý chi tiết.",
                        "outline_temp": outline_temp,
                        "xac_nhan_cho_phep": None,
                        "da_chon_dan_y": None
                    })
                    
                    wait_seconds = 600
                    start_wait = time.time()
                    ghi_nhat_ky("Tạm dừng: Chờ người dùng duyệt chi tiết các mục con cấp 2...")
                    while CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is None:
                        time.sleep(1.0)
                        if CONG_VIEC[ma_cv].get("huy_bo"):
                            CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                            break
                        if time.time() - start_wait > wait_seconds:
                            CONG_VIEC[ma_cv]["xac_nhan_cho_phep"] = False
                            break
                            
                    if CONG_VIEC[ma_cv].get("xac_nhan_cho_phep") is True:
                        da_chon_dan_y = CONG_VIEC[ma_cv].get("da_chon_dan_y") or []
                        ghi_nhat_ky(f"Người dùng đã duyệt dàn ý. Số mục được tích chọn: {len(da_chon_dan_y)}.")
                        
                        new_raw_outline = []
                        filtered_sec_global_idx = 0
                        for ch in raw_outline:
                            new_ch_sections = []
                            for sec in ch.get("sections", []):
                                if filtered_sec_global_idx in da_chon_dan_y:
                                    new_ch_sections.append(sec)
                                filtered_sec_global_idx += 1
                            new_raw_outline.append({
                                "title": ch.get("title", ""),
                                "sections": new_ch_sections
                            })
                        
                        raw_outline = new_raw_outline
                        outline_data["outline"] = raw_outline
                        
                        total_sections = sum([len(c.get("sections", [])) for c in raw_outline])
                        CONG_VIEC[ma_cv]["tong_chuong"] = len(raw_outline)
                        CONG_VIEC[ma_cv]["tong_muc"] = total_sections
                        
                        ghi_nhat_ky(f"Dàn ý đã được chọn lọc lại: {len(raw_outline)} chương, {total_sections} mục con.")
                        logger.info(f"{prefix}: Filtered outline has {len(raw_outline)} chapters and {total_sections} sections.")
                        
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "dang_chay",
                            "loi": None,
                            "xac_nhan_cho_phep": None,
                            "da_chon_dan_y": None,
                            "outline_temp": None
                        })
                    else:
                        ghi_nhat_ky("Tiến trình bị hủy do từ chối duyệt dàn ý hoặc quá thời gian chờ.")
                        CONG_VIEC[ma_cv].update({
                            "trang_thai": "that_bai",
                            "loi": "Người dùng từ chối duyệt dàn ý hoặc quá thời gian chờ (10 phút). Tiến trình bị hủy."
                        })
                        return

                # --- 🚀 LEVEL 3 EXPANSION: Outline-Driven Iterative Retrieval (Gap-Filling V42) ---
                from dich_vu.gap_filler import identify_knowledge_gaps, fill_knowledge_gaps
                
                check_cancel()
                CONG_VIEC[ma_cv].update({"tien_do": 45, "buoc": "Đang kiểm toán lỗ hổng tri thức (Knowledge Gap-Filling)..."})
                ghi_nhat_ky("Bắt đầu quét dàn ý để phát hiện lỗ hổng tri thức so với cơ sở dữ liệu.")
                
                # Quét lỗ hổng dựa trên từng section thay vì chỉ chapter
                gaps = identify_knowledge_gaps(raw_outline, ctx.passages_db, CauHinh.OPENAI_API_KEY, ctx.tieu_de, custom_section_words=custom_section_words, check_cancel=check_cancel)
                # Xác định is_custom_flow sớm để Gap Filler biết dùng truy vấn trực tiếp
                _custom_map = getattr(ctx, 'custom_sections_map', None)
                is_custom_flow = False
                if _custom_map:
                    is_custom_flow = any(len(secs) > 0 for secs in _custom_map.values())
                elif danh_sach_chuong:
                    is_custom_flow = len(danh_sach_chuong) > 0

                if gaps:
                    ghi_nhat_ky(f"Phát hiện {len(gaps)} lỗ hổng tri thức. Đang kích hoạt tìm kiếm bù đắp...")
                    new_passages = fill_knowledge_gaps(gaps, ctx.api_keys_list, CauHinh.OPENAI_API_KEY, ctx.tieu_de, is_custom_flow=is_custom_flow, custom_section_words=custom_section_words, check_cancel=check_cancel)
                    
                    if new_passages:
                        ghi_nhat_ky(f"Đang lập chỉ mục (Vectorizing) {len(new_passages)} đoạn văn bổ sung...")
                        try:
                            next_id = len(ctx.passages_db) + 1
                            new_passages_db = tao_vector_db(new_passages, api_key=CauHinh.OPENAI_API_KEY, start_id=next_id, check_cancel=check_cancel)
                            
                            if new_passages_db:
                                # Update thread-safe
                                with PASSAGES_LOCK:
                                    current_db = list(ctx.passages_db)
                                    current_db.extend(new_passages_db)
                                    ctx.passages_db = current_db
                                    
                                    ctx.passages.extend(new_passages)
                                    for p in new_passages_db:
                                        ctx.global_map[str(p['id'])] = p
                                        
                                ghi_nhat_ky(f"Đã cập nhật Vector DB. Tổng cộng: {len(ctx.passages_db)} passages.")
                        except Exception as ex_embed:
                            logger.error(f"{prefix}: Gap Filler embedding failed: {ex_embed}")
                            ghi_nhat_ky("Lỗi khi lập chỉ mục dữ liệu bù đắp.")
                    else:
                        ghi_nhat_ky("Không tìm thấy dữ liệu bù đắp hữu ích do Timeout hoặc lỗi nguồn.")
                else:
                    ghi_nhat_ky("Cơ sở dữ liệu đã phủ đủ chi tiết, không có lỗ hổng.")

                # Bước 3: Biên soạn nội dung (V17.0+ Turbo Resilience)
                check_cancel()
                
                # --- ASYNCHRONOUS PRE-FETCHING (V35) ---
                ghi_nhat_ky("Đang phân tích tải Reranking (Asynchronous Pre-fetching)...")
                from concurrent.futures import ThreadPoolExecutor, as_completed
                from dich_vu.vector_search import tim_kiem_vector_with_llm_rerank
                
                all_sections = []
                for chap in raw_outline:
                    chap_t = chap.get("title", "")
                    for sec in chap.get("sections", []):
                        sec_t = sec.get("title", "") if isinstance(sec, dict) else sec
                        all_sections.append((chap_t, sec_t))
                
                ghi_nhat_ky(f"[Batch Rerank] Đang tải tài liệu song song cho {len(all_sections)} mục (Max Workers: 5)...")
                
                has_custom = False
                custom_map = getattr(ctx, 'custom_sections_map', None)
                if custom_map:
                    has_custom = any(len(secs) > 0 for secs in custom_map.values())
                candidate_count = 50 if has_custom else 30
                
                def _prefetch_worker(chap_title, sec_title):
                    try:
                        dynamic_top_k = {"can_ban": 7, "tieu_chuan": 12, "chuyen_sau": 18}.get(ctx.quy_mo, 12)
                        p = tim_kiem_vector_with_llm_rerank(
                            query=f"{chap_title} {sec_title}",
                            passages_db=ctx.passages_db,
                            api_key=CauHinh.OPENAI_API_KEY,
                            top_k=dynamic_top_k,
                            candidate_n=min(candidate_count, len(ctx.passages_db)),
                            chapter_title=chap_title,
                            section_title=sec_title
                        )
                        return (chap_title, sec_title, p)
                    except Exception as e:
                        logger.error(f"[Prefetch Error] {chap_title} - {sec_title}: {e}")
                        return (chap_title, sec_title, [])

                ctx.prefetched_passages = {}
                with ThreadPoolExecutor(max_workers=5) as pf_executor:
                    pf_futures = [pf_executor.submit(_prefetch_worker, ct, st) for ct, st in all_sections]
                    for f in as_completed(pf_futures):
                        ct, st, p = f.result()
                        ctx.prefetched_passages[(ct, st)] = p
                
                ghi_nhat_ky("Khâu lọc tài liệu hoàn tất. Bắt đầu đẩy vào AI Writer.")

                CONG_VIEC[ma_cv].update({"tien_do": 50, "buoc": "Đang biên soạn nội dung (Parallel Writing)..."})
                ghi_nhat_ky("Bắt đầu giai đoạn biên soạn nội dung song song (Multi-threaded Micro-Writer).")
                final_chapters = parallel_generate(ctx, raw_outline, outline_data)
                ghi_nhat_ky(f"Biên soạn xong {len(final_chapters)} chương.")

                # --- CHUNKED BATCH CRITIC (V28) ---
                check_cancel()
                CONG_VIEC[ma_cv].update({"tien_do": 60, "buoc": "Đang chạy Chunked Batch Critic..."})
                ghi_nhat_ky("Bắt đầu kiểm toán gộp (Chunked Batch Critic).")
                from dich_vu.audit_service import ScholarlyAuditEngine
                audit_engine = ScholarlyAuditEngine(openai_key=CauHinh.OPENAI_API_KEY, gemini_keys=ctx.api_keys_list)
                
                chunk_size = 3
                audited_chapters = []
                for i in range(0, len(final_chapters), chunk_size):
                    check_cancel()
                    chunk = final_chapters[i:i+chunk_size]
                    import logging
                    logging.getLogger(__name__).info(f"{prefix}: Auditing chunk {i//chunk_size + 1}...")
                    audited_chunk = audit_engine.batch_audit_chunks(chunk, tieu_de)
                    audited_chapters.extend(audited_chunk)
                    time.sleep(1)
                    
                final_chapters = audited_chapters
                ghi_nhat_ky("Kiểm toán gộp hoàn tất.")

                # --- V33: CHAPTER SUMMARIES & GLOSSARY ---
                check_cancel()
                CONG_VIEC[ma_cv].update({"tien_do": 70, "buoc": "Đang sinh tóm tắt chương & bảng thuật ngữ..."})
                from dich_vu.openai_da_buoc import sinh_tom_tat_chuong, sinh_bang_thuat_ngu
                
                # Sinh tóm tắt cho từng chương (song song)
                from concurrent.futures import ThreadPoolExecutor as _TPE
                def _gen_summary(chap):
                    if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
                        raise Exception("Tiến trình đã bị người dùng hủy.")
                    sections_text = "\n".join(sec.get("content", "") for sec in chap.get("sections", []))
                    return sinh_tom_tat_chuong(tieu_de, chap.get("title", ""), sections_text, CauHinh.OPENAI_API_KEY, OPENAI_SEMAPHORE, ngon_ngu=ctx.ngon_ngu)
                
                with _TPE(max_workers=4) as _ex:
                    summaries = list(_ex.map(_gen_summary, final_chapters))
                for i, chap in enumerate(final_chapters):
                    chap["summary"] = summaries[i] if i < len(summaries) else ""
                ghi_nhat_ky(f"Đã sinh tóm tắt cho {sum(1 for s in summaries if s)} chương.")
                
                # V33.2: Sinh bài tập & câu hỏi ôn tập chương cho luồng tự custom
                is_custom_flow = False
                custom_map = getattr(ctx, 'custom_sections_map', None)
                if custom_map:
                    is_custom_flow = any(len(secs) > 0 for secs in custom_map.values())
                elif danh_sach_chuong:
                    is_custom_flow = len(danh_sach_chuong) > 0

                if is_custom_flow:
                    ghi_nhat_ky("Đang sinh bài tập & câu hỏi ôn tập cho từng chương...")
                    from dich_vu.openai_da_buoc import sinh_bai_tap_on_tap_chuong
                    
                    def _gen_exercises(chap):
                        if CONG_VIEC.get(ctx.ma_cv, {}).get("huy_bo"):
                            raise Exception("Tiến trình đã bị người dùng hủy.")
                        sections_text = "\n".join(sec.get("content", "") for sec in chap.get("sections", []))
                        return sinh_bai_tap_on_tap_chuong(tieu_de, chap.get("title", ""), sections_text, CauHinh.OPENAI_API_KEY, OPENAI_SEMAPHORE, ngon_ngu=ctx.ngon_ngu)
                    
                    with _TPE(max_workers=4) as _ex:
                        exercises_list = list(_ex.map(_gen_exercises, final_chapters))
                    
                    for i, chap in enumerate(final_chapters):
                        ex_content = exercises_list[i] if i < len(exercises_list) else ""
                        if ex_content:
                            sec_title = "Bài tập & Câu hỏi Ôn tập Chương" if ctx.ngon_ngu != "en" else "Review Questions & Practical Exercises"
                            chap["sections"].append({
                                "title": sec_title,
                                "content": ex_content,
                                "citations": []
                            })
                    ghi_nhat_ky(f"Đã sinh bài tập & ôn tập cho {sum(1 for e in exercises_list if e)} chương.")
                
                # Lấy toàn bộ tiêu đề Chương và Mục để làm Thuật ngữ
                title_terms = []
                for chap in final_chapters:
                    title_terms.append(chap.get("title", ""))
                    for sec in chap.get("sections", []):
                        title_terms.append(sec.get("title", ""))
                        
                # Kết hợp: Tiêu đề dàn ý + Các thuật ngữ khoa học cốt lõi ban đầu
                combined_terms = title_terms + [t.get("term", "") for t in ctx.terms]

                # Sinh bảng thuật ngữ (glossary)
                glossary = sinh_bang_thuat_ngu(combined_terms, tieu_de, CauHinh.OPENAI_API_KEY, OPENAI_SEMAPHORE, ngon_ngu=ctx.ngon_ngu)
                ghi_nhat_ky(f"Bảng thuật ngữ: {len(glossary)} định nghĩa.")
                CONG_VIEC[ma_cv]["glossary"] = glossary

                ghi_nhat_ky("Bắt đầu hậu xử lý trích dẫn.")

                # Bước 4: Hậu xử lý trích dẫn & Tham khảo (V37 APA 7th Edition)
                all_original_passages = {str(p['id']): p for p in ctx.passages_db}
                url_to_new_id = {}
                ordered_refs = []
                
                # V37: Ngày truy cập động cho APA
                from datetime import datetime as _dt_apa
                _apa_access_date = _dt_apa.now().strftime("%B %d, %Y")  # "April 28, 2026"
                _apa_year = _dt_apa.now().year
                
                def _shorten_apa_title(title, max_chars=40):
                    """Rút gọn tiêu đề bài viết cho inline citation APA."""
                    if not title: return "Wikipedia"
                    if len(title) <= max_chars: return title
                    truncated = title[:max_chars - 3].rsplit(' ', 1)[0]
                    return truncated + '…'
                
                # Scan 1: Xây dựng bản đồ trích dẫn duy nhất theo thứ tự xuất hiện (APA)
                for chap in final_chapters:
                    for sec in chap.get("sections", []):
                        found_groups = re.findall(r'\[(\d+(?:\s*,\s*\d+)*)\]', sec.get("content", ""))
                        for group in found_groups:
                            for oid in group.split(","):
                                oid = oid.strip()
                                if oid in all_original_passages:
                                    p = all_original_passages[oid]
                                    url = p.get('url', '')
                                    if url and url not in url_to_new_id:
                                        new_id = len(url_to_new_id) + 1
                                        title_full = p.get('title', 'Nguồn')
                                        short_title = _shorten_apa_title(title_full)
                                        url_to_new_id[url] = new_id
                                        ordered_refs.append({
                                            "id": new_id, "url": url,
                                            "title": title_full,
                                            "short_title": short_title,
                                            "access_date": _apa_access_date,
                                            "year": _apa_year
                                        })

                # Scan 2: Cập nhật mã trích dẫn mới vào nội dung và metadata
                for chap in final_chapters:
                    for sec in chap.get("sections", []):
                        content = sec.get("content", "")
                        sec_citations = []
                        added_urls = set()
                        
                        def _citation_replacer(match):
                            raw_ids = match.group(1).split(",")
                            tags = []
                            for oid in raw_ids:
                                oid = oid.strip()
                                if oid in all_original_passages:
                                    url = all_original_passages[oid].get('url', '')
                                    if url in url_to_new_id:
                                        ref_entry = next(r for r in ordered_refs if r["url"] == url)
                                        short_title = ref_entry.get("short_title", ref_entry["title"])
                                        title_escaped = ref_entry["title"].replace('"', '&quot;')
                                        year = ref_entry.get("year", _apa_year)
                                        tags.append(f'<span class="citation-apa">(<a href="{url}" title="{title_escaped}" target="_blank" rel="noopener noreferrer">{short_title}, {year}</a>)</span>')
                                        if url not in added_urls:
                                            sec_citations.append(ref_entry)
                                            added_urls.add(url)
                            return " ".join(tags) if tags else "" # Gỡ bỏ token [ID] lỗi không có trong hệ thống

                        # V37 APA: Cập nhật mã trích dẫn trong văn bản (hỗ trợ cả dạng gom cụm [1, 2])
                        content = re.sub(r'\[(\d+(?:\s*,\s*\d+)*)\]', _citation_replacer, content)
                        
                        sec["content"] = content
                        sec["citations"] = sec_citations

                # --- V33: GROUNDING SCORE CALCULATION ---
                grounding_stats = {"chapters": [], "overall": {}}
                total_paras_all = 0
                grounded_paras_all = 0

                for chap in final_chapters:
                    chap_title_gs = chap.get("title", "")
                    chap_total = 0
                    chap_grounded = 0
                    for sec in chap.get("sections", []):
                        # Skip chapter-level exercises section entirely
                        sec_title = sec.get("title", "")
                        if sec_title in ["Bài tập & Câu hỏi Ôn tập Chương", "Review Questions & Practical Exercises"]:
                            continue
                        content_gs = sec.get("content", "")
                        # Loại bỏ phần Bài tập/Ôn tập khỏi Grounding (không yêu cầu trích dẫn)
                        # Dùng regex bắt mọi biến thể: ### Câu hỏi..., **Câu hỏi...**, dòng thuần
                        review_pattern = re.compile(
                            r'(###\s*)?\*{0,2}(Câu hỏi Ôn tập|Bài tập\s*[&＆]\s*Ôn tập|Review Questions)\*{0,2}',
                            re.IGNORECASE
                        )
                        for line in content_gs.split("\n"):
                            if review_pattern.search(line.strip()):
                                content_gs = content_gs.split(line)[0]
                                break

                        paragraphs = [
                            p.strip() for p in content_gs.split("\n")
                            if p.strip()
                            and not p.strip().startswith("### ")   # Bỏ heading tiểu mục
                            and not re.match(r'^\d+\.\s', p.strip())  # Bỏ danh sách câu hỏi số
                            and len(p.strip()) >= 40               # Chỉ đoạn đủ dài
                        ]
                        for para in paragraphs:
                            chap_total += 1
                            # V37: Nhận diện cả APA format và IEEE legacy
                            if re.search(r'<span class="citation-apa">', para) or re.search(r'<sup class="citation">', para) or re.search(r'\[\d+\]', para):
                                chap_grounded += 1
                    ratio_gs = (chap_grounded / chap_total * 100) if chap_total > 0 else 0
                    grounding_stats["chapters"].append({
                        "title": chap_title_gs, "total": chap_total,
                        "grounded": chap_grounded, "ratio": round(ratio_gs, 1)
                    })
                    total_paras_all += chap_total
                    grounded_paras_all += chap_grounded

                overall_ratio = (grounded_paras_all / total_paras_all * 100) if total_paras_all > 0 else 0
                grounding_stats["overall"] = {
                    "total": total_paras_all, "grounded": grounded_paras_all,
                    "ratio": round(overall_ratio, 1)
                }
                CONG_VIEC[ma_cv]["grounding"] = grounding_stats
                ghi_nhat_ky(f"Grounding Score: {overall_ratio:.1f}% ({grounded_paras_all}/{total_paras_all} đoạn có nguồn)")


                # 3. Đóng gói Giáo trình Final
                from dich_vu.kiem_tra_cau_truc_json import clean_title_numbering
                final_book = {"title": tieu_de, "sections": [], "references": []}
                book_export = {"title": tieu_de, "chapters": [], "is_custom_flow": is_custom_flow}

                for chap in final_chapters:
                    c_title = clean_title_numbering(chap.get("title", "Không tên"))
                    new_chap = {"title": c_title, "summary": chap.get("summary", ""), "sections": []}
                    final_book["sections"].append({"title": c_title, "is_chapter": True, "content": ""})
                    
                    for sec in chap.get("sections", []):
                        s_title = clean_title_numbering(sec.get("title", "Mục"))
                        s_content = sec.get("content", "")
                        new_chap["sections"].append({"title": s_title, "content": s_content, "citations": sec.get("citations", [])})
                        final_book["sections"].append({"title": s_title, "is_chapter": False, "content": s_content})
                    book_export["chapters"].append(new_chap)

                # Bước 5: Xuất bản PDF/DOCX
                check_cancel()
                CONG_VIEC[ma_cv].update({"tien_do": 85, "buoc": "Đang đóng gói tài liệu..."})
                # Lấy danh sách Heading từ Wikipedia để làm nguồn minh chứng
                kb_headings = []
                for p in ctx.passages:
                    text = p.get("text", "")
                    for line in text.split("\n"):
                        stripped = line.strip()
                        if stripped.startswith("==") and stripped.endswith("=="):
                            heading = stripped.strip("= ").strip()
                            if heading and len(heading) > 2 and heading not in kb_headings:
                                kb_headings.append(heading)

                all_refs = ordered_refs

                # Bước 6: Kiểm định cuối cùng (Final Audit for Scale Compliance)
                from dich_vu.gemini_giam_sat import giam_sat_quy_mo
                audit_quy_mo = giam_sat_quy_mo(
                    tieu_de, 
                    final_chapters, 
                    quy_mo, 
                    api_keys=CauHinh.GEMINI_API_KEYS,
                    so_chuong_yeu_cau=(len(final_chapters) if (so_chuong_custom or danh_sach_chuong) else 0),
                    custom_section_words=getattr(ctx, 'custom_section_words', 0)
                )
                if audit_quy_mo.get("status") == "fail":
                    logger.warning(f"Quy mô chưa đạt kỳ vọng: {audit_quy_mo.get('issues')}")
                    ghi_nhat_ky(f"Cảnh báo quy mô: {audit_quy_mo.get('status')}. Vẫn tiếp tục đóng gói bản tốt nhất.")

                ket_qua = {"topic": tieu_de, "book_vi": book_export, "references": all_refs, "ui_book": final_book, "glossary": glossary, "grounding": grounding_stats, "extracted_terms": [{"term": t.get("term", t) if isinstance(t, dict) else str(t), "meaning": t.get("meaning", "") if isinstance(t, dict) else ""} for t in ctx.terms], "kb_headings": kb_headings, "audit_quy_mo": audit_quy_mo, "is_custom_flow": is_custom_flow}
                
                # --- PHIÊN BẢN SẠCH (KHÔNG TRÍCH DẪN) - V23.5 ---
                import copy
                def strip_citations(text):
                    if not text: return ""
                    # V37: Loại bỏ cả APA <span> và legacy <sup> trích dẫn
                    clean = re.sub(r'<span class="citation-apa">.*?</span>', '', text)
                    clean = re.sub(r'<sup class="citation">.*?</sup>', '', clean)
                    # Loại bỏ các tag [ID] thô nếu còn sót
                    clean = re.sub(r'\[\w+\]', '', clean)
                    return clean

                book_plain = copy.deepcopy(book_export)
                for chap in book_plain.get("chapters", []):
                    for sec in chap.get("sections", []):
                        sec["content"] = strip_citations(sec.get("content", ""))
                        sec["citations"] = [] # Xóa metadata trích dẫn

                ket_qua_plain = {
                    "topic": tieu_de, 
                    "book_vi": book_plain, 
                    "references": [], # Xóa danh mục tham khảo
                    "terms": ket_qua.get("terms", []), # Giữ lại thuật ngữ
                    "is_custom_flow": is_custom_flow
                }

                p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
                p_docx = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}.docx")
                p_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}.pdf")
                
                p_docx_plain = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_cv}_plain.docx")
                p_pdf_plain = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_cv}_plain.pdf")

                # Xuất các phiên bản
                _luu_json(ket_qua, p_json)
                xuat_docx(ket_qua, p_docx); xuat_pdf(ket_qua, p_pdf)
                xuat_docx(ket_qua_plain, p_docx_plain); xuat_pdf(ket_qua_plain, p_pdf_plain)

                # Upload to Azure Blob Storage (V34+)
                try:
                    from dich_vu.azure_blob import upload_to_blob
                    upload_to_blob(p_json, f"json/{ma_cv}.json")
                    upload_to_blob(p_docx, f"docx/{ma_cv}.docx")
                    upload_to_blob(p_pdf, f"pdf/{ma_cv}.pdf")
                    upload_to_blob(p_docx_plain, f"docx/{ma_cv}_plain.docx")
                    upload_to_blob(p_pdf_plain, f"pdf/{ma_cv}_plain.pdf")
                except Exception as blob_err:
                    logger.error(f"Failed to upload generated files to Azure Blob Storage: {blob_err}")

                CONG_VIEC[ma_cv].update({
                    "trang_thai": "hoan_thanh", "tien_do": 100, "nguon": all_refs,
                    "tai_docx": f"/tai/docx/{ma_cv}", 
                    "tai_pdf": f"/tai/pdf/{ma_cv}",
                    "tai_docx_plain": f"/tai/docx/{ma_cv}_plain",
                    "tai_pdf_plain": f"/tai/pdf/{ma_cv}_plain",
                    "giam_sat": {"chapters": len(final_chapters), "circuit_breaker": getattr(ctx, 'use_gemini_only', False)},
                    "audit_quy_mo": audit_quy_mo
                })
                
                ghi_nhat_ky(f"Quy trình hoàn tất thành công trong {time.time() - start_time:.1f} giây.")
                logger.info(f"Job {ma_cv}: COMPLETED in {time.time() - start_time:.2f} seconds.")

                # Lưu lịch sử database (Luôn lưu cho cả khách lẻ và thành viên bằng HTML tối giản)
                try:
                    from mo_hinh import LichSuGiaoTrinh
                    
                    noi_dung_html = tao_html_hoc_thuat(tieu_de, book_export, all_refs)
                    
                    # Tính tổng số ký tự nội dung giáo trình (V32 Fix)
                    tong_ky_tu = 0
                    for chap in book_export.get("chapters", []):
                        for sec in chap.get("sections", []):
                            tong_ky_tu += len(sec.get("content", ""))
                    
                    history_entry = LichSuGiaoTrinh(
                        nguoi_dung_id=user_id, # None nếu là khách vãng lai
                        chu_de=tieu_de, 
                        noi_dung_html=noi_dung_html, 
                        duong_dan_file=p_pdf, 
                        da_xuat_file=True,
                        do_dai_ky_tu=tong_ky_tu
                    )
                    db.session.add(history_entry)
                    db.session.commit()
                    logger.info(f"Successfully saved curriculum history for user_id={user_id}, topic='{tieu_de}'")
                except Exception as db_err:
                    db.session.rollback()
                    logger.error(f"DB History Error: {db_err}")

            except Exception as e:
                logger.error(f"Job {ma_cv} failed: {e}\n{traceback.format_exc()}")
                CONG_VIEC[ma_cv].update({"trang_thai": "that_bai", "loi": str(e)})
            finally:
                # V21.6 Hardening: Close DB session and ensure status isn't stuck
                db.session.remove()
                if ma_cv in CONG_VIEC and CONG_VIEC[ma_cv]["trang_thai"] == "dang_chay":
                    CONG_VIEC[ma_cv]["trang_thai"] = "that_bai"
                    CONG_VIEC[ma_cv]["loi"] = "Pipeline terminated unexpectedly."

    # Chạy Background
    t = threading.Thread(target=run_pipeline, name="pipeline_thread", args=(u_id, so_chuong_custom, danh_sach_chuong, custom_section_words, approve_outline))
    t.start()
    return jsonify({"ma_cv": ma_cv, "trang_thai": "dang_chay"})

@app.get("/trang_thai/<ma_cv>")
def trang_thai(ma_cv):
    data = CONG_VIEC.get(ma_cv, {"trang_thai": "khong_tim_thay"})
    # V30 Fix: Convert numpy types to native Python types for JSON serialization
    return jsonify(json.loads(json.dumps(data, default=_json_safe_default)))

def sanitize_filename(filename):
    """Làm sạch tên file, giữ lại tiếng Việt có dấu (modern browsers support it)."""
    return re.sub(r'[\\/*?:"<>|]', '', filename).strip()

def kiem_tra_quyen_tai(ma_goc):
    from mo_hinh import LichSuGiaoTrinh
    from flask_login import current_user
    
    # 1. Nếu là Admin, luôn được phép tải
    if current_user.is_authenticated and current_user.la_admin:
        return True
        
    # 2. Tìm giáo trình trong lịch sử
    ls = LichSuGiaoTrinh.query.filter(LichSuGiaoTrinh.duong_dan_file.contains(ma_goc)).first()
    if not ls:
        # Nếu không có trong lịch sử (chưa lưu, khách lẻ hoặc chạy local không qua DB), cho phép tải tự do để tránh lỗi
        return True
        
    # 3. Nếu không có người sở hữu (khách lẻ), cho phép tải tự do
    if not ls.nguoi_dung_id:
        return True
        
    # 4. Nếu là giáo trình do chính họ tạo ra, cho phép tải
    if current_user.is_authenticated and ls.nguoi_dung_id == current_user.id:
        return True
        
    # 5. Nếu giáo trình này được bật trưng bày (noi_bat = True), cho phép BẤT KỲ AI tải xuống!
    if ls.noi_bat:
        return True
        
    return False

def khoi_phuc_ket_qua_tu_html(html_content, tieu_de_fallback):
    from bs4 import BeautifulSoup
    import re
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 1. Trích xuất tiêu đề
    title = tieu_de_fallback
    h1 = soup.find('h1')
    if h1:
        text = h1.get_text()
        if ":" in text:
            title = text.split(":", 1)[1].strip()
        else:
            title = text.replace("GIÁO TRÌNH", "").replace("Giáo trình", "").strip()
            
    chapters = []
    
    # 2. Lấy các chương
    chuong_mucs = soup.find_all(class_='chuong-muc')
    if not chuong_mucs:
        current_chap = None
        current_sec = None
        
        container = soup.find(class_='giao-trinh-container') or soup
        for elem in container.find_all(['h2', 'h3', 'div', 'p', 'ul', 'ol', 'hr']):
            if elem.name == 'h2':
                h2_text = elem.get_text().strip()
                ch_title = h2_text
                if "." in h2_text:
                    parts = h2_text.split(".", 1)
                    ch_title = parts[1].strip()
                elif "Chương" in h2_text or "chuong" in h2_text.lower():
                    ch_title = re.sub(r'^(?i)chương\s*\d+\s*:?\s*', '', h2_text).strip()
                    
                current_chap = {
                    "title": ch_title,
                    "summary": "",
                    "sections": []
                }
                chapters.append(current_chap)
                current_sec = None
            elif elem.name == 'h3' and current_chap is not None:
                h3_text = elem.get_text().strip()
                sec_title = re.sub(r'^\d+(\.\d+)+\.?\s*', '', h3_text).strip()
                current_sec = {
                    "title": sec_title,
                    "content": ""
                }
                current_chap["sections"].append(current_sec)
            elif elem.name in ['div', 'p', 'ul', 'ol', 'hr'] and current_sec is not None:
                elem_text = elem.get_text().strip()
                if "Tóm tắt chương:" in elem_text and current_chap is not None:
                    current_chap["summary"] = elem_text.split("Tóm tắt chương:", 1)[1].strip()
                    continue
                    
                lines = []
                if elem.name in ['ul', 'ol']:
                    for li in elem.find_all('li'):
                        li_text = li.get_text().strip()
                        if li_text:
                            lines.append(f"- {li_text}")
                elif elem.name == 'hr':
                    lines.append("---")
                else:
                    if elem_text:
                        lines.append(elem_text)
                        
                if lines:
                    content_str = "\n".join(lines)
                    if current_sec["content"]:
                        current_sec["content"] += "\n" + content_str
                    else:
                        current_sec["content"] = content_str
    else:
        for ch_div in chuong_mucs:
            h2 = ch_div.find('h2')
            if not h2:
                continue
            h2_text = h2.get_text().strip()
            ch_title = h2_text
            if "." in h2_text:
                parts = h2_text.split(".", 1)
                ch_title = parts[1].strip()
            elif "Chương" in h2_text or "chuong" in h2_text.lower():
                ch_title = re.sub(r'^(?i)chương\s*\d+\s*:?\s*', '', h2_text).strip()
                
            summary = ""
            summary_div = ch_div.find('div', style=lambda s: s and 'background' in s and 'f8f9fa' in s)
            if not summary_div:
                for d in ch_div.find_all('div'):
                    if "Tóm tắt chương" in d.get_text():
                        summary_div = d
                        break
            if summary_div:
                summary_text = summary_div.get_text().strip()
                if "Tóm tắt chương:" in summary_text:
                    summary = summary_text.split("Tóm tắt chương:", 1)[1].strip()
                else:
                    summary = summary_text
                    
            sections = []
            muc_cons = ch_div.find_all(class_='muc-con')
            for sec_div in muc_cons:
                h3 = sec_div.find('h3')
                if not h3:
                    continue
                h3_text = h3.get_text().strip()
                sec_title = re.sub(r'^\d+(\.\d+)+\.?\s*', '', h3_text).strip()
                
                content_div = sec_div.find(class_='content')
                if content_div:
                    content_lines = []
                    for child in content_div.children:
                        if child.name in ['p', 'div']:
                            p_text = child.get_text().strip()
                            if p_text:
                                content_lines.append(p_text)
                        elif child.name in ['ul', 'ol']:
                            for li in child.find_all('li'):
                                li_text = li.get_text().strip()
                                if li_text:
                                    content_lines.append(f"- {li_text}")
                        elif child.name in ['h4', 'h5', 'h6']:
                            h_text = child.get_text().strip()
                            if h_text:
                                content_lines.append(f"### {h_text}")
                        elif child.name == 'hr':
                            content_lines.append("---")
                    sec_content = "\n".join(content_lines)
                else:
                    sec_content = sec_div.get_text().strip()
                    
                sections.append({
                    "title": sec_title,
                    "content": sec_content
                })
                
            chapters.append({
                "title": ch_title,
                "summary": summary,
                "sections": sections
            })
            
    # 3. Tài liệu tham khảo
    references = []
    ref_container = soup.find(class_='tai-lieu-tham-khao') or soup.find(id='tai-lieu-tham-khao')
    if not ref_container:
        for h in soup.find_all(['h2', 'h3']):
            if "tham khảo" in h.get_text().lower():
                parent = h.parent
                ref_container = parent
                break
                
    if ref_container:
        p_tags = ref_container.find_all('p')
        for p in p_tags:
            ref_text = p.get_text().strip()
            if "tài liệu tham khảo" in ref_text.lower():
                continue
            match = re.match(r'^\[(\d+)\]\s*(.*)', ref_text)
            if match:
                ref_id = int(match.group(1))
                rest = match.group(2).strip()
                year_match = re.search(r'\.\s*\(([^)]+)\)', rest)
                if year_match:
                    ref_title = rest[:year_match.start()].strip()
                    ref_year_str = year_match.group(1)
                    
                    url_a = p.find('a')
                    ref_url = url_a['href'] if url_a and url_a.has_attr('href') else ""
                    if not ref_url:
                        url_search = re.search(r'https?://[^\s]+', rest)
                        if url_search:
                            ref_url = url_search.group(0)
                            
                    references.append({
                        "id": ref_id,
                        "title": ref_title,
                        "year": ref_year_str,
                        "url": ref_url
                    })
                else:
                    references.append({
                        "id": ref_id,
                        "title": rest,
                        "url": ""
                    })
            else:
                if ref_text:
                    references.append({
                        "title": ref_text,
                        "url": ""
                    })
                    
    return {
        "topic": title,
        "book_vi": {
            "title": title,
            "chapters": chapters
        },
        "references": references,
        "terms": [],
        "glossary": []
    }

def dam_bao_file_ton_tai(ma_goc, ext, is_plain, tieu_de):
    suffix = "_plain" if is_plain else ""
    folder = CauHinh.THU_MUC_PDF if ext == "pdf" else CauHinh.THU_MUC_DOCX
    path = os.path.join(folder, f"{ma_goc}{suffix}.{ext}")
    
    if os.path.exists(path):
        return True
        
    try:
        from dich_vu.azure_blob import download_from_blob
        blob_name = f"{ext}/{ma_goc}{suffix}.{ext}"
        download_from_blob(blob_name, path)
        if os.path.exists(path):
            return True
    except Exception as blob_err:
        logger.error(f"Failed to download {ext} from Azure Blob Storage: {blob_err}")
        
    # Tự động biên dịch lại từ JSON nếu file PDF/Word bị thiếu
    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_goc}.json")
    if not os.path.exists(p_json):
        try:
            from dich_vu.azure_blob import download_from_blob
            download_from_blob(f"json/{ma_goc}.json", p_json)
        except Exception as json_err:
            logger.error(f"Failed to download JSON from Azure Blob Storage: {json_err}")
            
    # Phục hồi từ DB HTML nếu JSON vẫn bị thiếu
    if not os.path.exists(p_json):
        try:
            from mo_hinh import LichSuGiaoTrinh
            ls = LichSuGiaoTrinh.query.filter(LichSuGiaoTrinh.duong_dan_file.contains(ma_goc)).first()
            if ls and ls.noi_dung_html:
                logger.info(f"JSON missing for {ma_goc}. Reconstructing ket_qua from DB html_content...")
                ket_qua = khoi_phuc_ket_qua_tu_html(ls.noi_dung_html, tieu_de)
                # Lưu JSON phục hồi để tránh parse lại
                os.makedirs(os.path.dirname(p_json), exist_ok=True)
                with open(p_json, 'w', encoding='utf-8') as f:
                    import json
                    json.dump(ket_qua, f, ensure_ascii=False, indent=4)
                # Upload lên Blob Storage để sao lưu
                try:
                    from dich_vu.azure_blob import upload_to_blob
                    upload_to_blob(p_json, f"json/{ma_goc}.json")
                except Exception as upload_err:
                    logger.error(f"Failed to upload reconstructed JSON to Azure Blob Storage: {upload_err}")
        except Exception as db_err:
            logger.error(f"Failed to reconstruct ket_qua from DB HTML: {db_err}")

    if os.path.exists(p_json):
        try:
            import json
            with open(p_json, 'r', encoding='utf-8') as f:
                ket_qua = json.load(f)
                
            if is_plain:
                import copy
                def strip_citations(text):
                    if not text: return ""
                    clean = re.sub(r'<span class="citation-apa">.*?</span>', '', text)
                    clean = re.sub(r'<sup class="citation">.*?</sup>', '', clean)
                    clean = re.sub(r'\[\w+\]', '', clean)
                    return clean

                book_export = ket_qua.get("book_vi", {})
                book_plain = copy.deepcopy(book_export)
                for chap in book_plain.get("chapters", []):
                    for sec in chap.get("sections", []):
                        sec["content"] = strip_citations(sec.get("content", ""))
                        sec["citations"] = []

                ket_qua_data = {
                    "topic": ket_qua.get("topic", tieu_de), 
                    "book_vi": book_plain, 
                    "references": [],
                    "terms": ket_qua.get("terms", []),
                    "is_custom_flow": ket_qua.get("is_custom_flow", False)
                }
            else:
                ket_qua_data = ket_qua
                
            os.makedirs(os.path.dirname(path), exist_ok=True)
            
            if ext == "pdf":
                from dich_vu.xuat_tai_lieu.xuat_pdf import xuat_pdf
                xuat_pdf(ket_qua_data, path)
            else:
                from dich_vu.xuat_tai_lieu.xuat_docx import xuat_docx
                xuat_docx(ket_qua_data, path)
                
            # Upload lên Blob Storage để lưu trữ lâu dài
            try:
                from dich_vu.azure_blob import upload_to_blob
                upload_to_blob(path, f"{ext}/{ma_goc}{suffix}.{ext}")
            except Exception as upload_err:
                logger.error(f"Failed to upload regenerated file to Azure Blob Storage: {upload_err}")
                
            return os.path.exists(path)
        except Exception as regen_err:
            logger.error(f"Failed to regenerate {ext} from JSON: {regen_err}")
            
    return False

@app.get("/tai/<loai>/<ma>")
def tai_file(loai, ma):
    # Hỗ trợ phiên bản 'plain' bằng cách bóc tách hậu tố để tra cứu info (V23.5.2)
    is_plain = "_plain" in ma
    ma_goc = ma.replace("_plain", "")
    if not kiem_tra_quyen_tai(ma_goc):
        return "Bạn không có quyền tải xuống giáo trình này.", 403
        
    info = CONG_VIEC.get(ma_goc)
    tieu_de = info.get("tieu_de", "giao_trinh") if info else "giao_trinh"

    if not info:
        from mo_hinh import LichSuGiaoTrinh
        ls = LichSuGiaoTrinh.query.filter(LichSuGiaoTrinh.duong_dan_file.contains(ma_goc)).first()
        if ls:
            tieu_de = ls.chu_de
        else:
            p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_goc}.json")
            if os.path.exists(p_json):
                try:
                    import json
                    with open(p_json, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        tieu_de = data.get("topic", "giao_trinh")
                except:
                    pass

    ext = "pdf" if loai == "pdf" else "docx"
    
    if dam_bao_file_ton_tai(ma_goc, ext, is_plain, tieu_de):
        folder = CauHinh.THU_MUC_PDF if loai == "pdf" else CauHinh.THU_MUC_DOCX
        path = os.path.join(folder, f"{ma}.{ext}")
        filename = f"{sanitize_filename(tieu_de)}.{ext}"
        return send_file(os.path.abspath(path), as_attachment=True, download_name=filename)
    
    return "File not found on server", 404

@app.get("/tai/zip/<ma>")
def tai_zip(ma):
    """Tải cả Word + PDF trong 1 file ZIP (V32 - Bundle Export)."""
    is_plain = "_plain" in ma
    ma_goc = ma.replace("_plain", "")
    if not kiem_tra_quyen_tai(ma_goc):
        return "Bạn không có quyền tải xuống giáo trình này.", 403
        
    info = CONG_VIEC.get(ma_goc)
    tieu_de = info.get("tieu_de", "giao_trinh") if info else "giao_trinh"

    if not info:
        from mo_hinh import LichSuGiaoTrinh
        ls = LichSuGiaoTrinh.query.filter(LichSuGiaoTrinh.duong_dan_file.contains(ma_goc)).first()
        if ls:
            tieu_de = ls.chu_de
        else:
            p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_goc}.json")
            if os.path.exists(p_json):
                try:
                    import json
                    with open(p_json, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                        tieu_de = data.get("topic", "giao_trinh")
                except:
                    pass

    tieu_de = sanitize_filename(tieu_de)
    suffix = "_plain" if is_plain else ""
    
    # Đảm bảo cả 2 file Word và PDF đều tồn tại
    has_docx = dam_bao_file_ton_tai(ma_goc, "docx", is_plain, tieu_de)
    has_pdf = dam_bao_file_ton_tai(ma_goc, "pdf", is_plain, tieu_de)
    
    p_docx = os.path.join(CauHinh.THU_MUC_DOCX, f"{ma_goc}{suffix}.docx")
    p_pdf = os.path.join(CauHinh.THU_MUC_PDF, f"{ma_goc}{suffix}.pdf")
    
    if not has_docx and not has_pdf:
        return "Files not found on server", 404
    
    # Tạo ZIP trong bộ nhớ (RAM) để tránh ghi đĩa
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        if os.path.exists(p_docx):
            zf.write(p_docx, f"{tieu_de}.docx")
        if os.path.exists(p_pdf):
            zf.write(p_pdf, f"{tieu_de}.pdf")
    
    zip_buffer.seek(0)
    
    label = "ban_sach" if is_plain else "trich_dan"
    zip_filename = f"{tieu_de}_{label}.zip"
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )

@app.get("/ket_qua/<ma_cv>")
def ket_qua(ma_cv):
    thong_tin = CONG_VIEC.get(ma_cv)
    if not thong_tin: return "Not found", 404
    if thong_tin.get("trang_thai") == "hoan_thanh":
        p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma_cv}.json")
        if os.path.exists(p_json):
            with open(p_json, 'r', encoding='utf-8') as f:
                data = json.load(f)
            return render_template("result.html", ma_cv=ma_cv, thong_tin=thong_tin, book=data.get('book_vi', {}), references=data.get('references', []), glossary=data.get('glossary', []), grounding=data.get('grounding', {}), extracted_terms=data.get('extracted_terms', []), kb_headings=data.get('kb_headings', []))
    return render_template("result.html", ma_cv=ma_cv, thong_tin=thong_tin)

# --- V33: EXPORT RIÊNG GLOSSARY & SUMMARY ---
@app.get("/tai/glossary/<ma>")
def tai_glossary(ma):
    """Xuất file DOCX chỉ chứa Bảng thuật ngữ."""
    ma_goc = ma.replace("_plain", "")
    if not kiem_tra_quyen_tai(ma_goc):
        return "Bạn không có quyền tải xuống giáo trình này.", 403
        
    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma}.json")
    if not os.path.exists(p_json):
        try:
            from dich_vu.azure_blob import download_from_blob
            download_from_blob(f"json/{ma}.json", p_json)
        except Exception as blob_err:
            logger.error(f"Failed to download JSON file for glossary from Azure Blob Storage: {blob_err}")

    if not os.path.exists(p_json): return "JSON not found", 404

    with open(p_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    glossary = data.get("glossary", [])
    if not glossary: return "No glossary data", 404

    # Lấy tên chủ đề từ CONG_VIEC hoặc từ JSON
    info = CONG_VIEC.get(ma, {})
    tieu_de = info.get("tieu_de") or data.get("topic", "Giáo trình")

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    for _ in range(3): doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("BẢNG THUẬT NGỮ")
    run_t.font.size = Pt(20)
    run_t.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(tieu_de)
    run_s.font.size = Pt(14)
    run_s.italic = True

    doc.add_page_break()

    # Glossary entries
    for item in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run_term = p.add_run(f"{item.get('term', '')}: ")
        run_term.font.size = Pt(13)
        run_term.bold = True
        run_def = p.add_run(item.get('definition', ''))
        run_def.font.size = Pt(13)

    # Save to BytesIO
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name=f"{sanitize_filename(tieu_de)} - Bảng thuật ngữ.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.get("/tai/summary/<ma>")
def tai_summary(ma):
    """Xuất file DOCX chỉ chứa Tóm tắt các chương."""
    ma_goc = ma.replace("_plain", "")
    if not kiem_tra_quyen_tai(ma_goc):
        return "Bạn không có quyền tải xuống giáo trình này.", 403
        
    p_json = os.path.join(CauHinh.THU_MUC_JSON, f"{ma}.json")
    if not os.path.exists(p_json):
        try:
            from dich_vu.azure_blob import download_from_blob
            download_from_blob(f"json/{ma}.json", p_json)
        except Exception as blob_err:
            logger.error(f"Failed to download JSON file for summary from Azure Blob Storage: {blob_err}")

    if not os.path.exists(p_json): return "JSON not found", 404

    with open(p_json, 'r', encoding='utf-8') as f:
        data = json.load(f)

    book = data.get("book_vi", {})
    chapters = book.get("chapters", [])
    has_any = any(ch.get("summary") for ch in chapters)
    if not has_any: return "No summary data", 404

    # Lấy tên chủ đề từ CONG_VIEC hoặc từ JSON
    info = CONG_VIEC.get(ma, {})
    tieu_de = info.get("tieu_de") or data.get("topic", "Giáo trình")

    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    for _ in range(3): doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run("TÓM TẮT CÁC CHƯƠNG")
    run_t.font.size = Pt(20)
    run_t.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run(tieu_de)
    run_s.font.size = Pt(14)
    run_s.italic = True

    doc.add_page_break()

    # Chapter summaries
    for idx, ch in enumerate(chapters, 1):
        h = doc.add_heading(f"Chương {idx}: {ch.get('title', '')}", level=2)
        summary = ch.get("summary", "")
        if summary:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(summary)
            run.font.size = Pt(13)
        else:
            p = doc.add_paragraph("Chưa có tóm tắt cho chương này.")
            p.runs[0].italic = True

    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True,
                     download_name=f"{sanitize_filename(tieu_de)} - Tóm tắt chương.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")



def cleanup_old_jobs():
    """
    Tự động dọn dẹp các Job cũ để tránh chiếm dụng RAM vô thời hạn (V18.3).
    Chạy định kỳ mỗi 15 phút, xóa job đã xong/thất bại nếu danh sách quá dài (>50).
    """
    while True:
        try:
            time.sleep(900) # 15 minutes
            to_delete = []
            for ma_cv, info in list(CONG_VIEC.items()):
                if info.get("trang_thai") in ["hoan_thanh", "that_bai"]:
                    to_delete.append(ma_cv)
            
            # Chỉ cleanup nếu tổng số job vượt ngưỡng an toàn
            if len(CONG_VIEC) > 50:
                for ma_cv in to_delete[:20]:
                    CONG_VIEC.pop(ma_cv, None)
                if to_delete:
                    logger.info(f"[Cleanup] Pruned old jobs. Current count: {len(CONG_VIEC)}")
        except Exception as e:
            logger.error(f"[Cleanup Error] {e}")

# Khởi chạy thread dọn dẹp ngầm
threading.Thread(target=cleanup_old_jobs, daemon=True).start()

def migrate_database_ngay_tao():
    from mo_hinh import db_type, NguoiDung
    import datetime
    
    if db_type == "mongodb":
        try:
            import pymongo
            mongo_uri = os.getenv("MONGO_URI", "mongodb://localhost:27017/giao_trinh_ai")
            parsed = pymongo.uri_parser.parse_uri(mongo_uri)
            db_name = parsed.get("database") or "giao_trinh_ai"
            client = pymongo.MongoClient(mongo_uri)
            col = client[db_name]["nguoi_dung"]
            default_date = datetime.datetime(2026, 6, 15, 0, 0, 0)
            res = col.update_many({"ngay_tao": {"$exists": False}}, {"$set": {"ngay_tao": default_date}})
            logger.info(f"MongoDB: Updated missing ngay_tao for nguoi_dung table. Modified: {res.modified_count}")
        except Exception as e:
            logger.error(f"MongoDB migration error: {e}")
    else:
        try:
            from sqlalchemy import text
            db.session.execute(text("SELECT ngay_tao FROM nguoi_dung LIMIT 1"))
        except Exception:
            db.session.rollback()
            try:
                db.session.execute(text("ALTER TABLE nguoi_dung ADD COLUMN ngay_tao DATETIME DEFAULT '2026-06-15 00:00:00'"))
                db.session.commit()
                logger.info("SQL database: Added column ngay_tao to nguoi_dung table.")
            except Exception as alter_err:
                db.session.rollback()
                logger.error(f"SQL migration error (ALTER TABLE): {alter_err}")

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
        migrate_database_ngay_tao()
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)), debug=True, use_reloader=False)
